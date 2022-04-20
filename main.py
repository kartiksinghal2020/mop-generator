from docx.enum.text import WD_COLOR_INDEX, WD_UNDERLINE
from bs4 import BeautifulSoup
import requests
from docx import *
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import sys
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.dml import MSO_THEME_COLOR_INDEX

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import inspect, os


document = Document()

#Change style of main headings
font = document.styles['Heading 1'].font
font.name = 'Times New Roman'
font.size = docx.shared.Pt(22)
font.underline = WD_UNDERLINE.THICK

font = document.styles['Heading 2'].font
font.name = 'Times New Roman'
font.size = docx.shared.Pt(16)

class Formatter:
    def __init__(self):
        pass
    @classmethod
    def funTable(self, content):
        f = Formatter()
        table = document.add_table(rows = 1, cols = 1, style = 'Table Grid')
        row = table.rows[0].cells
        row[0].text =content
        for row in table.rows:
            for cell in row.cells:
                shading_elm_2 = parse_xml(r'<w:shd {} w:fill="a4a5a6"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm_2)
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Courier New'
                        font.color.rgb = RGBColor(0, 0, 0)
                        
                            
        return 
    @classmethod
    def add_hyperlink(self,paragraph, text, url):
       
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        r = paragraph.add_run ()
        r._r.append (hyperlink)
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        return hyperlink
    # Add the Page number in the file
    def create_element(name):
        return OxmlElement(name)
    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)
    @classmethod
    def add_page_number(self,run):
        fldStart = self.create_element('w:fldChar')
        self.create_attribute(fldStart, 'w:fldCharType', 'begin')
        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"
        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'separate')
        fldChar2 = self.create_element('w:t')
        fldChar2.text = "2"
        fldEnd = self.create_element('w:fldChar')
        self.create_attribute(fldEnd, 'w:fldCharType', 'end')
        run._r.append(fldStart)
        run._r.append(instrText)
        run._r.append(fldChar1)
        run._r.append(fldChar2)
        run._r.append(fldEnd)
    @classmethod
    def findURL(self,string):
        regex=r"(?i)\b((?:https?://|www\d{0,3}[.]|[a-z0-9.\-]+[.][a-z]{2,4}/)(?:[^\s()<>]+|\(([^\s()<>]+|(\([^\s()<>]+\)))*\))+(?:\(([^\s()<>]+|(\([^\s()<>]+\)))*\)|[^\s`!()\[\]{};:'\".,<>?«»“”‘’]))"
        if(re.search(regex, string)):
            return True
        else:
            return False
    
allProducts = {
    "Platforms": {
        "Asr1000": [
            {
                "Devicename": "ASR 1000",
                "version": "17.03.03",
                "typeofdevice": "Router",
                "os": "IOS XE"
            },
            {
                "Devicename": "ASR 1000",
                "version": "17.03.02",
                "typeofdevice": "Router",
                "os": "IOS XE"
            }
        ],
        "Switch Cat 9600": [
            {
                "Devicename": "Switch Catalyst 9600",
                "version":"17.06.02",
                "typeofdevice": "Switch",
                "os": "IOS XE"
            },
            {
                "Devicename": "Switch Catalyst 9600",
                "version": "17.03.05",
                "typeofdevice": "Switch",
                "os": "IOS XE"
            },
        ],
        "Switch Cat 9300":[
            {
                "Devicename": "Switch Catalyst 9300",
                "version": "17.03.04",
                "typeofdevice": "Switch",
                "os": "IOS XE"
            },
        ],
        "Switch Cat 9500":[
            {
                "Devicename": "Switch Catalyst 9500",
                "version": "17.03.04",
                "typeofdevice": "Switch",
                "os": "IOS XE"
            }
        ],
        "Nexus 5000":[
            {
                "Devicename": "Nexus 5548",
                "version": "7.3(6)N1(1)",
                "typeofdevice": "Switch",
                "os": "NX-OS"
            }
        ]
    },
    "class_list": ["RiskAnalysis","HighLevelProjectOverview","TimeScales","ResourceRequirements","Responsibilities","Excalation","SuccessCriteria","Prerequisites","PreUpgradeTasks","ImageUpload","CheckIncompatibilityOfSystemImage","IdentifyTheUpgradeImpact","UserPassword","HwStateVerification","ServicesandIncidentVerification","PreUpgradeChecksCapture","UpgradeProcedure","RommonUpgrade","UpgradeSwitch","ConfigurationRegister","UpgradeInInstallMode","RelatedDocuments","LimitationsAndRestrictions", "TargetProcessSteps","InstallAnalysis","OsUpgrade","RollbackProcedure","DiffReport","TimeTaken","BackupCurrentConfiguration", "ImageDeploymentAndValidation","IssuUpgrade", "InstallAndReload", "PreInstallVerification", "AvailableDram", "IosVersion", "AcronymListing"],
    
}

#class Contents
class Contents:
    def __init__(self):
        pass
    @classmethod
    def contents(self,headings):
        document.add_heading(headings[0],1)
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need
    
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)
    
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
    
        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)
        p_element = paragraph._p
        document.add_page_break()
        return 'a'

class About:
    def __init__(self):
        pass
    @classmethod
    def about(self,author):
        table = document.add_table(rows=0, cols=2)
        row=table.add_row().cells
        p=row[0].add_paragraph('Author')
        p = row[1].add_paragraph(f'{author}')
        p.alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
        p=row[0].add_paragraph('Change Authority ')
        p = row[1].add_paragraph('Cisco Systems Customer Experience')
        p.alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
        p=row[0].add_paragraph('DCP Reference')
        p.alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
        p=row[0].add_paragraph('Project ID ')
        p.alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
        return 'a'

#class for History
class History:
    def __init__(self):
        pass
    @classmethod
    def history(self):
        table = document.add_table(rows=3, cols=4, style = 'Colorful Grid Accent 1')
        v = 'a'
        row = table.rows[0].cells
        row[0].text = 'Version No.'
        row[1].text = 'Issue Date'
        row[2].text = 'Status'
        row[3].text = 'Reason for Change'
        return v

#Review Class
class Review:
    def __init__(self):
        pass
    @classmethod
    def review(self):
        table = document.add_table(rows=3, cols=3, style = 'Colorful Grid Accent 1')
        v = 'a'
        row = table.rows[0].cells
        row[0].text = "Reviewer's Details"
        row[1].text = 'Version No.'
        row[2].text = 'Date'
        return v
        
class DocumentConventions:
    def __init__(self):
        pass
    @classmethod
    def documentConventions(self):
        try:
            r = document.add_paragraph().add_run()
            r.add_picture('caution.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(' Caution—Alerts readers to be careful. In this situation, you might do something that could  result in equipment damage or loss of data.')
            r = document.add_paragraph().add_run()
            r.add_picture('note.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(' Note—Alerts readers to take note. Notes contain helpful suggestions or references to  material not covered in the document.')
            r = document.add_paragraph().add_run()
            r.add_picture('timesaver.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(' Timesaver—Alerts the reader that they can save time by performing the action described in the paragraph affixed to this icon.')
            r = document.add_paragraph().add_run()
            r.add_picture('alert.png',width=Inches(0.3), height=Inches(.3))
            r.add_text(' Tip—Alerts the reader that the information affixed to this icon will help them solve a problem. The information might not be troubleshooting or even an action, but it could be useful information similar to a Timesaver.')
            r = document.add_paragraph().add_run()
            r.add_picture('warning.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(' Warning—Alerts readers of a situation that could cause bodily injury. They need to be aware of the hazards involved with electrical circuitry and familiarize themselves with standard practices for preventing accidents.')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found...Pleae refer description of this script.") 
        return 
class Preface:
    def __init__(self):
        pass
    @classmethod
    def preface(self,devicename,pid,typeofdevice):
        if typeofdevice == "Switch":
            return f"The purpose of this document is to define a standard process for the upgrade of software on Cisco {devicename.replace('Switch','')} Series {typeofdevice}es and to identify areas where the process may benefit from automation."
        else:
            return f"The purpose of this document is to define a standard process for the upgrade of software on Cisco {devicename} Series {typeofdevice}s and to identify areas where the process may benefit from automation."
        
class Audience:
    def __init__(self):
        pass
    @classmethod
    def audience(self,cname):
        return f"This document is intended to support {cname} who will plan, configure & install software upgrades."

class Scope:
    def __init__(self):
        pass
    @classmethod
    def scope(self, cname, devicename, version, os):
        return f'The purpose of this document is to provide to {cname} engineers/personnel the technical steps required to upgrade stacked and/or standalone {devicename} to {os} version {version}.'

class TargetProcessSteps:
    def __init__(self,os,version):
        self.os =os
        self.version = version
    def targetProcessSteps_ASR_1000_17_03_03(self):
        document.add_paragraph("The process has two main stages: Preparation and Upgrade. Each of these stages is further divided into several steps: these steps are described below in detail.")
        document.add_paragraph("Preparation", style='Heading 3')
        document.add_paragraph("Install Analysis", style='List Bullet')
        document.add_paragraph("Image Deployment and Validation", style='List Bullet')
        document.add_paragraph("Upgrade", style='Heading 3')
        document.add_paragraph("Pre-install Verification", style='List Bullet')
        document.add_paragraph("Install", style='List Bullet')
        document.add_paragraph("Post-install Verification", style='List Bullet')
        return
    def targetProcessSteps_Switch_Catalyst_9600_17_06_02(self):
        document.add_paragraph("The process has two main stages: Preparation and Upgrade. Each of these stages is further divided into several steps: these steps are described below in detail.")
        document.add_paragraph("Preparation", style='Heading 3')
        document.add_paragraph("Install Analysis", style='List Bullet')
        document.add_paragraph("Image Deployment and Validation", style='List Bullet')
        document.add_paragraph("Upgrade", style='Heading 3')
        document.add_paragraph("Pre-install Verification", style='List Bullet')
        document.add_paragraph("Install", style='List Bullet')
        document.add_paragraph("Post-install Verification", style='List Bullet')
        return
    def targetProcessSteps_Switch_Catalyst_9600_17_03_05(self):
        document.add_paragraph("The process has two main stages: Preparation and Upgrade. Each of these stages is further divided into several steps: these steps are described below in detail.")
        document.add_paragraph("Preparation", style='Heading 3')
        document.add_paragraph("Install Analysis", style='List Bullet')
        document.add_paragraph("Image Deployment and Validation", style='List Bullet')
        document.add_paragraph("Upgrade", style='Heading 3')
        document.add_paragraph("Pre-install Verification", style='List Bullet')
        document.add_paragraph("Install", style='List Bullet')
        document.add_paragraph("Post-install Verification", style='List Bullet')
        return
    def targetProcessSteps_Switch_Catalyst_9300_17_03_04(self):
        document.add_paragraph("The process has two main stages: Preparation and Upgrade. Each of these stages is further divided into several steps: these steps are described below in detail.")
        document.add_paragraph("Preparation", style='Heading 3')
        document.add_paragraph("Install Analysis", style='List Bullet')
        document.add_paragraph("Image Deployment and Validation", style='List Bullet')
        document.add_paragraph("Upgrade", style='Heading 3')
        document.add_paragraph("Pre-install Verification", style='List Bullet')
        document.add_paragraph("Install", style='List Bullet')
        document.add_paragraph("Post-install Verification", style='List Bullet')
        return
    def targetProcessSteps_Switch_Catalyst_9500_17_03_04(self):
        document.add_paragraph("The process has two main stages: Preparation and Upgrade. Each of these stages is further divided into several steps: these steps are described below in detail.")
        document.add_paragraph("Preparation", style='Heading 3')
        document.add_paragraph("Install Analysis", style='List Bullet')
        document.add_paragraph("Image Deployment and Validation", style='List Bullet')
        document.add_paragraph("Upgrade", style='Heading 3')
        document.add_paragraph("Pre-install Verification", style='List Bullet')
        document.add_paragraph("Install", style='List Bullet')
        document.add_paragraph("Post-install Verification", style='List Bullet')
        return
class Assumptions:
    def __init__(self):
        pass
    @classmethod
    def assumptions(self,cname):
        return f"The information in this document is based upon data collected through meetings, documentation, and data provided by {cname}. Any Cisco hardware and/or software information in this document is based upon current performance estimates and feature capabilities."


class RelatedDocuments(Formatter):
    def __init__(self,os,version, rd):
        self.os = os
        self.version = version
        self.rd = rd
    def addition_content(self):
        for i in self.rd:
            if RelatedDocuments.findURL(i[1]):
                p = document.add_paragraph(i[0]+'\n',style='List Bullet')
                RelatedDocuments.add_hyperlink(p,i[1],i[1])
            else:
                document.add_paragraph(i[0],style='List Bullet')
                document.add_paragraph(i[1])
        return 
    def relatedDocuments_ASR_1000_17_03_03(self):
        document.add_paragraph("This document should be read in association with the relevant publicly available documentation for these devices, including release notes and upgrade guides for the specific version of software being deployed. \n\nThese are available at the following location:")
        p = document.add_paragraph('ASR1000 release notes 17.3.3 \n',style = 'List Bullet')
        RelatedDocuments.add_hyperlink(p,"https://www.cisco.com/c/en/us/td/docs/routers/asr1000/release/notes/xe-17-3/asr1000-rel-notes-xe-17-3.html","https://www.cisco.com/c/en/us/td/docs/routers/asr1000/release/notes/xe-17-3/asr1000-rel-notes-xe-17-3.html")

        p = document.add_paragraph('ASR1000 Command Reference \n',style = 'List Bullet')
        RelatedDocuments.add_hyperlink(p,"https://www.cisco.com/c/en/us/td/docs/ios/fundamentals/command/reference/cf_book.html","https://www.cisco.com/c/en/us/td/docs/routers/asr1000/install/guide/1001-x/asr1hig-book/pwr_up_init_configuartion.html?referring_site=RE&pos=1&page=https://www.cisco.com/c/en/us/td/docs/routers/asr1000/quick/start/guide/asr1_qs1.html")

        p = document.add_paragraph('ASR1000 Rommon upgrade guide \n',style = 'List Bullet')
        RelatedDocuments.add_hyperlink(p,"https://www.cisco.com/c/en/us/td/docs/routers/asr1000/rommon/asr1000-rommon-upg-guide.html#con_46405","https://www.cisco.com/c/en/us/td/docs/routers/asr1000/rommon/asr1000-rommon-upg-guide.html#con_46405")

        p = document.add_paragraph('ASR1000 Issu upgrade guide\n',style = 'List Bullet')
        RelatedDocuments.add_hyperlink(p,"https://www.cisco.com/c/en/us/td/docs/routers/asr1000/configuration/guide/chassis/asr1000-software-config-guide/issu-asr.html?referring_site=RE&pos=2&page=https://www.cisco.com/c/en/us/support/routers/asr-1000-series-aggregation-services-routers/products-installation-guides-list.html","https://www.cisco.com/c/en/us/td/docs/routers/asr1000/configuration/guide/chassis/asr1000-software-config-guide/issu-asr.html?referring_site=RE&pos=2&page=https://www.cisco.com/c/en/us/support/routers/asr-1000-series-aggregation-services-routers/products-installation-guides-list.html")
        for i in self.rd:
            if RelatedDocuments.findURL(i[1]):
                p = document.add_paragraph(i[0]+'\n',style='List Bullet')
                RelatedDocuments.add_hyperlink(p,i[1],i[1])
            else:
                document.add_paragraph(i[0],style='List Bullet')
                document.add_paragraph(i[1])
        document.add_page_break()
        return
    def relatedDocuments_Switch_Catalyst_9600_17_06_02(self):
        p=document.add_paragraph('Cisco Catalyst 9600 Series Switches, Cisco IOS XE Bengaluru 17.6.x\n',style='List Bullet')
        RelatedDocuments.add_hyperlink(p,"https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-6/release_notes/ol-17-6-9600.html#concept_ycv_jdf_3mb","https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-6/release_notes/ol-17-6-9600.html#concept_ycv_jdf_3mb")
        p = document.add_paragraph('Cisco Catalyst 9600 Series Switches, Command Reference\n',style='List Bullet')
        RelatedDocuments.add_hyperlink(p,"https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-6/command_reference/b_176_9600_cr/system_management_commands.html#wp2862294214","https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-6/command_reference/b_176_9600_cr/system_management_commands.html#wp2862294214")
        for i in self.rd:
            if RelatedDocuments.findURL(i[1]):
                p = document.add_paragraph(i[0]+'\n',style='List Bullet')
                RelatedDocuments.add_hyperlink(p,i[1],i[1])
            else:
                document.add_paragraph(i[0],style='List Bullet')
                document.add_paragraph("\t"+i[1])
        document.add_page_break()
        return
    def relatedDocuments_Switch_Catalyst_9600_17_03_05(self):
        p=document.add_paragraph('Cisco Catalyst 9600 Series Switches Hardware Installation Guide\n',style='List Bullet')
        RelatedDocuments.add_hyperlink(p,"https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/hardware/install/b_9600_hig.html","https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/hardware/install/b_9600_hig.html")
        for i in self.rd:
            if RelatedDocuments.findURL(i[1]):
                p = document.add_paragraph(i[0]+'\n',style='List Bullet')
                RelatedDocuments.add_hyperlink(p,i[1],i[1])
            else:
                document.add_paragraph(i[0],style='List Bullet')
                document.add_paragraph("\t"+i[1])
        document.add_page_break()
        return
    def relatedDocuments_Switch_Catalyst_9300_17_03_04(self):
        p=document.add_paragraph('Release Notes for Cisco Catalyst 9300 Series Switches, Cisco IOS XE Amsterdam 17.3.x ',style='List Bullet')
        RelatedDocuments.add_hyperlink(p,'https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9300/software/release/17-3/release_notes/ol-17-3-9300.html','https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9300/software/release/17-3/release_notes/ol-17-3-9300.html')
        p=document.add_paragraph('Cisco Catalyst 9300 Series Switches Hardware Installation Guide\n',style='List Bullet')
        RelatedDocuments.add_hyperlink(p,"https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9300/hardware/install/b_c9300_hig.html","https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9300/hardware/install/b_c9300_hig.html")
        for i in self.rd:
            if RelatedDocuments.findURL(i[1]):
                p = document.add_paragraph(i[0]+'\n',style='List Bullet')
                RelatedDocuments.add_hyperlink(p,i[1],i[1])
            else:
                document.add_paragraph(i[0],style='List Bullet')
                document.add_paragraph("\t"+i[1])
        document.add_page_break()
        return
    def relatedDocuments_Switch_Catalyst_9500_17_03_04(self):
        p=document.add_paragraph('Release Notes for Cisco Catalyst 9500 Series Switches, Cisco IOS XE Amsterdam 17.3.x ',style='List Bullet')
        RelatedDocuments.add_hyperlink(p,'https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9500/software/release/17-3/release_notes/ol-17-3-9500.html','https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9500/software/release/17-3/release_notes/ol-17-3-9500.html')
        p=document.add_paragraph('Cisco Catalyst 9500 Series Switches Hardware Installation Guide\n',style='List Bullet')
        RelatedDocuments.add_hyperlink(p,"https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9500/hardware/install/b_catalyst_9500_hig.html","https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9500/hardware/install/b_catalyst_9500_hig.html")
        for i in self.rd:
            if RelatedDocuments.findURL(i[1]):
                p = document.add_paragraph(i[0]+'\n',style='List Bullet')
                RelatedDocuments.add_hyperlink(p,i[1],i[1])
            else:
                document.add_paragraph(i[0],style='List Bullet')
                document.add_paragraph("\t"+i[1])
        document.add_page_break()
        return
    def relatedDocuments_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph('This document should be read in association with the relevant publicly available documentation for Cisco NX-OS, including release notes and upgrade guides for the specific version of NX-OS being deployed. Please see attached for the relevant documentation for the version being targeted upon upgrade.')
        p=document.add_paragraph('\n\nN5K upgrade\n')
        RelatedDocuments.add_hyperlink(p,'https://www.cisco.com/c/en/us/td/docs/switches/datacenter/nexus5000/sw/upgrade/503_N1_1/n5000_upgrade_downgrade_503_n1_1.html','https://www.cisco.com/c/en/us/td/docs/switches/datacenter/nexus5000/sw/upgrade/503_N1_1/n5000_upgrade_downgrade_503_n1_1.html')
        document.add_paragraph('The below is the summary of the procedure mentioned in the above link and the same has been followed in this document:\n\n'
                               'A disruptive upgrade causes a switch and connected FEXs to reload. The time required for a FEX to '
                               'reload is less than the time required for a switch to reload. When hosts are connected to a dual-homed FEX, it is possible to keep the traffic disruption of the hosts to same time as required by FEX '
                               'to reload (approximately 120 seconds), instead of the time required for an upgrade of the entire '
                               'access layer.\n'
                               'The following figure shows a dual-homed FEX topology in which the access layer includes a vPC '
                               'configuration to hosts or downstream switches')
        p = document.add_paragraph()
        p.add_run('Note').bold = True
        p.add_run(' The following dual-homed FEX procedure is supported only for an upgrade and not for a downgrade.')
        p = document.add_paragraph('\n')
        p.add_run('Step 1:').bold = True
        p.add_run('Configure FEX module pre-provisioning for all the FEXs connected to both the '
                  'switches (vPC primary and vPC secondary switches).\n'
                  'Upgrade the vPC primary switch with the new image using the install all '
                  'kickstartimage systemimage command. During the upgrade process, the switch is '
                  'reloaded. When the switch is reloaded, only singled-homed FEXs connected to the switch '
                  'are reloaded and dual-homed FEXs are not reloaded. Servers connected to the dual-homed '
                  'FEXs retain network connectivity through the vPC secondary switch.')
        p = document.add_paragraph()
        p.add_run('Step 2:').bold = True
        p.add_run('Verify that the upgrade of the vPC primary switch is completed successfully. At the completion of the upgrade, the vPC primary switch restores vPC peering. However, dual-homed FEXs are connected only to the secondary vPC switch.')
        p = document.add_paragraph()
        p.add_run('Step 3:').bold = True
        p.add_run(' Reload the dual-homed FEXs using the reload fex command from the vPC secondary switch. Reload the FEXs one-by-one or in a bunch of two or three FEXs. The servers connected to the dual-homed FEXs will lose connectivity.')
        p = document.add_paragraph()
        p.add_run('Step 4:').bold = True
        p.add_run('Wait for the FEXs to reload. After the reload, the FEXs connect to the upgraded switch (vPC primary switch).')
        p = document.add_paragraph()
        p.add_run('Step 5:').bold = True
        p.add_run('Upgrade the vPC secondary switch with the new image using the install all kickstartimage systemimage command. During the upgrade process, the switch is reloaded. When the switch is reloaded, only singled-homed FEXs connected to the switch are reloaded and dual-homed FEXs are not reloaded')
        p = document.add_paragraph()
        p.add_run('Step 6:').bold = True
        p.add_run('Verify that the upgrade of the vPC secondary switch is completed successfully. At the completion of the upgrade, the vPC secondary switch restores vPC peering. Dual-homed FEXs connect to both the peer switches and start forwarding traffic.')
        document.add_paragraph("_______________________________________________________________\nThe releases we are upgrading from are impacted by below bug:")
        r = document.add_paragraph().add_run()
        r.add_text('\tCSCul22703 - N5K/6K: Upgrade between incompatible images can result in loss of config.')
        r.font.color.rgb = RGBColor(255, 0, 0)
        document.add_paragraph("\tWorkaround – FEX pre-provisioning which is explained in detail in the upgrade procedure.")
        document.add_paragraph("_______________________________________________________________")
        p= document.add_paragraph()
        p.add_run("Note: ").bold=True
        p.add_run("Issue “reload power-cycle” on required device/devices after both devices of the vpc-pair are upgraded to a particular release if you see the highlighted line:")
        RelatedDocuments.funTable("# reload power-cycle \n"
                                  "WARNING: There is unsaved configuration!!!\n"
                                  "WARNING: This command will reboot the system\n"
                                  "Do you want to continue? (y/n) [n] y\n"
                                  "[ 1750.957586] Shutdown Ports..")
        document.add_paragraph("Before the upgrade – “show install all impact” command Or"
                               "\tAfter the upgrade- “show install all status” command",style="List Bullet")
        RelatedDocuments.funTable("Images will be upgraded according to following table:\n"
                                  "Module Image Running-Version New-Version Upg-Required\n"
                                  "------ ---------- ---------------------- ---------------------- ------------\n"
                                  "     1     system           5.2(1)N1(1a)           5.2(1)N1(9b)          yes\n"
                                  "     1  kickstart           5.2(1)N1(1a)           5.2(1)N1(9b)          yes\n"
                                  "     1       bios     v3.6.0(05/09/2012)     v3.6.0(05/09/2012)           no\n"
                                  "     1     SFP-uC               v1.0.0.0               v1.0.0.0           no\n"
                                  "   100      fexth           5.2(1)N1(1a)           5.2(1)N1(9b)          yes\n"
                                  "     1  power-seq                   v1.0                   v3.0          yes\n"
                                  "     3  power-seq                   v5.0                   v5.0           no\n"
                                  "     1         uC               v1.2.0.1               v1.2.0.1           no")
        document.add_paragraph("Or During the install",style="List Bullet")
        RelatedDocuments.funTable("Do you want to continue with the installation (y/n)? [n] y\n"
                                  "Install is in progress, please wait.\n"
                                  "Performing runtime checks.\n"
                                  "[####################] 100% -- SUCCESS\n"
                                  "Setting boot variables.\n"
                                  "[####################] 100% -- SUCCESS\n"
                                  "Performing configuration copy.\n"
                                  "May 18, 2021 Software Upgrade Procedure: NX-OS Nexus devices\n"
                                  "Cisco Confidential. All printed copies and duplicate soft copies are considered uncontrolled\n"
                                  "And the original online version should be referred to for the latest version.\n"
                                  "Page 8 of 41\n"
                                  "[####################] 100% -- SUCCESS\n"
                                  "Module 1: Refreshing compact flash and upgrading bios/loader/bootrom/power-seq.\n"
                                  "Warning: please do not remove or power off the module at this time.\n"
                                  "Note: Power-seq upgrade needs a power-cycle to take into effect.\n"
                                  "On success of power-seq upgrade, SWITCH OFF THE POWER to the system and then, power \n"
                                  "it up.\n"
                                  "Note: Micro-controller upgrade needs a power-cycle to take into effect.\n"
                                  "On success of micro-controller upgrade, SWITCH OFF THE POWER to the system and \n"
                                  "then, power it up.\n"
                                  "[####################] 100% -- SUCCESS\n"
                                  "Pre-loading modules.\n"
                                  "[This step might take upto 20 minutes to complete - please wait.]\n"
                                  "[*Warning -- Please do not abort installation/reload or powercycle fexes*]\n"
                                  "[# ] 0%2017 Feb 1 10:24:17 5596-1 %SATCTRL-FEX102-2-\n"
                                  "SATCTRL_IMAGE: FEX102 Image update in progress. \n"
                                  "[#### ] 15%2017 Feb 1 10:29:53 5596-1 %SATCTRL-FEX102-2-\n"
                                  "SATCTRL_IMAGE: FEX102 Image update complete. \n"
                                  "[####################] 100% -- SUCCESS\n"
                                  "Finishing the upgrade, switch will reboot in 10 seconds.\n"
                                  "5596-1# 2017 Feb 1 10:30:28 5596-1 Feb 1 10:30:28 %KERN-0-SYSTEM_MSG: Shutdown \n"
                                  "Ports.. - kernel\n"
                                  "Broadcast message from root (Wed Feb 1 10:30:28 2017):\n"
                                  "The system is going down for reboot NOW!\n"
                                  "2017 Feb 1 10:30:28 5596-1 Feb 1 10:30:28 %KERN-0-SYSTEM_MSG: writing reset \n"
                                  "reason 49, - kernel\n"
                                  "______________________________________________________________________________")
        self.addition_content()
        return

#High Level Project Overview
class HighLevelProjectOverview:
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def highLevelProjectOverview_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph(f"The {cust} network worldwide consists of a number of globally dispersed data centers. The sites are currently considered legacy and run mostly on Cisco based DC switches, including the Nexus 5000 platforms.")
        return

#risk Analysis
class RiskAnalysis:
    def __init__(self,os,version):
        self.os = os
        self.version =version
    def riskAnalysis_Nexus_5548_7_3_6_N1_1_(self):
        document.add_heading('2.2.1 Risks of performing this upgrade',3)
        document.add_paragraph('It is recommended that all upgrade activities be completed in dedicated maintenance windows, '
                               f'agreed and understood within {cust}. During this time, no other planned or unplanned network or '
                               'configuration change should be anticipated without first verifying with the on-site and/or remote '
                               'team engineers. \n'
                               'During the preparation works for the upgrade, some risks that could prevent from upgrade were '
                               'identified. Those are listed and described below:')
        document.add_paragraph('Lack of detailed documentation. The lack of detailed Data Center network design '
                               'may lead to incorrect connectivity assessment and missing important steps in the' 
                               'execution plan and/or not fully understanding the existing caveats.',style='List Bullet 2')
        document.add_paragraph('Hardware health. The process of upgrading software is a stress to the individual'
                               'components of a network device. It is possible that power supplies, fans, fabric modules '
                               'or any other type of system card may fail during the process. It is recommended to have '
                               'a spare at all time. ',style='List Bullet 2')
        document.add_paragraph('Change window duration. It should always be ensured that the provided MW length '
                               'will accommodate for the time identified to be needed for the execution of the upgrade, '
                               'plus minimum 50% reserve',style='List Bullet 2')
        document.add_paragraph('Non-ISSU upgrade specifics. Due to the old running software and identified hardware '
                               'and software issues, an ISSU upgrade is either not possible or not recommended on '
                               'those devices, that means individual nodes would experience disruptive reboot, that '
                               'may affect inappropriately configured and/or connected endpoints',style='List Bullet 2')
        document.add_paragraph('During the software upgrade, the following situations can cause the installation to fail and should be prevented from happening:')
        document.add_paragraph('If the bootflash: file system does not have sufficient space to accept the updated image(s).',style="List Bullet 2")
        document.add_paragraph('If a module is removed while the upgrade is in progress.' ,style="List Bullet 2")
        document.add_paragraph('If the device has any power disruption while the upgrade is in progress.' ,style="List Bullet 2")
        document.add_paragraph('If the entire path for the remote server location is not specified accurately. ',style="List Bullet 2")
        document.add_paragraph('If images are incompatible after an upgrade. For example, an I/O module image may be incompatible with the system image. ',style="List Bullet 2")
        document.add_paragraph('If a Spanning Tree Protocol (STP) topology change occurs while the upgrade is in progress.',style="List Bullet 2")
        document.add_paragraph('If any drastic configuration and topology change occurs within the routing domain',style="List Bullet 2")
        
        document.add_heading('2.2.2 Risks of not performing this upgrade',3)
        document.add_paragraph(f'During the analysis of the current {cust} installed base it was identified that there is an enormous '
                               'software/hardware diversity in all environments, open to multiple bugs, scalability issues and PSIRT '
                               'advisories.\n\n'
                               f'Cisco, together with {cust} identified a process to harmonize the existing installed base and create '
                               'a robust mid- to long-term plan for both hardware and software. The upgrade of the existing Nexus '
                               '5000 devices to a more recent (NX-OS 7.3(6)N1(1)) software version is a medium impact process that' 
                               'will create a common starting point for the future target releases.\n\n'
                               'The selected software version was identified to be stable and fitting existing hardware and features in '
                               f'use at {cust} the best.')
        

        return

#Timescales
class TimeScales:
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def timeScales_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph('High level time scales can be found in the table below. The numbers are based on previous experience and can vary between environments')
        p=document.add_paragraph()
        p.add_run('Table 1 Migration Tasks timelines').italic = True
        data = (
        ("1","Check services and confirm no active incidents" ,"Pre Migration Window","15 mins"),
        ("2", "Configuration and License backup", "Pre Migration Window", "10 mins"),
        ("3", "Capture pre-check outputs from devices","Pre Migration Window", "15 mins"),
        ("4", "Execute software upgrade (each stage, disruptive)", "During Migration", "45 mins*4=180 mins"),
        ("6", "Post-upgrade verification and monitoring", "During Migration", "15 mins"),
        ("7" ,"(optional) Rollback", "During Migration","90 mins"),
        )
        table = document.add_table(rows=1, cols=4, style = 'Colorful Grid Accent 1')
        row = table.rows[0].cells
        row[0].text = "Task Number"
        row[1].text = 'Task Description'
        row[2].text = "Pre/During/Post Migration Window Activity"
        row[3].text = "Estimated time"
        for tn,td,pd,et in data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(tn)
            row_cells[1].text = str(td)
            row_cells[2].text = str(pd)
            row_cells[3].text = str(et)
        return

class ResourceRequirements():
    def __init__(self,os, version) -> None:
        self.os = os
        self.version = version
    def resourceRequirements_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph('The proposed team set up below serves as an example. Detailed needs will be provided separately for each maintenance window')
        p=document.add_paragraph()
        p.add_run('Table 2 Resource Requirement').italic= True
        table = document.add_table(rows=5, cols=4, style = 'Colorful Grid Accent 1')
        row = table.rows[0].cells
        row[0].text = 'Party'
        row[1].text = 'Role'
        row[2].text = 'Resource'
        row[3].text = 'Location'
        row = table.rows[1].cells
        row[1].text = 'NCCM Engineer'
        row[2].text = '1'
        row[3].text = 'Remote'
        row = table.rows[2].cells
        row[1].text = 'Change Manager'
        row[2].text = '1'
        row[3].text = 'Remote'
        row = table.rows[3].cells
        row[0].text = f'{cust}'
        row[1].text = 'Customer Regional SME'
        row[2].text = '1'
        row[3].text = 'Remote'
        row = table.rows[4].cells
        row[0].text = f'{cust}'
        row[1].text = 'Operations Team'
        row[2].text = 'Normal Team Staffing'
        row[3].text = 'On call/Onsite'
        a = table.cell(1, 0)
        b = table.cell(2, 0)
        A = a.merge(b)
        A.text = 'Cisco CX'

        return

class Responsibilities:
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def responsibilities_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph("Cisco, or their appointed Representatives, will be responsible for tasks designated with a Task Owner "
                               "of Cisco.\n\n"
                               f"{cust}, or their appointed Representatives, will provide resource to be responsible for tasks "
                               f"designated with a Task Owner of {cust}\n\n"
                               f"This document will be distributed to {cust} and the content should be agreed upon prior to any "
                               "party undertaking further work. Following this, additions will only be incorporated with agreement "
                               "from the Cisco Project Team. ")
        return

class Excalation:
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def excalation_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph("Should the need arise during the preparation or on-site migrationactivities, for technical issues related "
                               "to the upgrade, escalation through the Cisco TAC by raising support cases should be initiated. The "
                               "normal support escalation processes will then be invoked, as necessary.\n\n"
                               "For non-technical, commercial or other issues the Cisco escalation path will be to the Project "
                               f"Manager who will involve additional parties and invoke other Cisco and/or {cust} organizations, as "
                               "appropriate")
        return

class SuccessCriteria:
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def successCriteria_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph("The migrationwill be judged to have been successfully completed if the network is deemed fit to carry "
                               "live traffic and to be supported by Cisco TAC after its execution. If during or after the migration a "
                               "software or hardware defect is discovered that does not severely impact the operation of the network, "
                               "(for example any cosmetic bugs) the investigation and rectification of this defect will continue to be" 
                               "managed under the Cisco Support Process after the Cisco CXEngagement has been completed.\n\n"
                               "If any software or hardware defects are discovered that will severely impact the operation of the "
                               "network, during or after the migration, then any show stopping, or high severity issues will be "
                               "managed immediately under the Cisco TAC and support services.")
        return

class Prerequisites:
    def __init__(self,os,version):
        self.os =os
        self.version =version
    def prerequisites_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph('Following is a list of prerequisites for completing the procedure described in this document')
        document.add_paragraph("In-band/Out-of-Band and console access are available to all devices undergoing migration",style="List Bullet 2")
        document.add_paragraph("Console access must be available at any time for troubleshooting purposes",style="List Bullet 2")
        document.add_paragraph("Migration Team members have sufficient TACACS/RADIUS access to allow them to complete the migration activities.",style="List Bullet 2")
        document.add_paragraph("Equipment not discussed in the Equipment and Service Schedule section of the CR specific document will not be considered in the scope of this work.",style="List Bullet 2")
        document.add_paragraph(f"{cust} has agreed to provide suitable maintenance windows for upgrades. Such windows will be of a minimum of 4 hours or more, depending on the number of devices to be upgraded and the needed upgrade path.",style="List Bullet 2")
        document.add_paragraph(f"{cust} personnel will be available at the respective sites during pre-migration check and testing and migration execution",style="List Bullet 2")
        document.add_paragraph(f"{cust} on-call personnel from all the teams involved in the migration will be available at all sites during migration activities.",style="List Bullet 2")
        document.add_paragraph(f"{cust} will make staff with the ability to diagnose and identify network problems available during the Migration",style="List Bullet 2")
        document.add_paragraph(f"{cust} will inform Cisco immediately of any issues suspected to have arisen in the Migration or those that occur in the period following.",style="List Bullet 2")
        document.add_paragraph("A Cisco support contract is in place and active.",style="List Bullet 2")
        document.add_paragraph("A proactive TAC case has been opened and the contact details of the TAC case owner are known",style="List Bullet 2")
        document.add_paragraph("The Cisco project manager and technical leads must approve any deviation from these procedures that affect the Implementation and migration of the network.",style="List Bullet 2")
        document.add_paragraph(f"When all the procedures described herein are judged to have been completed and the Success Criteria met, {cust} will provide sign-off of the Upgrade Completion.")
        document.add_paragraph("Note: CLI output shown in this document is for reference purposes only and does not reflectactual production environment!")
        document.add_page_break()
        return

class PreUpgradeTasks:
    def __init__(self,os,version):
        self.os=os
        self.version=version
    def preUpgradeTasks_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph("Please follow the pre-upgrade steps prior to commencing into any upgrade activity(note all "
                               "commands should be ran on both switches of vPC-Pairto ensure no issues with upgrade). "
                               "Please investigate all issues and note for pre-checks. If something is impactful, we need to be able to "
                               f"alert {cust} team to remediate before proceeding.\n\n"
                               "The following steps need to be executed prior to the upgrade window.")
        return

class ImageUpload(Formatter):
    def __init__(self,os,version):
        self.os=os
        self.version=version
    def imageUpload_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph("The following images need to be acquired from software.cisco.com and uploaded to the device "
                              "boootflash. \n\n"
                              "Depending on the starting point for each device, only the next-in-queue releases will be" 
                              "uploaded.")
        p = document.add_paragraph()
        p.add_run("Table 3 Software images Nexus 5000").italic = True
        table = document.add_table(rows=3, cols=5, style = 'Colorful Grid Accent 1')
        row = table.rows[0].cells
        row[0].text = 'Platform name'
        row[1].text = 'Image version'
        row[2].text = 'Image file'
        row[3].text = 'Image size'
        row[4].text = 'MD5'
        row = table.rows[1].cells
        row[2].text = 'n5000-uk9.7.3.6.N1.1.bin'
        row[3].text = '320.09 MB (335634018 bytes)'
        row[4].text = 'bef58f988bcf4daef7896ac84c96c10d'
        row = table.rows[2].cells
        row[2].text = 'n5000-uk9-kickstart.7.3.6.N1.1.bin'
        row[3].text = '32.78 MB (34371584 bytes)'
        row[4].text = 'fb2f0bb05c0ddb035d4335a446f429af'
        a = table.cell(1, 0)
        b = table.cell(2, 0)
        A = a.merge(b)
        A.text = 'Nexus 5000'
        c = table.cell(1, 1)
        d = table.cell(2, 1)
        A = c.merge(d)
        A.text = '7.3(6)N1(1) '
        document.add_paragraph("Issue \"dir\" command to verify that sufficient space is available in the location where the images"
                               "will be copied to. This location includes the supervisor module bootflash: (internal to the device). "
                               "Internal bootflash: should have approximately 750 megabyte of free space available. Verify the "
                               "available space by issuing the commands below before each stage of upgrade:")
        try:
            r = document.add_paragraph().add_run()
            r.add_picture('note.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(f'Lines containing a hash sign (#) include commands that must be executed on the device. The hash sign itself is not part of the command sequence')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
        ImageUpload.funTable("# dir bootflash: | inc bytes\n"
                             " 4082638336 bytes used\n"
                             " 1733847552 bytes free\n"
                             " 4816485888 bytes total")
        document.add_paragraph("Verify OS Images on bootflash:")
        ImageUpload.funTable("# dir bootflash:\n"
                             " 4096       Jan 03 01:41:23 2009 .patch/\n"
                             " 98236      Apr 12 06:54:08 2020 20200411_005539_poap_3477_2.log\n"
                             " 1048680    Apr 11 03:12:06 2020 20200411_005539_poap_3477_init.log\n"
                             " 127598     Apr 12 07:14:21 2020 20200412_065736_poap_3485_init.log\n"
                             " 1436       Apr 30 11:23:45 2020 2d20c88a-90ab-4409-a7c2-652fece381a7.config\n"
                             " 3564       Apr 21 12:23:27 2020 86948514-4197-407d-9016-16b54054784c.config\n"
                             " 1453       Aug 16 10:49:48 2009 8d15d58c-515d-4501-b315-1cd091fdb7fe.config\n"
                             " 312        Feb 08 11:37:30 2009 SSI172904PX.lic\n"
                             " 312        Feb 08 11:22:54 2009 SSI172904PX_20190913042640731.lic\n"
                             " 264        Apr 20 12:58:26 2020 assoc_ascii_cnv.log\n"
                             " 5290       May 04 05:00:28 2020 mts.log\n"
                             " 31642624   Apr 04 00:26:07 2020 n5000-uk9-kickstart.5.2.1.N1.1a.bin\n"
                             " 31688192   May 04 06:03:51 2020 n5000-uk9-kickstart.5.2.1.N1.9b.bin\n"
                             " 37256704   May 04 07:21:54 2020 n5000-uk9-kickstart.7.0.8.N1.1.bin\n"
                             " 173099995  Apr 04 00:11:28 2020 n5000-uk9.5.2.1.N1.1a.bin\n"
                             " 174855143  May 04 05:53:25 2020 n5000-uk9.5.2.1.N1.9b.bin\n"
                             " 272308702  May 04 07:11:54 2020 n5000-uk9.7.0.8.N1.1.bin\n"
                             " 4096       Jan 03 01:43:02 2009 scripts/\n"
                             " 6803       Apr 20 12:47:24 2020 span.log\n"
                             " 4096       Jan 03 01:41:22 2009 vdc_2/\n"
                             " 4096       Jan 03 01:41:22 2009 vdc_3/\n"
                             " 4096       Jan 03 01:41:22 2009 vdc_4/\n"
                             " 4096       Jan 03 01:41:27 2009 virtual-instance/\n"
                             " 4096       Feb 25 00:42:31 2009 virtual-instance-stby-sync/\n\n"
                             "Usage for bootflash://sup-local\n"
                             " 840577024 bytes used\n"
                             " 810328064 bytes free\n"
                             "1650905088 bytes total\n")
        document.add_paragraph("Verify MD5 Checksum for the OS Images Integrity by issuing the “show file” command. Compare the results to Table 3 Software images for byte size & MD5 checksum before each stage of upgrade.")
        ImageUpload.funTable("N5K-1(config)# show file bootflash:///n5000-uk9-kickstart.7.3.6.N1.1.bin md5sum"
                             "\nfb2f0bb05c0ddb035d4335a446f429af\n")
        ImageUpload.funTable("N5K-1(config)# show file bootflash:///n5000-uk9.7.3.6.N1.1.bin md5sum"
                             "\nbef58f988bcf4daef7896ac84c96c10d")
        return

class CheckIncompatibilityOfSystemImage(Formatter):
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def checkIncompatibilityOfSystemImage_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph("Check the incompatibility of system image by using below command before each stage")
        CheckIncompatibilityOfSystemImage.funTable("# show incompatibility system <Image to be upgraded>\n"
                                                  "Ex: N5K-1(config)# show incompatibility system bootflash:n5000-uk9.7.3.6.N1.1.bin\n"
                                                  "No incompatible configurations")
        document.add_paragraph("Device upgrade is marked as failed if some error comes in above command output.")
        return

class IdentifyTheUpgradeImpact(Formatter):
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def identifyTheUpgradeImpact_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph("Identify the upgrade impact before each step of upgrade. Because L3 license is enabled and "
                               "SVI configured it will be disruptive. Please just copy the outputs to refer to later in case of "
                               "any issues. The upgrade table will be particularly useful for comparison if issues arise.")
        IdentifyTheUpgradeImpact.funTable("# show install all impact kickstart <Kickstart Image to be upgraded> system <System \n"
                                         "image to be upgraded>\n"
                                         "Ex: \n"
                                         "N5K-1(config)# show install all impact kickstart bootflash:n5000-uk9-\n"
                                         "kickstart.7.3.6.N1.1.bin system bootflash:n5000-uk9.7.3.6.N1.1.bin\n"
                                         "Verifying image bootflash:/n5000-uk9-kickstart.7.3.6.N1.1.bin for boot variable \n"
                                         "\"kickstart\".\n"
                                         "[####################] 100% -- SUCCESS\n"
                                         "Verifying image bootflash:/n5000-uk9.7.3.6.N1.1.bin for boot variable \"system\".\n"
                                         "[####################] 100% -- SUCCESS\n"
                                         "Verifying image type.\n"
                                         "[####################] 100% -- SUCCESS\n"
                                         "Extracting \"system\" version from image bootflash:/n5000-uk9.7.3.6.N1.1.bin.\n"
                                         "[####################] 100% -- SUCCESS\n"
                                         "Extracting \"kickstart\" version from image bootflash:/n5000-uk9-\n"
                                         "kickstart.7.3.6.N1.1.bin.\n"
                                         "[####################] 100% -- SUCCESS\n"
                                         "Extracting \"bios\" version from image bootflash:/n5000-uk9.7.3.6.N1.1.bin.\n"
                                         "[####################] 100% -- SUCCESS\n"
                                         "Extracting \"fexth\" version from image bootflash:/n5000-uk9.7.3.6.N1.1.bin.\n"
                                         "[####################] 100% -- SUCCESS\n"
                                         "Performing module support checks.\n"
                                         "[####################] 100% -- SUCCESS\n"
                                         "Notifying services about system upgrade.\n"
                                         "[####################] 100% -- SUCCESS\n"
                                         "Compatibility check is done:\n"
                                         "Module bootable Impact Install-type Reason\n"
                                         "------ -------- -------------- ------------ ------\n"
                                         " 1 yes disruptive reset Incompatible image\n"
                                         " 101 yes disruptive reset Incompatible image\n"
                                         "Images will be upgraded according to following table:\n"
                                         "Module      Image        Running-Version            New-Version Upg-Required\n"
                                         "------ ---------- ---------------------- ---------------------- ------------\n"
                                         " 1        system            5.2(1)N1(1a)            7.3(6)N1(1)          yes\n"
                                         " 1        kickstart         5.2(1)N1(1a)            7.3(6)N1(1)          yes\n"
                                         " 1        bios        v3.6.0(05/09/2012)     v3.6.0(05/09/2012)           no\n"
                                         " 1        SFP-uC                v1.1.0.0               v1.0.0.0           no\n"
                                         " 101      fexth             5.2(1)N1(1a)            7.3(6)N1(1)           yes\n"
                                         " 1        power-seq                 v3.0                   v3.0            no\n"
                                         " 1        uC                    v1.2.0.1               v1.2.0.1            no\n"
                                         "")
        document.add_paragraph("Note: In the upgrade cycle, if the power-sequencer micro controller upgrade is required, a “reload power-cycle” command need to be issued once both devices are upgraded to particular release. ")
        return

class UserPassword(Formatter):
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def userPassword_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph(f"As per {cust} request, we need to change the password for the local user account “admin” and share the password with {cust}.")
        UserPassword.funTable("# conf t\n"
                              "(config)# username admin password xxxxx role network-admin\n"
                              "(config)# exit\n"
                              "#\n")
        return
class HwStateVerification(Formatter):
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def hwStateVerification_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph("Verify the hardware state and inventory of the device (modules, power supplies, etc.)")
        HwStateVerification.funTable("# show version\n"
                                     "# show module\n"
                                     "# show hardware internal ?\n"
                                     "# show fex\n"
                                     "# show environment\n"
                                     "# show environment fex all\n"
                                     "# show system resource\n"
                                     "# show diagnostic result module all\n")
        try:
            r = document.add_paragraph().add_run()
            r.add_picture('note.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(f'Please look for any irregularities in modules, particularly the status shows “ok”.')
            r = document.add_paragraph().add_run()
            r.add_picture('note.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(f'From “show diagnostic result module all” output, Please make sure all tests pass, look for failures, incomplete, untested and abort notifications. ')
            r = document.add_paragraph().add_run()
            r.add_picture('note.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(f'Some of those commands are platform specific and there may be slight variations depending on software version and hardware model')
            r = document.add_paragraph("\t").add_run()
            r.add_picture('caution.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(f'If the Nexus 5000 switch has FEX modules attached, then FEX pre-provisioning must be enabled before any upgrade is started. This is to ensure that HIF configuration losses caused by the upgrade are automatically restored when the module comes back online.')
            HwStateVerification.funTable("# show fex\n"
                                         "# config terminal\n"
                                         "# slot <100-199>\n"
                                         "# provision model <MODEL> (example N2K-C2232P)\n"
                                         "# exit\n")
            r = document.add_paragraph("\t").add_run()
            r.add_picture('caution.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(f'Before issuing “show install all” command inactive policies/commands in the running configuration should be removed')
            r = document.add_paragraph("\t").add_run()
            r.add_picture('caution.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(f'Writing the active configuration to NVRAM is needed to avoid failures and configuration loss upon reboot. ')
            HwStateVerification.funTable("# copy running-config startup-config")
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
        return
class ServicesandIncidentVerification:
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def servicesandIncidentVerification_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph("Contact the service owners and verify that no active incidents are reported for the device to be "
                               "upgraded. Obtain their permission to start the maintenance works\n\n"
                               "Save, commit, or discard any active configuration sessions before upgrading or downgrading the "
                               "Cisco NX-OS software image on your device\n\n"
                               "Place devices in maintenance mode in all monitoring tools in order to suppress false positive alarms caused by activities described herein")
        return

class PreUpgradeChecksCapture(Formatter):
    def __init__(self,os,version):
        self.os=os
        self.version=version
    def preUpgradeChecksCapture_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph("Collect the following list of show commands. The same set of commands will be collected again at the "
                               "end of the maintenance window after the upgrade completion (or after a Rollback in case it will be"
                               "needed). After collecting the second round of these outputs they will be compared to identify any "
                               "potential control plane gaps which may have occurred during/after the upgrade completion. ")
        try:
            r = document.add_paragraph().add_run()
            r.add_picture('note.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(f'The verification command list may not be completely applicable to the particular set up. It is up to the sole discretion of the update engineer to choose which commands to execute based on the particular setup')
            r = document.add_paragraph().add_run()
            r.add_picture('note.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(f'Health Check commands in the list do not need comparison but are used to briefly check the system health. Before executing the commands from the list execute:')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
        PreUpgradeChecksCapture.funTable("# terminal length 0\n"
                                         "# show clock")
        p=document.add_paragraph("\n")
        p.add_run("Table 5 Pre-check show commands list").italic = True
        table = document.add_table(rows=22, cols=3, style = 'Colorful Grid Accent 1')
        row = table.rows[0].cells
        row[0].text = 'Category'
        row[1].text = 'show commands'
        row[2].text = 'Exact Match'
        row = table.rows[1].cells
        row[1].text = 'show version | inc BIOS|kickstart|system'
        row[2].text = ''
        row = table.rows[2].cells
        row[1].text = ''
        row[2].text = ''
        a = table.cell(1,0)
        b = table.cell(2,0)
        A = a.merge(b)
        a.text = 'Version'
        row = table.rows[3].cells
        row[0].text = 'BOOT'
        row[1].text = 'show boot'
        row = table.rows[4].cells
        row[0].text = 'Module'
        row[1].text = 'show module'
        row[2].text = 'X (include OS version check)'
        row = table.rows[5].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[6].cells
        row[1].text = 'show interface brief | count'
        row[2].text = 'X'
        row = table.rows[7].cells
        row[1].text = 'show interface brief'
        row[2].text = 'X'
        row = table.rows[8].cells
        row[1].text = 'show ip interface brief include-secondary vrf all '
        row[2].text = 'X'
        row = table.rows[9].cells
        row[1].text = 'show ip interface brief operational'
        row[2].text = ''
        a= table.cell(6,0)
        b= table.cell(7,0)
        c= table.cell(8,0)
        d= table.cell(9,0)
        A = a.merge(b).merge(c).merge(d)
        A.text ='Interfaces'
        row = table.rows[10].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[11].cells
        row[0].text = 'Port-Channel'
        row[1].text = 'show port-channel summary'
        row[2].text = 'X'
        row = table.rows[12].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[13].cells
        row[1].text = 'show vpc'
        row[2].text = ''
        row = table.rows[14].cells
        row[1].text = 'show vpc brief'
        row[2].text = ''
        row = table.rows[15].cells
        row[1].text = 'show vpc consistency-parameters global'
        row[2].text = ''
        row = table.rows[16].cells
        row[1].text = 'show vpc role'
        row[2].text = ''
        row = table.rows[17].cells
        row[1].text = 'show vpc peer-keepalive'
        row[2].text = ''
        a = table.cell(13,0)
        b= table.cell(14,0)
        c= table.cell(15,0)
        d= table.cell(16,0)
        e= table.cell(17,0)
        A = a.merge(b).merge(c).merge(d).merge(e)
        A.text ='vPC'

        row = table.rows[18].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[19].cells
        row[0].text = 'FEX'
        row[1].text = 'show fex'
        row[2].text = ''
        row = table.rows[20].cells
        row[0].text = ''
        row[1].text = 'show fex detail'
        row[2].text = ''
        row = table.rows[21].cells
        row[0].text = ''
        row[1].text = 'show diagnostic result fex all'
        row[2].text = ''
        document.add_paragraph()
        table = document.add_table(rows=26, cols=3, style = 'Colorful Grid Accent 1')
        row = table.rows[0].cells
        row[0].text = 'Category'
        row[1].text = 'show commands'
        row[2].text = 'Exact Match'
        row = table.rows[1].cells
        row[0].text = ''
        row[1].text = 'Show interface fex-fabric'
        row[2].text = ''
        row = table.rows[2].cells
        row[0].text = 'Configuration '
        row[1].text = 'show running-config'
        row[2].text = 'X1(observe release notes for default config, such as class maps)'
        row = table.rows[3].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[4].cells
        row[1].text = 'show feature'
        row[2].text = 'X'
        a = table.cell(4,0)
        b = table.cell(5,0)
        A = a.merge(b)
        a.text = 'Feature'
        row = table.rows[5].cells
        row[1].text = ''
        row[2].text = ''
        row = table.rows[6].cells
        row[0].text = 'Enviroment'
        row[1].text = 'show environment'
        row[2].text = ''
        row = table.rows[7].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[8].cells
        row[0].text = 'Version'
        row[1].text = 'show version'
        row[2].text = ''
        row = table.rows[9].cells
        row[0].text = '' 
        row[1].text = 'show fex <fex_ID> version'
        row[2].text = ''
        row = table.rows[10].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[11].cells
        row[0].text = 'MAC addresses '
        row[1].text = 'show mac address-table count'
        row[2].text = ''
        row = table.rows[12].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[13].cells
        row[1].text = 'show udld global'
        row[2].text = ''
        row = table.rows[14].cells
        row[1].text = 'show udld neighbors'
        row[2].text = 'X'
        a = table.cell(12,0)
        b = table.cell(13,0)
        A = a.merge(b)
        A.text = 'UDLD'
        row = table.rows[15].cells
        row[0].text = 'ARP'
        row[1].text = 'show ip arp summary vrf all'
        row[2].text = ''
        row = table.rows[16].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[17].cells
        row[1].text = 'show vlan summary '
        row[2].text = 'X'
        row = table.rows[18].cells
        row[0].text = ''
        row[1].text = 'show vlan brief'
        row[2].text = 'X'
        row = table.rows[19].cells
        row[1].text = 'show vtp status'
        row[2].text = 'X'
        row = table.rows[20].cells
        row[1].text = 'show system vlan reserved'
        row[2].text = ''
        a= table.cell(17,0)
        b= table.cell(18,0)
        c= table.cell(19,0)
        d= table.cell(20,0)
        A = a.merge(b).merge(c).merge(d)
        A.text ='VLANs'
        row = table.rows[21].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[22].cells
        row[1].text = 'show fabricpath isis adjacency summary '
        row[2].text = 'X'
        row = table.rows[23].cells
        row[1].text = 'show fabricpath isis database'
        row[2].text = ''
        row = table.rows[24].cells
        row[1].text = 'show fabricpath isis ip mroute'
        row[2].text = ''
        row = table.rows[25].cells
        row[1].text = 'show fabricpath isis route'
        row[2].text = ''
        a= table.cell(22,0)
        b= table.cell(23,0)
        c= table.cell(24,0)
        d= table.cell(25,0)
        A = a.merge(b).merge(c).merge(d)
        A.text = 'Fabric-path state'
        document.add_paragraph("____________________________________________________"
                               "1\n"
                               "Caution: this upgrade adds lacp suspend-individual\n"
                               "Changes seen with upgrade from 5.2(1)N1(4)\n"
                               "class-map type qos \n"
                               "class-fcoe \n"
                               "class-map type queuing \n"
                               "class-fcoe match qos-group 1 \n"
                               "class-map type queuing \n"
                               "class-all-flood match qos-group 2 \n"
                               "class-map type queuing \n"
                               "class-ip-multicast match qos-group 2 \n"
                               "class-map type network-qos \n"
                               "class-fcoe match qos-group 1 \n"
                               "class-map type network-qos \n"
                               "class-all-flood match qos-group 2 \n"
                               "class-map type network-qos \n"
                               "class-ip-multicast match qos-group 2")
        table = document.add_table(rows=36, cols=3, style = 'Colorful Grid Accent 1')
        row = table.rows[0].cells
        row[0].text = 'Category'
        row[1].text = 'show commands'
        row[2].text = 'Exact Match'
        row = table.rows[1].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[2].cells
        row[1].text = 'show spanning-tree summary total '
        row[2].text = 'X'
        row = table.rows[3].cells
        row[1].text = 'show spanning-tree | inc VLAN|Spanning'
        row[2].text = ''
        
        row = table.rows[4].cells
        row[1].text = 'show spanning-tree inconsistentports'
        row[2].text = ''
        row = table.rows[5].cells
        row[1].text = 'show spanning-tree blockedports'
        row[2].text = ''
        row = table.rows[6].cells
        row[1].text = 'show spanning-tree internal info global | inc ports'
        row[2].text = ''
        a = table.cell(2,0)
        b = table.cell(3,0)
        c = table.cell(4,0)
        d = table.cell(5,0)
        e = table.cell(6,0)
        A = a.merge(b).merge(c).merge(d).merge(e)
        A.text = 'Spanning tree'
        row = table.rows[7].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[8].cells
        row[1].text = ''
        row[2].text = ''
        row = table.rows[9].cells
        row[1].text = 'show forwarding ipv4 route summary vrf all'
        row[2].text = ''
        row = table.rows[10].cells
        row[1].text = 'show forwarding ipv4 route vrf all | count'
        row[2].text = ''
        row = table.rows[11].cells
        row[1].text = 'show forwarding ipv4 route module [Module] | count'
        row[2].text = ''
        row = table.rows[12].cells
        row[1].text = 'show forwarding ipv6 adjacency detail '
        row[2].text = 'X'
        row = table.rows[13].cells
        row[1].text = 'show forwarding ipv6 route vrf all | count'
        row[2].text = ''
        row = table.rows[14].cells
        row[1].text = 'show forwarding ipv6 route vrf all module [Module] | count'
        row[2].text = ''
        row = table.rows[15].cells
        row[1].text = 'show forwarding route summary vrf all '
        row[2].text = ''
        row = table.rows[16].cells
        row[1].text = 'show ip route vrf all summary'
        row[2].text = ''
        a = table.cell(8,0)
        b = table.cell(9,0)
        c = table.cell(10,0)
        d = table.cell(11,0)
        e = table.cell(12,0)
        f = table.cell(13,0)
        g = table.cell(14,0)
        h = table.cell(15,0)
        i = table.cell(16,0)
        A = a.merge(b).merge(c).merge(d).merge(e).merge(f).merge(g).merge(h).merge(i)
        A.text= 'FIB unicast'
        row = table.rows[17].cells
        row[0].text = ''
        row[1].text = 'show ipv6 route summary vrf all'
        row[2].text = ''
        row = table.rows[18].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[19].cells
        row[1].text = 'show forwarding multicast route summary vrf all '
        row[2].text = '*2'
        row = table.rows[20].cells
        row[1].text = 'show ip mroute summary count vrf all'
        row[2].text = ''
        row = table.rows[21].cells
        row[1].text = 'show ip mroute summary count vrf all | egrep "Total number of routes|Group|VRF" | egrep 1'
        row[2].text = ''
        a = table.cell(19,0)
        b = table.cell(20,0)
        c = table.cell(21,0)
        A = a.merge(b).merge(c)
        A.text = 'FIB multicast '
        row = table.rows[22].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[23].cells
        row[1].text = ''
        row[2].text = ''
        row = table.rows[24].cells
        row[1].text = 'show ip access-lists summary | inc ACL|Total'
        row[2].text = ''
        row = table.rows[25].cells
        row[1].text = 'show ip access-lists summary'
        row[2].text = 'X'
        a = table.cell(23,0)
        b = table.cell(24,0)
        c = table.cell(25,0)
        A = a.merge(b).merge(c)
        A.text = 'ACLs/TCAM'
        row = table.rows[26].cells
        row[0].text = ''
        row[1].text = ''
        row[2].text = ''
        row = table.rows[27].cells
        row[1].text = 'show resource internal info resource'
        row[2].text = ''
        row = table.rows[28].cells
        row[1].text = 'show system resources'
        row[2].text = ''
        row = table.rows[29].cells
        row[1].text = 'show hardware internal cpu-mac inband stats'
        row[2].text = ''
        row = table.rows[30].cells
        row[1].text = 'show policy-map type control-plane'
        row[2].text = 'x'
        row = table.rows[31].cells
        row[1].text = 'show diagnostic result module all'
        row[2].text = ''
        row = table.rows[32].cells
        row[1].text = 'show logging last 100'
        row[2].text = ''
        row = table.rows[33].cells
        row[1].text = 'show logging '
        row[2].text = ''
        row = table.rows[34].cells
        row[1].text = 'show process cpu sort'
        row[2].text = ''
        row = table.rows[35].cells
        row[1].text = 'show logging internal info'
        row[2].text = ''
        a = table.cell(27,0)
        b = table.cell(28,0)
        c = table.cell(29,0)
        d = table.cell(30,0)
        e = table.cell(31,0)
        f = table.cell(32,0)
        g = table.cell(33,0)
        h = table.cell(34,0)
        i = table.cell(35,0)
        A = a.merge(b).merge(c).merge(d).merge(e).merge(f).merge(g).merge(h).merge(i)
        A.text= 'Health check'
        document.add_page_break()
        return

class UpgradeProcedure(Formatter):
    def __init__(self,os,version):
        self.os=os
        self.version=version
    def upgradeProcedure_Nexus_5548_7_3_6_N1_1_(self):
        upgrade = (
        ('Task Number', '1'),
        ('Task Description ', 'Configuration and License backup'),
        ('Task Scheduled Start', ''),
        ('Task Duration','10 minutes'),
        ('Task Owner Pri/Sec',''),
        ('Task Dependencies','Network is operational')
        )
        table = document.add_table(rows=1, cols=2, style='Colorful Grid Accent 1')
        for term, definition in upgrade:
            row_cells = table.add_row().cells
            row_cells[0].text = str(term)
            row_cells[1].text = str(definition)

        p = document.add_paragraph("\n")
        p.add_run("Task Detail ").bold = True
        p = document.add_paragraph()

        p.add_run("Step 1 ").bold = True
        p.add_run('Save the configuration of the switch')
        document.add_paragraph('Log in to Cisco switch.',style='List Bullet 2')
        document.add_paragraph('Write the configuration to the NVRAM using:',style='List Bullet 2')
        UpgradeProcedure.funTable('# copy running-config startup-config\n'
                                  '\t N5K-2(config)# copy run start\n'
                                  '\t [########################################] 100%\n'
                                  '\t Copy complete, now saving to disk (please wait)...\n'
                                  '\t N5K-1(config)# copy run start\n'
                                  '\t ########################################] 100%\n'
                                  '\t Copy complete, now saving to disk (please wait)...\n')
        UpgradeProcedure.funTable('# checkpoint file bootflash:show_pre_7.3_config\n'
                                  '\t N5K-2(config)# checkpoint file bootflash:show_pre_7.3_config\n'
                                  '\t Done\n'
                                  '\t N5K-1(config)# checkpoint file bootflash:show_pre_7.3_config\n'
                                  '\t Done')
        document.add_paragraph('Save the configuration on boothflash and/or TFTP/FTP server and/or USB stick:',style='List Bullet 2')
        UpgradeProcedure.funTable('\t# copy startup-config bootflash://<switchname>_run_conf_dd_mm\n'
                                  '\t# copy startup-config usb1: <switchname>_run_conf_dd_mm\n'
                                  '\t# copy startup-config tftp: <switchname>_run_conf_dd_mm')
        r = document.add_paragraph().add_run()
        r.add_picture('note.jpg', width=Inches(0.3), height=Inches(.3))
        r.add_text('Please note that the USB which will be used has to be previously formatted on a Nexus 5000 switch running NX-OS 7.x release or later or having FAT32 file system. This is due to issues that may arise if another file system was used for formatting')

        p = document.add_paragraph()
        p.add_run("Step 2 ").bold = True
        p.add_run('Back up license files of the switch:')
        document.add_paragraph('In order to create a backup of the Core switches license files to TFTP and/or USB memory stick the following command should be used:',style='List Bullet 2')
        UpgradeProcedure.funTable('\t# copy licenses bootflash://<switchname>_licenses_backup.tar\n'
                                  '\t# copy licenses usb1:<switchname>_licenses_backup.tar \n'
                                  '\t# copy bootflash://<switchname>_licenses_backup.tar tftp://<IP>')
        p = document.add_paragraph()
        p.add_run("Task Success Criteria").bold = True
        document.add_paragraph('All configuration, license files and Fex configurations are successfully stored and copied to the relevant backup locations. ')
        p = document.add_paragraph()
        p.add_run("Failure Procedure").bold = True
        document.add_paragraph('\t1. Resolve the problem, if one discovered, and if possible.\n'
                               '\t2. Provision spare UBS memory sticks with 2GB or larger size formatted in NX-OS version 7.x \n'
                               '\t3. Call Cisco TAC for support')
        upgrade1 = (
            ('Task Number', '2'),
            ('Task Description ', 'Execute upgrade to target release – 7.3(6)N1(1)'),
            ('Task Scheduled Start', ''),
            ('Task Duration','45 minutes'),
            ('Task Owner Pri/Sec',''),
            ('Task Dependencies','Network is operational,INB/OOB/Console access to the device')
        )
        table = document.add_table(rows=1, cols=2, style='Colorful Grid Accent 1')
        for term, definition in upgrade1:
            row_cells = table.add_row().cells
            row_cells[0].text = str(term)
            row_cells[1].text = str(definition)
        try:
            r = document.add_paragraph().add_run()
            r.add_picture('note.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('Please note that the actual upgrade time may vary ')

            r = document.add_paragraph("\t").add_run()
            r.add_picture('caution.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('\tExecute this step only on devices running software image prior to 7.3(6)N1(1) ')

            r = document.add_paragraph("\t").add_run()
            r.add_picture('caution.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('\tWhen doing this upgrade on the pair of Nexus switches always start with vPC Primary in the topology. ')

            r = document.add_paragraph("\t").add_run()
            r.add_picture('caution.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('\tThis step may require power sequencer or microcontroller firmware upgrades.If it is required, “reload power-cycle” will need to be issued once both devices are upgraded.')
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found...")
        p = document.add_paragraph()
        p.add_run("Task Detail: ").bold = True
        p.add_run('This section covers upgrading from 5.2(1)N1(1a) to final release 7.3(6)N1(1).The upgrade may be disruptive in this step.')

        p = document.add_paragraph("\n")
        p.add_run("Step 1 ").bold = True
        p.add_run('Gain SSH access to the device:')

        p = document.add_paragraph("\n")
        p.add_run("Step 2 ").bold = True
        p.add_run('Verify that you have no active configuration sessions:')
        UpgradeProcedure.funTable('\tN5K-2(config)# show configuration session summary\n'
                                  '\tThere are no active configuration sessions \n'
                                  '\tN5K-1(config)# show configuration session summary\n'
                                  '\tThere are no active configuration sessions ' )

        p = document.add_paragraph("\n")
        p.add_run("Step 3 ").bold = True
        p.add_run(' Configure FEX module pre-provisioning for all the FEXs connected to both the witches (vPC primary and vPC secondary switches).')

        r = document.add_paragraph().add_run()
        r.add_text('This is done to work around bug CSCul22703.')
        r.font.color.rgb = RGBColor(255, 0, 0)
        document.add_paragraph('\tEnsure that the running-configuration has been saved to startup-configuration. Also, keep a backup copy of the running-configuration on Flash. These measures when coupled with FEX pre-provisioning enable an engineer to side-step a large part of the impact of FEX configuration loss following manual upgrades. The configuration loss occurs because the system reapplies the FEX configuration as soon as the Cisco Nexus switch comes back up, without waiting for the FEX to come online. As a result, the FEX HIF configuration is sometimes rejected because the system does not see the FEX online and the FEX interfaces do not appear to be available. Perform these on Both the switches')
        UpgradeProcedure.funTable('\tN5K-2(config)# slot 100\n'
                                  '\tN5K-2(config-slot)# provision model N2K-C2248T \n'
                                  '\t# N5K-2(config)# copy run start\n'
                                  '\t[########################################] 100%\n'
                                  '\tCopy complete, now saving to disk (please wait)...\n'
                                  '\tNN5K-1(config)# slot 100 \n'
                                  '\t# N5K-1(config-slot)# provision model N2K-C2248T\n'
                                  '\tN5K-1(config)# copy run start\n'
                                  '\t[########################################] 100%\n'
                                  '\tCopy complete, now saving to disk (please wait)...\n')

        p = document.add_paragraph("\n")
        p.add_run("Step 4 ").bold = True
        p.add_run('Proceed to perform the upgrade using the install all command at the command prompt on vPC primary switch.')
        UpgradeProcedure.funTable('N5K-2(config)# install all kickstart n5000-uk9-kickstart.7.3.6.N1.1.bin system\n'
                                  'n5000-uk9.7.3.6.N1.1.bin\n'
                                  'Verifying image bootflash:/n5000-uk9-kickstart.7.3.6.N1.1.bin for boot variable\n'
                                  'kickstart.\n'
                                  '[####################] 100% -- SUCCESS\n'
                                  'Verifying image bootflash:/n5000-uk9.7.3.6.N1.1.bin for boot variable "system".\n'
                                  '[####################] 100% -- SUCCESS\n'
                                  'Verifying image type.\n'
                                  '[####################] 100% -- SUCCESS\n'
                                  'Extracting "system" version from image bootflash:/n5000-uk9.7.3.6.N1.1.bin.\n'
                                  '[####################] 100% -- SUCCESS\n'
                                  'Extracting "kickstart" version from image bootflash:/n5000-uk9-\n'
                                  'kickstart.7.3.6.N1.1.bin.\n'
                                  '[####################] 100% -- SUCCESS\n'
                                  'Extracting "bios" version from image bootflash:/n5000-uk9.7.3.6.N1.1.bin.\n'
                                  '[####################] 100% -- SUCCESS\n'
                                  'Extracting "fexth" version from image bootflash:/n5000-uk9.7.3.6.N1.1.bin.\n'
                                  '[####################] 100% -- SUCCESS\n'
                                  'Performing module support checks.\n'
                                  '[####################] 100% -- SUCCESS\n'
                                  'Notifying services about system upgrade.\n'
                                  '[####################] 100% -- SUCCESSe\n'
                                  'Compatibility check is done:e\n'
                                  'Module bootable Impact Install-type Reasone\n'
                                  '------ -------- -------------- ------------ ------e\n'
                                  '1 yes disruptive reset Incompatible imagee\n'
                                  '101 yes disruptive reset Incompatible image\n'
                                  'Images will be upgraded according to following table:e\n'
                                  'Module Image Running-Version New-Version Upg-Requirede\n'
                                  '------ ---------- ---------------------- ---------------------- ------------\n'
                                  '     1     system           5.2(1)N1(1a)            7.3(6)N1(1)          yes\n'
                                  '     1  kickstart           5.2(1)N1(1a)            7.3(6)N1(1)          yes\n'
                                  '     1       bios     v3.6.0(05/09/2012)     v3.6.0(05/09/2012)           no\n'
                                  '     1     SFP-uC               v1.1.0.0               v1.0.0.0           no\n'
                                  '     101    fexth           5.2(1)N1(1a)            7.3(6)N1(1)          yes\n'
                                  '     1  power-seq                   v3.0                   v3.0           no\n'
                                  '     1         uC               v1.2.0.1               v1.2.0.1           no\n'
                                  'Switch will be reloaded for disruptive upgrade.\n'
                                  'Do you want to continue with the installation (y/n)? [n] y\n'
                                  'Install is in progress, please wait.\n'
                                  'Performing runtime checks.\n'
                                  '[####################] 100% -- SUCCESS\n'
                                  'Setting boot variables.\n'
                                  '[####################] 100% -- SUCCESS\n'
                                  'Performing configuration copy.\n'
                                  '[####################] 100% -- SUCCESS\n'
                                  'Pre-loading modules.]\n'
                                  '[This step might take upto 20 minutes to complete - please wait.]\n'
                                  '[*Warning -- Please do not abort installation/reload or powercycle fexes*]\n'
                                  '[####################] 100% -- SUCCESS\n'
                                  'Finishing the upgrade, switch will reboot in 10 seconds.\n'
                                  '5K-1# Shutdown Ports..\n'
                                  'writing reset reason 49,'
                                  )

        p = document.add_paragraph("\n")
        p.add_run("Step 5 ").bold = True
        p.add_run(' Verify that the upgrade of the vPC primary switch is completed successfully. At the completion of the upgrade, the vPC primary switch restores vPC peering. Verify the status of version of the upgraded device, VPC and FEX status ')

        p = document.add_paragraph("\n")
        p.add_run("Note: ").bold = True
        p.add_run(' During the upgrade process, when the switch is reloaded, only singled-homed FEXs connected to the switch are reloaded and dual-homed FEXs are not reloaded. Servers connected to the dual-homed FEXs retain network connectivity through the vPC secondary switch.')


        UpgradeProcedure.funTable('N5K-2# sh version\n'
                                  'Cisco Nexus Operating System (NX-OS) Software\n'
                                  'TAC support: http://www.cisco.com/tac\n'
                                  'Documents: \n'
                                  'http://www.cisco.com/en/US/products/ps9372/tsd_products_support_series_home.html\n'
                                  'Copyright (c) 2002-2019, Cisco Systems, Inc. All rights reserved.\n'
                                  'The copyrights to certain works contained herein are owned by\n'
                                  'other third parties and are used and distributed under license.\n'
                                  'Some parts of this software are covered under the GNU Public\n'
                                  'License. A copy of the license is available at\n'
                                  'http://www.gnu.org/licenses/gpl.html.\n'
                                  'Software\n'
                                  ' BIOS: version 3.6.0\n'
                                  ' Power Sequencer Firmware: \n'
                                  ' Module 1: v3.0\n'
                                  ' Module 2: v5.0\n'
                                  ' Microcontroller Firmware: version v1.2.0.1\n'
                                  ' QSFP Microcontroller Firmware: \n'
                                  ' Module not detected\n'
                                  ' CXP Microcontroller Firmware: \n'
                                  ' Module not detected\n'
                                  ' kickstart: version 7.3(6)N1(1)\n'
                                  ' system: version 7.3(6)N1(1)\n'
                                  ' BIOS compile time: 05/09/2012\n'
                                  ' kickstart image file is: bootflash:///n5000-uk9-kickstart.7.3.6.N1.1.bin\n'
                                  ' kickstart compile time: 9/11/2019 23:00:00 [09/12/2019 07:56:00]\n'
                                  ' system image file is: bootflash:///n5000-uk9.7.3.6.N1.1.bin\n'
                                  ' system compile time: 9/11/2019 23:00:00 [09/12/2019 12:10:27]\n'
                                  'May 18, 2021 Software Upgrade Procedure: NX-OS Nexus devices\n'
                                  'Cisco Confidential. All printed copies and duplicate soft copies are considered uncontrolled\n'
                                  'And the original online version should be referred to for the latest version.\n'
                                  'Page 26 of 41\n'
                                  'Hardware\n'
                                  ' cisco Nexus5548 Chassis ("O2 32X10GE/Modular Universal Platform Supervisor")\n'
                                  ' Intel(R) Xeon(R) CPU with 8253792 kB of memory.\n'
                                  ' Processor Board ID FOC17350D7M\n'
                                  ' Device name: N5K-2\n'
                                  ' bootflash: 2007040 kB\n'
                                  'Kernel uptime is 0 day(s), 1 hour(s), 15 minute(s), 22 second(s)\n'
                                  'Last reset at 367578 usecs after Mon May 4 18:57:48 2020\n'
                                  ' Reason: Disruptive upgrade\n'
                                  ' System version: 7.1(5)N1(1b)\n'
                                  ' Service: \n'
                                  'plugin\n'
                                  ' Core Plugin, Ethernet Plugin\n'
                                  'Active Package(s)\n'
                                  'N5K-2# sh vpc\n'
                                  'Legend:\n'
                                  ' (*) - local vPC is down, forwarding via vPC peer-link\n'
                                  'vPC domain id : 100 \n'
                                  'Peer status : peer adjacency formed ok \n'
                                  'vPC keep-alive status : peer is alive \n'
                                  'Configuration consistency status : success \n'
                                  'Per-vlan consistency status : success \n'
                                  'Type-2 consistency status : success \n'
                                  'vPC role : primary, operational secondary\n'
                                  'Number of vPCs configured : 2 \n'
                                  'Peer Gateway : Disabled\n'
                                  'Dual-active excluded VLANs : -\n'
                                  'Graceful Consistency Check : Enabled\n'
                                  'Operational Layer3 Peer-router : Disabled\n'
                                  'Auto-recovery status : Enabled (timeout = 240 seconds)\n'
                                  'vPC Peer-link status\n'
                                  '---------------------------------------------------------------------\n'
                                  'id Port Status Active vlans \n'
                                  '-- ---- ------ --------------------------------------------------\n'
                                  '1  Po20 up     1,10,100 \n'
                                  'vPC status\n'
                                  '----------------------------------------------------------------------------\n'
                                  'id     Port        Status Consistency Reason                     Active vlans\n'
                                  '------ ----------- ------ ----------- -------------------------- -----------\n'
                                  '100    Po100       up     success     success                    - \n'
                                  '1000   Po1000      down*  success     success                    - \n\n')
        UpgradeProcedure.funTable('N5K-2# sh fex detail\n'
                                  'FEX: 100 Description: FEX0100 state: AA Version Mismatch\n'
                                  'FEX version: 7.1(5)N1(1b) [Switch version: 7.3(6)N1(1)]\n'
                                  ' FEX Interim version: 7.1(5)N1(1b)\n'
                                  ' Switch Interim version: 7.3(6)N1(1)\n'
                                  'Logs:\n'
                                  '05/04/2020 20:13:06.421014: Module register received\n'
                                  '05/04/2020 20:13:06.423569: Image Version Mismatch\n'
                                  '05/04/2020 20:13:06.423719: A-A Version Mismatch\n'
                                  'May 18, 2021 Software Upgrade Procedure: NX-OS Nexus devices\n'
                                  'Cisco Confidential. All printed copies and duplicate soft copies are considered uncontrolled\n'
                                  'And the original online version should be referred to for the latest version.\n'
                                  'Page 27 of 41\n\n')
        UpgradeProcedure.funTable('N5K-1# sh vpc\n'
                                  'Legend:\n'
                                  ' (*) - local vPC is down, forwarding via vPC peer-link\n'
                                  'vPC domain id : 100 \n'
                                  'Peer status : peer adjacency formed ok \n'
                                  'vPC keep-alive status : peer is alive \n'
                                  'Configuration consistency status : success \n'
                                  'Per-vlan consistency status : success \n'
                                  'Type-2 consistency status : success \n'
                                  'vPC role : secondary, operational primary\n'
                                  'Number of vPCs configured : 50 \n'
                                  'Peer Gateway : Disabled\n'
                                  'Dual-active excluded VLANs : -\n'
                                  'Graceful Consistency Check : Enabled\n'
                                  'Auto-recovery status : Enabled (timeout = 240 seconds)\n'
                                  'vPC Peer-link status\n'
                                  '---------------------------------------------------------------------\n'
                                  'id Port Status Active vlans \n'
                                  '-- ---- ------ --------------------------------------------------\n'
                                  '1  Po20 up     1,10,100 \n'
                                  'vPC status\n'
                                  '----------------------------------------------------------------------------\n'
                                  'id     Port        Status Consistency Reason                     Active vlans\n'
                                  '------ ----------- ------ ----------- -------------------------- -----------\n'
                                  '100    Po100       up     success     success                    - \n'
                                  '1000   Po1000      down*  Not         Consistency Check Not - \n'
                                  '                          Applicable  Performed \n\n')
        UpgradeProcedure.funTable('N5K-1(config)# sh fex\n'
                                  ' FEX FEX FEX FEX \n'
                                  'Number Description State Model Serial \n'
                                  '------------------------------------------------------------------------\n'
                                  '100 FEX0100 Online N2K-C2248TP-1GE SSI15310B1N\n')

        p = document.add_paragraph("\n")
        p.add_run("Step 6: ").bold = True
        p.add_run(' Reload the dual-homed FEXs using the reload fex command from the vPC secondary switch. Reload the FEXs one-by-one or in a bunch of two or three FEXs. The servers connected to the dual-homed FEXs will lose connectivity.Conneecondary switches')
        UpgradeProcedure.funTable('N5K-1(config)# reload fex 100\n'
                                  'WARNING: This command will reboot FEX 100\n'
                                  'Do you want to continue? (y/n) [n] y\n'
                                  'N5K-1(config)')
        p = document.add_paragraph("\n")
        p.add_run('Step 7: ').bold = True
        p.add_run('Wait for the FEXs to reload. After the reload, the FEXs connect to the upgraded switch (vPC primary switch).')
        UpgradeProcedure.funTable('N5K-2# sh fex\n'
                                  '  FEX       FEX        FEX                 FEX \n'
                                  'Number  Description   State       Model            Serial\n'
                                  '---------------------------------------------------------------\n'
                                  '100     FEX0100       Online      N2K-C2248TP-1GE  SSI15310B1N\n\n'
                                  'sh fex detail snap below\n'
                                  'FEX version: 7.3(6)N1(1) [Switch version: 7.3(6)N1(1)]\n'
                                  'FEX Interim version: 7.3(6)N1(1)\n'
                                  'Switch Interim version: 7.3(6)N1(1)\n'
                                  'Logs:\n'
                                  '05/04/2020 20:22:39.877144: Module register received\n'
                                  '05/04/2020 20:22:39.881018: Registration response sent\n'
                                  '05/04/2020 20:22:40.022366: create module inserted event.\n'
                                  '05/04/2020 20:22:40.023434: Module Online Sequence\n'
                                  '05/04/2020 20:22:46.602091: Module Online')
        document.add_paragraph("\n")
        UpgradeProcedure.funTable('N5K-1(config)# sh fex\n'
                                  ' FEX       FEX          FEX                    FEX \n'
                                  'Number Description     State         Model         Serial \n'
                                  '--------------------------------------------------------------------\n'
                                  '100      FEX0100       Online   N2K-C2248TP-1GE   SSI15310B1N\n\n'
                                  'sh fex detail snap below\n'
                                  'FEX: 100 Description: FEX0100 state: AA Version Mismatch\n'
                                  ' FEX version: 7.3(6)N1(1) [Switch version: 7.1(5)N1(1b)]\n'
                                  ' FEX Interim version: 7.3(6)N1(1)\n'
                                  ' Switch Interim version: 7.1(5)N1(1b)\n'
                                  'Logs:\n'
                                  '05/04/2020 20:22:27.383298: Deleting route to FEX\n'
                                  '05/04/2020 20:22:27.391562: Module disconnected\n'
                                  '05/04/2020 20:22:27.394288: Offlining Module\n'
                                  '05/04/2020 20:22:33.845421: Module register received\n'
                                  '05/04/2020 20:22:33.847425: Image Version Mismatch\n'
                                  '05/04/2020 20:22:33.847611: A-A Version Mismatch\n')

        p = document.add_paragraph("\n")
        p.add_run("Step 8: ").bold = True
        p.add_run('Upgrade the vPC secondary switch with the new image using the install all kickstartimage systemimage command.')
        UpgradeProcedure.funTable('N5K-1(config)# install all kickstart n5000-uk9-kickstart.7.3.6.N1.1.bin system\n'
                                  'n5000-uk9.7.3.6.N1.1.bin\n\n'
                                  'Verifying image bootflash:/n5000-uk9-kickstart.7.3.6.N1.1.bin for boot variable\n'
                                  'kickstart\n'
                                  '[####################] 100% -- SUCCESS\n\n'
                                  'Verifying image bootflash:/n5000-uk9.7.3.6.N1.1.bin for boot variable "system".\n'
                                  '[####################] 100% -- SUCCESS\n\n'
                                  'Verifying image type.\n'
                                  '[####################] 100% -- SUCCESS\n\n'
                                  'Extracting "system" version from image bootflash:/n5000-uk9.7.3.6.N1.1.bin.\n'
                                  '[####################] 100% -- SUCCESS\n\n'
                                  'Extracting "kickstart" version from image bootflash:/n5000-uk9-\n'
                                  'kickstart.7.3.6.N1.1.bin.\n'
                                  '[####################] 100% -- SUCCESS\n\n'
                                  'Extracting "bios" version from image bootflash:/n5000-uk9.7.3.6.N1.1.bin.\n'
                                  '[####################] 100% -- SUCCESS\n\n'
                                  'Extracting "fexth" version from image bootflash:/n5000-uk9.7.3.6.N1.1.bin.\n'
                                  '[####################] 100% -- SUCCESS\n\n'
                                  'Performing module support checks.\n'
                                  '[####################] 100% -- SUCCESS\n\n'
                                  'Notifying services about system upgrade.\n'
                                  '[####################] 100% -- SUCCESS\n\n'
                                  'Compatibility check is done:\n'
                                  'Module bootable         Impact  Install-type  Reason\n'
                                  '------ -------- --------------  ------------  ------\n'
                                  ' 1          yes     disruptive         reset  Incompatible image\n'
                                  ' 101        yes     disruptive         reset  Incompatible image\n\n'
                                  'Images will be upgraded according to following table:\n'
                                  'Module   Image       Running-Version    New-Version    Upg-Required\n'
                                  '------   -------     --------------    ------------    ------------\n'
                                  ' 1       system        5.2(1)N1(1a)     7.3(6)N1(1)           yes\n'
                                  ' 1     kickstart       5.2(1)N1(1a)     7.3(6)N1(1)           yes\n'
                                  ' 1         bios     v3.6.0(05/09/2012)  v3.6.0(05/09/201 2)    no\n'
                                  ' 1       SFP-uC            v1.1.0.0        v1.0.0.0            no\n'
                                  ' 101     fexth         5.2(1)N1(1a)      7.3(6)N1(1)           yes\n'
                                  ' 1     power-seq            v3.0            v3.0               no\n'
                                  ' 1         uC             v1.2.0.1         v1.2.0.1           no\n\n'
                                  'Switch will be reloaded for disruptive upgrade.\n'
                                  'Do you want to continue with the installation (y/n)? [n] y\n\n'
                                  'Install is in progress, please wait.\n\n'
                                  'Performing runtime checks.\n'
                                  '[####################] 100% -- SUCCESS\n\n'
                                  'Setting boot variables.\n'
                                  '[####################] 100% -- SUCCESS\n\n'
                                  'Performing configuration copy.\n'
                                  '[####################] 100% -- SUCCESS\n\n'
                                  'Pre-loading modules.\n'
                                  '[This step might take upto 20 minutes to complete - please wait.]\n'
                                  '[*Warning -- Please do not abort installation/reload or powercycle fexes*]\n'
                                  '[####################] 100% -- SUCCESS\n\n'
                                  'Finishing the upgrade, switch will reboot in 10 seconds.\n'
                                  '5K-1# Shutdown Ports..\n'
                                  'writing reset reason 49, \n')

        p = document.add_paragraph("\n")
        p.add_run("Step 9: ").bold = True
        p.add_run(' Verify that the upgrade of the vPC secondary switch is completed successfully. At the completion of the upgrade, the vPC secondary switch restores vPC peering')
        document.add_paragraph('Note: When the switch is reloaded, only singled-homed FEXs connected to the switch are reloaded and dual-homed FEXs are not reloaded\n' 
        'In case of dual- homed FEX, there is no need for issuing additional reload command from any switch. Automatically Dual-homed FEXs connect to both the peer switches and start forwarding traffic.')
        UpgradeProcedure.funTable('N5K-1# sh version\n'
                                  'Cisco Nexus Operating System (NX-OS) Software\n'
                                  'TAC support: http://www.cisco.com/tac\n'
                                  'Documents: \n'
                                  'http://www.cisco.com/en/US/products/ps9372/tsd_products_support_series_home .html\n'
                                  'Copyright (c) 2002-2019, Cisco Systems, Inc. All rights reserved.\n'
                                  'The copyrights to certain works contained herein are owned by\n'
                                  'other third parties and are used and distributed under license.\n'
                                  'Some parts of this software are covered under the GNU Public\n'
                                  'License. A copy of the license is available at\n'
                                  'http://www.gnu.org/licenses/gpl.html.\n\n'
                                  'Software\n'
                                  ' BIOS: version 3.6.0\n'
                                  ' Power Sequencer Firmware:\n'
                                  '           Module 1: v3.0\n'
                                  '           Module 2: v5.0\n'
                                  ' Microcontroller Firmware:  version v1.2.0.1\n'
                                  ' QSFP Microcontroller Firmware:\n'
                                  '           Module not detected\n'
                                  ' CXP Microcontroller Firmware:\n'
                                  '           Module not detected\n'
                                  ' kickstart: version 7.3(6)N1(1)\n'
                                  ' system:    version 7.3(6)N1(1)\n'
                                  ' BIOS compile time:        05/09/2012\n'
                                  ' kickstart image file is:  bootflash:///n5000-uk9-kickstart.7.3.6.N1.1.bin\n'
                                  ' kickstart compile time:   9/11/2019 23:00:00 [09/12/2019 07:56:00]\n'
                                  ' system image file is:     bootflash:///n5000-uk9.7.3.6.N1.1.bin\n'
                                  ' system compile time:      9/11/2019 23:00:00 [09/12/2019 12:10:27]\n\n'
                                  'Hardware\n'
                                  ' cisco Nexus5548 Chassis ("O2 32X10GE/Modular Universal Platform Supervisor")\n'
                                  ' Intel(R) Xeon(R) CPU with 8253792 kB of memory.\n'
                                  ' Processor Board ID FOC19151KY3\n\n'
                                  ' Device name: N5K-1\n'
                                  ' bootflash: 2007040 kB\n\n'
                                  'Kernel uptime is 0 day(s), 0 hour(s), 42 minute(s), 24 second(s)\n\n'
                                  'Last reset at 705024 usecs after Mon May 4 20:30:53 2020\n\n'
                                  'Reason: Disruptive upgrade\n'
                                  ' System version: 7.1(5)N1(1b)\n'
                                  ' Service:\n\n'
                                  'plugin\n'
                                  ' Core Plugin, Ethernet Plugin\n\n'
                                  'Active Package(s)\n\n')
        document.add_paragraph("\n")
        UpgradeProcedure.funTable('N5K-1# sh fex\n'
                                  ' FEX       FEX         FEX              FEX\n''\n'
                                  'Number Description    State      Model          Serial\n'
                                  '---------------------------------------------------------\n'
                                  '100        FEX0100    Online  N2K-C2248TP-1GE   SSI15310B1N\n'
                                  'sh fex detail snap below\n'
                                  'FEX version: 7.3(6)N1(1) [Switch version: 7.3(6)N1(1)]\n'
                                  ' FEX Interim version: 7.3(6)N1(1)\n'
                                  ' Switch Interim version: 7.3(6)N1(1)\n'
                                  'Logs:\n'
                                  '05/04/2020 18:04:55.582231: Module register received\n'
                                  '05/04/2020 20:40:02.881723: Module register received\n'
                                  '05/04/2020 20:40:02.893951: Registration response sent\n'
                                  '05/04/2020 20:40:02.917536: create module inserted event.\n'
                                  '05/04/2020 20:40:02.918619: Module Online Sequence\n'
                                  '05/04/2020 20:40:07.024766: Module Online\n\n')
        document.add_paragraph("\n")
        UpgradeProcedure.funTable('N5K-1# sh vpc Legend:\n'
                                  ' (*) - local vPC is down, forwarding via vPC peer-link\n'
                                  'vPC domain id                    : 100\n'
                                  'Peer status                      : peer adjacency formed ok\n'
                                  'vPC keep-alive status            : peer is alive\n'
                                  'Configuration consistency status : success\n'
                                  'Per-vlan consistency status      : success\n'
                                  'Type-2 consistency status        : success\n'
                                  'vPC role                         : secondary\n'
                                  'Number of vPCs configured        : 50\n'
                                  'Peer Gateway                     : Disabled\n'
                                  'Dual-active excluded VLANs       : -\n'
                                  'Graceful Consistency Check       : Enabled\n'
                                  'Operational Layer3 Peer-router   : Disabled\n'
                                  'Auto-recovery status             : Enabled (timeout = 240 seconds)\n\n'
                                  'vPC Peer-link status\n'
                                  '----------------------------------------------------------\n'
                                  'id  Port Status  Active vlans\n'
                                  '--  ---- ------ ------------------------------------------\n'
                                  '1   Po20   up      1,10,100\n\n'
                                  'vPC status\n'
                                  '-----------------------------------------------------------\n'
                                  'id    Port   Status Consistency Reason               Active vlans\n'
                                  '---- ------- ------ ----------- --------------------- ---------\n'
                                  '100   Po100     up   success     success                 -\n'
                                  '1000  Po1000   down*   Not       Consistency Check Not   -\n'
                                  '                       Applicable  Performed\n\n'
                                  'N5K-2# sh fex\n'''
                                  ' FEX    FEX            FEX           FEX\n'
                                  'Number  Description  State    Model          Serial\n'
                                  '-----------------------------------------------------------------\n'
                                  '100     FEX0100      Online  N2K-C2248TP-1GE  SSI15310B1N\n\m'
                                  'sh fex detail snap below\n'
                                  'FEX: 100 Description: FEX0100 state: Online\n'
                                  'FEX version: 7.3(6)N1(1) [Switch version: 7.3(6)N1(1)]\n\n'
                                  '  FEX Interim version: 7.3(6)N1(1)\n'
                                  '  Switch Interim version: 7.3(6)N1(1)\n'
                                  'Logs:\n'
                                  '05/04/2020 20:22:39.877144: Module register received\n'
                                  '05/04/2020 20:22:39.881018: Registration response sent\n'
                                  '05/04/2020 20:22:40.022366: create module inserted event.\n'
                                  '05/04/2020 20:22:40.023434: Module Online Sequence\n'
                                  '05/04/2020 20:22:46.602091: Module Online\n')
        UpgradeProcedure.funTable('N5K-2# sh vpc Legend:\n'
                                  ' (*) - local vPC is down, forwarding via vPC peer-link\n'
                                  'vPC domain id                    : 100\n'
                                  'Peer status                      : peer adjacency formed ok\n'
                                  'vPC keep-alive status            : peer is alive\n'
                                  'Configuration consistency status : success\n'
                                  'Per-vlan consistency status      : success\n'
                                  'Type-2 consistency status        : success\n'
                                  'vPC role                         : primary\n'
                                  'Number of vPCs configured        : 50\n'
                                  'Peer Gateway                     : Disabled\n'
                                  'Dual-active excluded VLANs       : -\n'
                                  'Graceful Consistency Check       : Enabled\n'
                                  'Operational Layer3 Peer-router   : Disabled\n'
                                  'Auto-recovery status             : Enabled (timeout = 240 seconds)\n\n'
                                  'vPC Peer-link status\n'
                                  '-----------------------------------------------------------------\n'
                                  'id  Port Status Active vlans \n'
                                  '--  ---- ------ -------------------------------------------------\n'
                                  '1   Po20  up     1,10,100 \n\n'
                                  'vPC status\n'
                                  '-----------------------------------------------------------------\n'
                                  'id     Port        Status Consistency Reason         Active vlans\n'
                                  '------ ----------- ------ ----------- -------------  -----------\n'
                                  '100    Po100       up     success     success              - \n'
                                  '1000   Po1000      down*  success     success              - \n\n')


        p = document.add_paragraph("\n")
        p.add_run("Step 10: ").bold = True
        p.add_run(' Verify the process has finished by issuing the commands')
        UpgradeProcedure.funTable('\t# show install all status\n'
                                  '\t\tThis is the log of last installation.\n'
                                  '\t\t...\n'
                                  '\t\tInstall has been successful.\n'
                                  '\t# show boot')
        p=document.add_paragraph()
        p.add_run("Note: ").bold=True
        document.add_paragraph('Issue “reload power-cycle” on required device/devices after both devices of the vpc-pair are upgraded successfully to a particular release if you see the following highlighted lines in output of \n'
        '“show install all status” command:')
        UpgradeProcedure.funTable('# reload power-cycle\n'
                                  'WARNING: There is unsaved configuration!!!\n'
                                  'WARNING: This command will reboot the system\n'
                                  'Do you want to continue? (y/n) [n] y\n'
                                  '[ 1750.957586] Shutdown Ports..\n'
                                  'Images will be upgraded according to following table:\n'
                                  'Module Image      Running-Version      New-Version   Upg-Required\n'
                                  '------ ------    -----------------  ---------------  -----------\n'
                                  '   1   system        7.3(6)N1(1)        7.3(6)N1(1)       no\n'
                                  '   1   kickstart     7.3(6)N1(1)        7.3(6)N1(1)       no\n'
                                  '   1   bios      v3.6.0(05/09/2012)   v6.0(05/09/2012)    no\n'
                                  '   1   SFP-uC             v1.0.0.0         v1.0.0.0       no\n'
                                  '  100  fexth            7.3(6)N1(1)      7.3(6)N1(1)      no\n'
                                  '   1   power-seq              v1.0              v3.0     yes\n'
                                  '   3   power-seq              v5.0              v5.0      no\n'
                                  '   1   uC v                 .2.0.1           1.2.0.1      no\n\n')

        try:
            r = document.add_paragraph().add_run()
            r.add_picture('note.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('The install all command is used for both ISSU and non-ISSU scenarios')
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found...")
        document.add_paragraph('Upgrades chassis with a single command',style='List Bullet 2')
        document.add_paragraph('Returns descriptive information about the intended changes to the system before you continue with the installation',style='List Bullet 2')
        document.add_paragraph('Provides a choice whether to install in a disruptive or non-disruptive (when available) manner',style='List Bullet 2')
        document.add_paragraph('Automatically checks the image integrity, including the running kickstart and system images. ',style='List Bullet 2')
        document.add_paragraph('Sets the kickstart and system boot variables',style='List Bullet 2')
        document.add_paragraph('The command performs a platform validity check to verify that a wrong image is not used.',style='List Bullet 2')

        p = document.add_paragraph("\n")
        p.add_run("Perform the Pre-Post Check:").bold = True
        document.add_paragraph('In order to verify the connectivity between the Nexus device and other services, we should make sure that there is no change in the below show commands from pre-post diff report after upgrading both the switches.')
        p = document.add_paragraph()
        p.add_run("\tshow interface status").bold = True
        p = document.add_paragraph()
        p.add_run("\tshow port-channel summary").bold = True
        p = document.add_paragraph()
        p.add_run("\tshow vsan").bold = True
        p = document.add_paragraph()
        p.add_run("\tshow fex").bold = True
        p = document.add_paragraph()
        p.add_run("\tshow vpc").bold = True
        p = document.add_paragraph()
        p.add_run("\tshow mac address-table").bold = True
        p = document.add_paragraph()
        p.add_run("\tshow ip arp").bold = True
        p = document.add_paragraph()
        p.add_run("\tshow running-config").bold = True
        p = document.add_paragraph()
        p.add_run("Important: Don’t proceed the closer of CR if there are any mis-match in the pre-post diff report.").bold = True
        p = document.add_paragraph()
        p.add_run("Vsan mismatch:").bold = True
        p = document.add_paragraph()
        p.add_run("If there is an issue in vsan database where the existing vsan values are replaced with default value 1, implement below commands to restore the vsan values.").bold = True
        p = document.add_paragraph()
        p.add_run("# conft").bold = True
        p = document.add_paragraph()
        p.add_run("# vsan database").bold = True
        p = document.add_paragraph()
        p.add_run("# vsan <no>").bold = True
        p = document.add_paragraph()
        p.add_run("# vsan <no> interface <int_name>").bold = True
        p = document.add_paragraph()
        p.add_run("#copy running-config starup-config").bold = True

        p= document.add_paragraph("\n\n")
        p.add_run("Task Success Criteria").bold = True
        document.add_paragraph('The upgrade procedure completes successfully.')

        p = document.add_paragraph("\n")
        p.add_run("Failure Procedure").bold = True
        document.add_paragraph('\t1. Resolve the problem, if one discovered, and if possible.\n'
                               '\t2. Call Cisco TAC for support \n'
                               '\t3. Proceed to Rollback procedure')
        upgrade2 = (
            ('Task Number', '3'),
            ('Task Description ', 'Verification'),
            ('Task Scheduled Start', ''),
            ('Task Duration','15 minutes'),
            ('Task Owner Pri/Sec',''),
            ('Task Dependencies','Network is operational,OOB/Console access to the device')
        )
        table = document.add_table(rows=1, cols=2, style='Colorful Grid Accent 1')
        for term, definition in upgrade1:
            row_cells = table.add_row().cells
            row_cells[0].text = str(term)
            row_cells[1].text = str(definition)


        p = document.add_paragraph("\n")
        p.add_run("Step 1 ").bold = True
        p.add_run('Gain SSH access to the device:')

        p = document.add_paragraph("\n")
        p.add_run("Step 2 ").bold = True
        p.add_run('Verify new software version')
        UpgradeProcedure.funTable('\t# show install all status\n'
                                  '\t# show version \n'
                                  '\t# show fex <fex_ID> version\n'
                                  '\t# show boot \n'
                                  '\t# show system vlan reserved'
                                  )

        p = document.add_paragraph("\n")
        p.add_run("Step 3 ").bold = True
        p.add_run('Execute basic verification commands to check hardware and operational status')
        UpgradeProcedure.funTable('\t# show module\n'
                                  '\t# show running-config \n'
                                  '\t# show ip interface brief include-secondary vrf all\n'
                                  '\t# show  port-channel summary \n')

        p = document.add_paragraph("\n")
        p.add_run("Step 4 ").bold = True
        p.add_run('Check fex pre-provisioning')
        UpgradeProcedure.funTable('\t# show provision failed-config <FEX-ID>\n')
        try:
            r = document.add_paragraph().add_run()
            r.add_picture('caution.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('Response should contain “Config was applied without any errors.”')
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found...")

        p = document.add_paragraph("\n")
        p.add_run("Step 5 ").bold = True
        p.add_run('Execute all other pre-/-post checks as defined in Table 5')
        try:
            r = document.add_paragraph().add_run()
            r.add_picture('caution.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('The above results from Step 5 should be exact matches as per last column in Table 5 re-check show commands list equal the Pre install show results, if any of the results is not equal a major alarm should be raised and the next steps must be postponed until a local engineer has verified the impact occurred on the upgraded device.')
            r = document.add_paragraph().add_run()
            r.add_picture('note.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('Type-2 inconsistency reason will show VTP type-2 configuration incompatible in show vpc output while only the VPC Primary has been upgraded. Once both peers have been upgraded this Type-2 inconsistency will be removed.')
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found...")
        

        p = document.add_paragraph("\n\n")
        p.add_run("Task Success Criteria").bold = True
        document.add_paragraph('Device (supervisors and modules) is operational.\n'
        'All interfaces have resumed operation in the expected state.\n'
        'Configuration of the device has the same lines as prior to the upgrade\n'
        'All Protocols are operational, and their parameters and values are the same as before the upgrade.')

        p = document.add_paragraph()
        p.add_run("Failure Procedure").bold = True
        document.add_paragraph('\t1. Resolve the problem, if one discovered, and if possible.\n'
                               '\t2. Call Cisco TAC for support \n'
                               '\t3.  Proceed to Rollback procedure')
        document.add_page_break()

        return
class InstallAnalysis:
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def installAnalysis_ASR_1000_17_03_03(self):
        document.add_paragraph("Certain requirements need to be met before deploying an image onto a device. The Install Analysis phase is a Pass or Fail checklist before the image can be uploaded to a device. This reduces any potential risks that may cause an install failure.\n")
        r = document.add_paragraph().add_run()
        try:
            r.add_picture('caution.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(f'Verify the Flash: or Bootflashmemory have enough space to keep the new {self.os} version before to proceed {self.os} upgrade.')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")

    def installAnalysis_Switch_Catalyst_9600_17_06_02(self):
        document.add_paragraph("Certain requirements need to be met before deploying an image onto a device. The Install Analysis phase is a Pass or Fail checklist before the image can be uploaded to a device. This reduces any potential risks that may cause an install failure.\n")
        r = document.add_paragraph().add_run()
        try:
            r.add_picture('caution.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(f'Verify the Flash: or Bootflashmemory have enough space to keep the new {self.os} version before to proceed {self.os} upgrade.')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")

    def installAnalysis_Switch_Catalyst_9600_17_03_05(self):
        document.add_paragraph("Certain requirements need to be met before deploying an image onto a device. The Install Analysis phase is a Pass or Fail checklist before the image can be uploaded to a device. This reduces any potential risks that may cause an install failure.\n")
        r = document.add_paragraph().add_run()
        try:
            r.add_picture('caution.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(f'Verify the Flash: or Bootflashmemory have enough space to keep the new {self.os} version before to proceed {self.os} upgrade.')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
        return
    def installAnalysis_Switch_Catalyst_9300_17_03_04(self):
        document.add_paragraph("Certain requirements need to be met before deploying an image onto a device. The Install Analysis phase is a Pass or Fail checklist before the image can be uploaded to a device. This reduces any potential risks that may cause an install failure.\n")
        r = document.add_paragraph().add_run()
        try:
            r.add_picture('caution.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(f'Verify the Flash: or Bootflashmemory have enough space to keep the new {self.os} version before to proceed {self.os} upgrade.')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
        return
    def installAnalysis_Switch_Catalyst_9500_17_03_04(self):
        document.add_paragraph("Certain requirements need to be met before deploying an image onto a device. The Install Analysis phase is a Pass or Fail checklist before the image can be uploaded to a device. This reduces any potential risks that may cause an install failure.\n")
        r = document.add_paragraph().add_run()
        try:
            r.add_picture('caution.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text(f'Verify the Flash: or Bootflashmemory have enough space to keep the new {self.os} version before to proceed {self.os} upgrade.')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
        return

#class for 2.1.1 Verify the mode of IOS version
class IosVersion(Formatter):
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def iosVersion_ASR_1000_17_03_03(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/routers/asr1000/install/guide/1001-x/asr1hig-book/pwr_up_init_configuartion.html?referring_site=RE&pos=1&page=https://www.cisco.com/c/en/us/td/docs/routers/asr1000/quick/start/guide/asr1_qs1.html')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#con_1111409 > section > pre:nth-child(11) > code')
            IosVersion.funTable(elem[0].text)
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'IOS Version' section.")
        document.add_paragraph('\n')
        IosVersion.funTable('Router#show file system\n'
                                      'File Systems:\n'
                                      '       Size(b)       Free(b)      Type  Flags  Prefixes\n'
                                      '             -             -    opaque     rw   system:\n'
                                      '             -             -    opaque     rw   tmpsys:\n'
                                      '*  16420106240   15099944960      disk     rw   bootflash: flash\n'
                                      '    1524695040    1445007360      disk     ro   webui:\n'
                                      '             -             -    opaque     rw   null:\n'
                                      '             -             -    opaque     ro   tar:\n'
                                      '             -             -   network     rw   tftp:\n'
                                      '      33554432      33551308     nvram     rw   nvram:\n'
                                      '             -             -    opaque     wo   syslog:\n'
                                      '             -             -   network     rw   rcp:\n'
                                      '             -             -   network     rw   pram:\n'
                                      '             -             -   network     rw   http:\n'
                                      '             -             -   network     rw   ftp:\n'
                                      '             -             -   network     rw   scp:\n'
                                      '             -             -   network     rw   sftp:\n'
                                      '             -             -   network     rw   https:\n'
                                      '             -             -    opaque     ro   cns:\n')
    def iosVersion_Switch_Catalyst_9600_17_06_02(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-6/release_notes/ol-17-6-9600.html#concept_ycv_jdf_3mb')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(6) > td:nth-child(2) > div > div:nth-child(3) > pre > code')
            IosVersion.funTable(elem[0].text.strip())
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'IOS Version' section.")

    def iosVersion_Switch_Catalyst_9600_17_03_05(self):
        IosVersion.funTable('Switch#show version \n'
                            'Cisco IOS XE Software, Version 16.12.04 \n'
                            'Cisco IOS Software [Gibraltar], Catalyst L3 Switch Software (CAT9K_IOSXE), Version 16.12.4, RELEASE SOFTWARE (fc5) \n'
                            'Technical Support: http://www.cisco.com/techsupport \n'
                            'Copyright (c) 1986-2020 by Cisco Systems, Inc. \n'
                            'Compiled Thu 09-Jul-20 21:49 by mcpre \n'
                            ' \n'
                            ' \n'
                            'Cisco IOS-XE software, Copyright (c) 2005-2020 by cisco Systems, Inc. \n'
                            'All rights reserved. Certain components of Cisco IOS-XE software are \n'
                            'licensed under the GNU General Public License ("GPL") Version 2.0. The \n'
                            'software code licensed under GPL Version 2.0 is free software that comes \n'
                            'with ABSOLUTELY NO WARRANTY. You can redistribute and/or modify such \n'
                            'GPL code under the terms of GPL Version 2.0. For more details, see the \n'
                            'documentation or "License Notice" file accompanying the IOS-XE software, \n'
                            'or the applicable URL provided on the flyer accompanying the IOS-XE \n'
                            'software. \n'
                            ' \n'
                            ' \n'
                            'ROM: IOS-XE ROMMON \n'
                            'BOOTLDR: System Bootstrap, Version 17.1.1[FC2], RELEASE SOFTWARE (P) \n'
                            ' \n'
                            'Switch uptime is 46 weeks, 4 days, 37 minutes \n'
                            'Uptime for this control processor is 46 weeks, 4 days, 39 minutes \n'
                            'System returned to ROM by PowerOn \n'
                            'System restarted at 23:48:27 cdt Wed Apr 14 2021 \n'
                            'System image file is "bootflash:packages.conf" \n'
                            'Last reload reason: PowerOn \n'
                            ' \n'
                            ' \n'
                            ' \n'
                            'This product contains cryptographic features and is subject to United \n'
                            'States and local country laws governing import, export, transfer and \n'
                            'use. Delivery of Cisco cryptographic products does not imply \n'
                            'third-party authority to import, export, distribute or use encryption. \n'
                            'Importers, exporters, distributors and users are responsible for \n'
                            'compliance with U.S. and local country laws. By using this product you \n'
                            'agree to comply with applicable laws and regulations. If you are unable \n'
                            'to comply with U.S. and local laws, return this product immediately. \n'
                            ' \n'
                            'A summary of U.S. laws governing Cisco cryptographic products may be found at: \n'
                            'http://www.cisco.com/wwl/export/crypto/tool/stqrg.html \n'
                            ' \n'
                            'If you require further assistance please contact us by sending email to \n'
                            'export@cisco.com. \n'
                            ' \n'
                            ' \n'
                            'Technology Package License Information: \n'
                            ' \n'
                            '------------------------------------------------------------------------------ \n'
                            'Technology-package Technology-package \n'
                            'Current Type Next reboot \n'
                            '------------------------------------------------------------------------------ \n'
                            'network-advantage Smart License network-advantage \n'
                            'dna-advantage Subscription Smart License dna-advantage \n'
                            'AIR License Level: AIR DNA Advantage \n'
                            'Next reload AIR license Level: AIR DNA Advantage \n'
                            ' \n'
                            ' \n'
                            'Smart Licensing Status: UNREGISTERED/EVAL EXPIRED \n'
                            ' \n'
                            'cisco C9606R (X86) processor (revision V01) with 1867864K/6147K bytes of memory. \n'
                            'Processor board ID FXS2423Q75Y \n'
                            '100 Virtual Ethernet interface \n'
                            '24 Forty/Hundred Gigabit Ethernet interfaces \n'
                            '96 TwentyFive Gigabit Ethernet interfaces \n'
                            '32768K bytes of non-volatile configuration memory. \n'
                            '16009160K bytes of physical memory. \n'
                            '11161600K bytes of Bootflash at bootflash:. \n'
                            '1638400K bytes of Crash Files at crashinfo:. \n'
                            '0K bytes of WebUI ODM Files at webui:. \n'
                            ' \n'
                            'Base Ethernet MAC Address : 3c:57:31:05:08:00 \n'
                            'Motherboard Assembly Number : 4C57 \n'
                            'Motherboard Serial Number : FXS242101JU \n'
                            'Model Revision Number : V02 \n'
                            'Motherboard Revision Number : 4 \n'
                            'Model Number : C9606R \n'
                            'System Serial Number : FXS2423Q75Y \n')
        return
    def iosVersion_Switch_Catalyst_9300_17_03_04(self):
        IosVersion.funTable('Switch#show version\n'
                            'Cisco IOS XE Software, Version 16.12.04\n'
                            'Cisco IOS Software [Amsterdam], Catalyst L3 Switch Software (CAT9K_IOSXE), Version 16.12.4, RELEASE SOFTWARE (fc7)\n'
                            'Technical Support: http://www.cisco.com/techsupport\n'
                            'Copyright (c) 1986-2021 by Cisco Systems, Inc.\n'
                            'Compiled Thu 04-Mar-21 12:32 by mcpre\n'
                            '\n'
                            '\n'
                            'Cisco IOS-XE software, Copyright (c) 2005-2021 by cisco Systems, Inc.\n'
                            'All rights reserved. Certain components of Cisco IOS-XE software are\n'
                            'licensed under the GNU General Public License ("GPL") Version 2.0. The\n'
                            'software code licensed under GPL Version 2.0 is free software that comes\n'
                            'with ABSOLUTELY NO WARRANTY. You can redistribute and/or modify such\n'
                            'GPL code under the terms of GPL Version 2.0. For more details, see the\n'
                            'documentation or "License Notice" file accompanying the IOS-XE software,\n'
                            'or the applicable URL provided on the flyer accompanying the IOS-XE\n'
                            'software.\n'
                            '\n'
                            '\n'
                            'ROM: IOS-XE ROMMON\n'
                            'BOOTLDR: System Bootstrap, Version 17.5.2r, RELEASE SOFTWARE (P)\n'
                            '\n'
                            'Switch uptime is 0 minutes\n'
                            'Uptime for this control processor is 2 minutes\n'
                            'System returned to ROM by Image Install at 06:40:32 EDT Wed Mar 30 2022\n'
                            'System image file is "flash:packages.conf"\n'
                            'Last reload reason: Image Install\n'
                            '\n'
                            '\n'
                            '\n'
                            'This product contains cryptographic features and is subject to United\n'
                            'States and local country laws governing import, export, transfer and\n'
                            'use. Delivery of Cisco cryptographic products does not imply\n'
                            'third-party authority to import, export, distribute or use encryption.\n'
                            'Importers, exporters, distributors and users are responsible for\n'
                            'compliance with U.S. and local country laws. By using this product you\n'
                            'agree to comply with applicable laws and regulations. If you are unable\n'
                            'to comply with U.S. and local laws, return this product immediately.\n'
                            '\n'
                            'A summary of U.S. laws governing Cisco cryptographic products may be found at:\n'
                            'http://www.cisco.com/wwl/export/crypto/tool/stqrg.html\n'
                            '\n'
                            'If you require further assistance please contact us by sending email to\n'
                            'export@cisco.com.\n'
                            '\n'
                            '\n'
                            'Technology Package License Information:\n'
                            '\n'
                            '------------------------------------------------------------------------------\n'
                            'Technology-package Technology-package\n'
                            'Current Type Next reboot\n'
                            '------------------------------------------------------------------------------\n'
                            'network-advantage Smart License network-advantage\n'
                            'dna-advantage Subscription Smart License dna-advantage\n'
                            'AIR License Level: AIR DNA Advantage\n'
                            'Next reload AIR license Level: AIR DNA Advantage\n'
                            '\n'
                            '\n'
                            'Smart Licensing Status: Registration Not Applicable/Not Applicable\n'
                            '\n'
                            'cisco C9300-48U (X86) processor with 1331521K/6147K bytes of memory.\n'
                            'Processor board ID FOC2428LAPN\n'
                            '1 Virtual Ethernet interface\n'
                            '52 Gigabit Ethernet interfaces\n'
                            '8 Ten Gigabit Ethernet interfaces\n'
                            '2 TwentyFive Gigabit Ethernet interfaces\n'
                            '2 Forty Gigabit Ethernet interfaces\n'
                            '2048K bytes of non-volatile configuration memory.\n'
                            '8388608K bytes of physical memory.\n'
                            '1638400K bytes of Crash Files at crashinfo:.\n'
                            '11264000K bytes of Flash at flash:.\n'
                            '\n'
                            'Base Ethernet MAC Address : 3c:13:cc:19:05:80\n'
                            'Motherboard Assembly Number : 73-18275-05\n'
                            'Motherboard Serial Number : FOC24276PWK\n'
                            'Model Revision Number : A0\n'
                            'Motherboard Revision Number : A0\n'
                            'Model Number : C9300-48U\n'
                            'System Serial Number : FOC2428LAPN\n'
                            'CLEI Code Number :\n'
                            '\n'
                            '\n'
                            'Switch Ports Model SW Version SW Image Mode\n'
                            '------ ----- ----- ---------- ---------- ----\n'
                            '* 1 65 C9300-48U 16.12.04 CAT9K_IOSXE INSTALL\n'
                            '\n'
                            '\n'
                            'Configuration register is 0x102\n')
    def iosVersion_Switch_Catalyst_9500_17_03_04(self):
        IosVersion.funTable('Switch#show version\n'
                            'Cisco IOS XE Software, Version 16.12.04\n'
                            'Cisco IOS Software [Amsterdam], Catalyst L3 Switch Software (CAT9K_IOSXE), Version 16.12.4, RELEASE SOFTWARE (fc2)\n'
                            'Technical Support: http://www.cisco.com/techsupport\n'
                            'Copyright (c) 1986-2022 by Cisco Systems, Inc.\n'
                            'Compiled Wed 09-Feb-22 10:41 by mcpre\n'
                            '\n'
                            '\n'
                            'Cisco IOS-XE software, Copyright (c) 2005-2022 by cisco Systems, Inc.\n'
                            'All rights reserved. Certain components of Cisco IOS-XE software are\n'
                            'licensed under the GNU General Public License ("GPL") Version 2.0. The\n'
                            'software code licensed under GPL Version 2.0 is free software that comes\n'
                            'with ABSOLUTELY NO WARRANTY. You can redistribute and/or modify such\n'
                            'GPL code under the terms of GPL Version 2.0. For more details, see the\n'
                            'documentation or "License Notice" file accompanying the IOS-XE software,\n'
                            'or the applicable URL provided on the flyer accompanying the IOS-XE\n'
                            'software.\n'
                            '\n'
                            '\n'
                            'ROM: IOS-XE ROMMON\n'
                            'BOOTLDR: System Bootstrap, Version 16.12.2r, RELEASE SOFTWARE (P)\n'
                            '\n'
                            'Switch uptime is 10 hours, 34 minutes\n'
                            'Uptime for this control processor is 10 hours, 35 minutes\n'
                            'System returned to ROM by Reload Command\n'
                            'System image file is "flash:packages.conf"\n'
                            'Last reload reason: Reload Command\n'
                            '\n'
                            '\n'
                            '\n'
                            'This product contains cryptographic features and is subject to United\n'
                            'States and local country laws governing import, export, transfer and\n'
                            'use. Delivery of Cisco cryptographic products does not imply\n'
                            'third-party authority to import, export, distribute or use encryption.\n'
                            'Importers, exporters, distributors and users are responsible for\n'
                            'compliance with U.S. and local country laws. By using this product you\n'
                            'agree to comply with applicable laws and regulations. If you are unable\n'
                            'to comply with U.S. and local laws, return this product immediately.\n'
                            '\n'
                            'A summary of U.S. laws governing Cisco cryptographic products may be found at:\n'
                            'http://www.cisco.com/wwl/export/crypto/tool/stqrg.html\n'
                            '\n'
                            'If you require further assistance please contact us by sending email to\n'
                            'export@cisco.com.\n'
                            '\n'
                            '\n'
                            'Technology Package License Information:\n'
                            '\n'
                            '------------------------------------------------------------------------------\n'
                            'Technology-package Technology-package\n'
                            'Current Type Next reboot\n'
                            '------------------------------------------------------------------------------\n'
                            'network-advantage Smart License network-advantage\n'
                            'dna-advantage Subscription Smart License dna-advantage\n'
                            'AIR License Level: AIR DNA Advantage\n'
                            'Next reload AIR license Level: AIR DNA Advantage\n'
                            '\n'
                            '\n'
                            'Smart Licensing Status: Registration Not Applicable/Not Applicable\n'
                            '\n'
                            'cisco C9500-16X (X86) processor with 1331284K/6147K bytes of memory.\n'
                            'Processor board ID FOC2504LC2L\n'
                            '1 Virtual Ethernet interface\n'
                            '24 Ten Gigabit Ethernet interfaces\n'
                            '2 Forty Gigabit Ethernet interfaces\n'
                            '2048K bytes of non-volatile configuration memory.\n'
                            '16777216K bytes of physical memory.\n'
                            '1638400K bytes of Crash Files at crashinfo:.\n'
                            '11264000K bytes of Flash at flash:.\n'
                            '\n'
                            'Base Ethernet MAC Address : bc:e7:12:5d:0f:00\n'
                            'Motherboard Assembly Number : 73-18709-01\n'
                            'Motherboard Serial Number : FOC25032TXH\n'
                            'Model Revision Number : E0\n'
                            'Motherboard Revision Number : C0\n'
                            'Model Number : C9500-16X\n'
                            'System Serial Number : FOC2504LC2L\n'
                            'CLEI Code Number :\n'
                            '\n'
                            '\n'
                            'Switch Ports Model SW Version SW Image Mode\n'
                            '------ ----- ----- ---------- ---------- ----\n'
                            '* 1 26 C9500-16X 16.12.04 CAT9K_IOSXE BUNDLE\n')

class ConfigurationRegister(Formatter):
    def __init__(self,os, version):
        self.os = os
        self.version = version
    def configurationRegister_Switch_Catalyst_9600_17_06_02(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-6/command_reference/b_176_9600_cr/system_management_commands.html#wp2862294214')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#wp2862294214__EXAMPLE_522ECF84B2BC468EA580E7EAA26862CF > pre:nth-child(3) > code')
            ConfigurationRegister.funTable(elem[0].text.strip())
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Verify the Configuration Register' section")

#AvailableDram
class AvailableDram(Formatter):
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def availableDram_ASR_1000_17_03_03(self):
        return AvailableDram.funTable('Router#show file system\n'
                                      'File Systems:\n'
                                      '       Size(b)       Free(b)      Type  Flags  Prefixes\n'
                                      '             -             -    opaque     rw   system:\n'
                                      '             -             -    opaque     rw   tmpsys:\n'
                                      '*  16420106240   15099944960      disk     rw   bootflash: flash\n'
                                      '    1524695040    1445007360      disk     ro   webui:\n'
                                      '             -             -    opaque     rw   null:\n'
                                      '             -             -    opaque     ro   tar:\n'
                                      '             -             -   network     rw   tftp:\n'
                                      '      33554432      33551308     nvram     rw   nvram:\n'
                                      '             -             -    opaque     wo   syslog:\n'
                                      '             -             -   network     rw   rcp:\n'
                                      '             -             -   network     rw   pram:\n'
                                      '             -             -   network     rw   http:\n'
                                      '             -             -   network     rw   ftp:\n'
                                      '             -             -   network     rw   scp:\n'
                                      '             -             -   network     rw   sftp:\n'
                                      '             -             -   network     rw   https:\n'
                                      '             -             -    opaque     ro   cns:\n')
    def availableDram_Switch_Catalyst_9600_17_06_02(self):
        document.add_paragraph('The below command displays the amount space available within flash. When copying over the new image to the device this is where the image is kept. It may be necessary to delete existing files from the device to free up more space. Use the show file system privileged EXEC command to list all file systems')
        AvailableDram.funTable('Switch# show file systems\n'
                               'File Systems:\n'
                               '       Size(b)       Free(b)      Type  Flags  Prefixes\n'
                               '             -             -    opaque     rw   system:\n'
                               '             -             -    opaque     rw   tmpsys:\n'
                               '    1651314688    1320402944      disk     rw   crashinfo:\n'
                               ' * 11353194496    9832026112      disk     rw   flash: bootflash:\n'
                               '    8062349312    7948447744      disk     ro   webui:\n'
                               '             -             -      opaque   rw   null:\n'
                               '             -             -      opaque   ro   tar:\n'
                               '             -             -      network  rw   tftp:\n'
                               '       2097152       1758064      nvram    rw   nvram:\n'
                               '             -             -      network  rw   rcp:\n'
                               '             -             -      network  rw   http:\n'
                               '             -             -      network  rw   ftp:\n'
                               '             -             -      network  rw   scp:\n'
                               '             -             -      network  rw   sftp:\n'
                               '             -             -      network  rw   https:\n'
                               '             -             -      opaque   ro   cns:\n'
                               'Switch#')
        return
    def availableDram_Switch_Catalyst_9600_17_03_05(self):
        AvailableDram.funTable('Switch#show version\n'
                               'Cisco IOS XE Software, Version 16.12.04\n'
                               'Cisco IOS Software [Gibraltar], Catalyst L3 Switch Software (CAT9K_IOSXE), Version 16.12.4, RELEASE SOFTWARE (fc5)\n'
                               'Technical Support: http://www.cisco.com/techsupport\n'
                               'Copyright (c) 1986-2020 by Cisco Systems, Inc.\n'
                               'Compiled Thu 09-Jul-20 21:49 by mcpre\n'
                               '\n'
                               '\n'
                               'Cisco IOS-XE software, Copyright (c) 2005-2020 by cisco Systems, Inc.\n'
                               'All rights reserved. Certain components of Cisco IOS-XE software are\n'
                               'licensed under the GNU General Public License ("GPL") Version 2.0. The\n'
                               'software code licensed under GPL Version 2.0 is free software that comes\n'
                               'with ABSOLUTELY NO WARRANTY. You can redistribute and/or modify such\n'
                               'GPL code under the terms of GPL Version 2.0. For more details, see the\n'
                               'documentation or "License Notice" file accompanying the IOS-XE software,\n'
                               'or the applicable URL provided on the flyer accompanying the IOS-XE\n'
                               'software.\n'
                               '\n'
                               '\n'
                               'ROM: IOS-XE ROMMON\n'
                               'BOOTLDR: System Bootstrap, Version 17.1.1[FC2], RELEASE SOFTWARE (P)\n'
                               '\n'
                               'Switch uptime is 46 weeks, 4 days, 37 minutes\n'
                               'Uptime for this control processor is 46 weeks, 4 days, 39 minutes\n'
                               'System returned to ROM by PowerOn\n'
                               'System restarted at 23:48:27 cdt Wed Apr 14 2021\n'
                               'System image file is "bootflash:packages.conf"\n'
                               'Last reload reason: PowerOn\n'
                               '\n'
                               '\n'
                               '\n'
                               'This product contains cryptographic features and is subject to United\n'
                               'States and local country laws governing import, export, transfer and\n'
                               'use. Delivery of Cisco cryptographic products does not imply\n'
                               'third-party authority to import, export, distribute or use encryption.\n'
                               'Importers, exporters, distributors and users are responsible for\n'
                               'compliance with U.S. and local country laws. By using this product you\n'
                               'agree to comply with applicable laws and regulations. If you are unable\n'
                               'to comply with U.S. and local laws, return this product immediately.\n'
                               '\n'
                               'A summary of U.S. laws governing Cisco cryptographic products may be found at:\n'
                               'http://www.cisco.com/wwl/export/crypto/tool/stqrg.html\n'
                               '\n'
                               'If you require further assistance please contact us by sending email to\n'
                               'export@cisco.com.\n'
                               '\n'
                               '\n'
                               'Technology Package License Information:\n'
                               '\n'
                               '------------------------------------------------------------------------------\n'
                               'Technology-package Technology-package\n'
                               'Current Type Next reboot\n'
                               '------------------------------------------------------------------------------\n'
                               'network-advantage Smart License network-advantage\n'
                               'dna-advantage Subscription Smart License dna-advantage\n'
                               'AIR License Level: AIR DNA Advantage\n'
                               'Next reload AIR license Level: AIR DNA Advantage\n'
                               '\n'
                               '\n'
                               'Smart Licensing Status: UNREGISTERED/EVAL EXPIRED\n'
                               '\n'
                               'cisco C9606R (X86) processor (revision V01) with 1867864K/6147K bytes of memory.\n'
                               'Processor board ID FXS2423Q75Y\n'
                               '100 Virtual Ethernet interface\n'
                               '24 Forty/Hundred Gigabit Ethernet interfaces\n'
                               '96 TwentyFive Gigabit Ethernet interfaces\n'
                               '32768K bytes of non-volatile configuration memory.\n'
                               '16009160K bytes of physical memory.\n'
                               '11161600K bytes of Bootflash at bootflash:.\n'
                               '1638400K bytes of Crash Files at crashinfo:.\n'
                               '0K bytes of WebUI ODM Files at webui:.\n'
                               '\n'
                               'Base Ethernet MAC Address : 3c:57:31:05:08:00\n'
                               'Motherboard Assembly Number : 4C57\n'
                               'Motherboard Serial Number : FXS242101JU\n'
                               'Model Revision Number : V02\n'
                               'Motherboard Revision Number : 4\n'
                               'Model Number : C9606R\n'
                               'System Serial Number : FXS2423Q75Y\n')
        return
    def availableDram_Switch_Catalyst_9300_17_03_04(self):
        document.add_paragraph('The below command displays the amount space available within flash. When copying over the new image to the device this is where the image is kept. It may be necessary to delete existing files from the device to free up more space. Use the show file system privileged EXEC command to list all file systems\n')
        AvailableDram.funTable('Switch#show file systems\n'
                               'File Systems:\n'
                               '\n'
                               'Size(b) Free(b) Type Flags Prefixes\n'
                               '- - opaque rw system:\n'
                               '- - opaque rw tmpsys:\n'
                               '1651314688 1107984384 disk rw crashinfo:\n'
                               '* 11353194496 7791243264 disk rw flash: bootflash:\n'
                               '3840733184 3726831616 disk ro webui:\n'
                               '- - opaque rw null:\n'
                               '- - opaque ro tar:\n'
                               '- - network rw tftp:\n'
                               '2097152 2074635 nvram rw nvram:\n'
                               '- - opaque wo syslog:\n'
                               '- - network rw rcp:\n'
                               '- - network rw http:\n'
                               '- - network rw ftp:\n'
                               '- - network rw scp:\n'
                               '- - network rw sftp:\n'
                               '- - network rw https:\n'
                               '- - opaque ro cns:\n')
    def availableDram_Switch_Catalyst_9500_17_03_04(self):
        document.add_paragraph('The below command displays the amount space available within flash. When copying over the new image to the device this is where the image is kept. It may be necessary to delete existing files from the device to free up more space. Use the show file system privileged EXEC command to list all file systems\n')
        AvailableDram.funTable('Switch#show file system\n'
                               'File Systems:\n'
                               '\n'
                               'Size(b) Free(b) Type Flags Prefixes\n'
                               '- - opaque rw system:\n'
                               '- - opaque rw tmpsys:\n'
                               '1651314688 1501265920 disk rw crashinfo:\n'
                               '* 11353194496 7187222528 disk rw flash: bootflash:\n'
                               '8062414848 7948283904 disk ro webui:\n'
                               '- - opaque rw null:\n'
                               '- - opaque ro tar:\n'
                               '- - network rw tftp:\n'
                               '2097152 2025602 nvram rw nvram:\n'
                               '- - opaque wo syslog:\n'
                               '- - network rw rcp:\n'
                               '- - network rw http:\n'
                               '- - network rw ftp:\n'
                               '- - network rw scp:\n'
                               '- - network rw sftp:\n'
                               '- - network rw https:\n'
                               '- - opaque ro cns:\n')
class ImageDeploymentAndValidation(Formatter):
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def imageDeploymentAndValidation_ASR_1000_17_03_03(self):
        document.add_paragraph("The Image Deployment is when the NCE engineer has verified that the device/devices have passed their relevant analysis checks and are ready to perform an image deployment onto the device.")
        p = document.add_paragraph('Log onto:', style='List Number 2')
        ImageDeploymentAndValidation.add_hyperlink(p, 'https://software.cisco.com/download/navigator.html?mode=home',' https://software.cisco.com/download/navigator.html?mode=home ')
        document.add_paragraph('Go to software download page and map to Download Home and map to the relevant image for the upgrade ', style='List Number 2')
        document.add_paragraph('Take a note of the MD5 checksum and download the new image file. ',   style='List Number 2')
        document.add_paragraph('Run the below command: ',   style='List Number 2')
        ImageDeploymentAndValidation.funTable('Router# verify /md5 flash: asr1000rpx86-universalk9.17.03.03.SPA.bin'
                                                '...................................................................................'
                                                '...........................................'
                                                '........................................'
                                                '...................................................................................'
                                                '................................................'
                                                '................................... '
                                                '...............................................'
                                                '...................................'
                                                '................................ ...........Done! \n'
                                                'verify /md5 (flash: asr1000rpx86-universalk9.17.03.03.SPA.bin) =08c4c732e70ae272980fbbe092d8debb')
        document.add_page_break()

    def imageDeploymentAndValidation_Switch_Catalyst_9600_17_06_02(self):
        document.add_paragraph("The Image Deployment is when the NCE engineer has verified that the device/devices have passed their relevant analysis checks and are ready to perform an image deployment onto the device.")
        p = document.add_paragraph('Log onto:', style='List Number 2')
        ImageDeploymentAndValidation.add_hyperlink(p, 'https://software.cisco.com/download/navigator.html?mode=home',' https://software.cisco.com/download/navigator.html?mode=home ')
        document.add_paragraph('Go to software download page and map to Download Home and map to the relevant image for the upgrade ', style='List Number 2')
        document.add_paragraph('Take a note of the MD5 checksum and download the new image file. ',   style='List Number 2')
        document.add_paragraph('Run the below command: ',   style='List Number 2')
        ImageDeploymentAndValidation.funTable('Switch# verify /md5 flash:cat9k_iosxe.17.03.03.SPA.bin\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '................................\n'
                                              '...........Done!\n'
                                              'verify /md5 (flash:cat9k_iosxe.17.03.03.SPA.bin) = 89c98b1ed44cf6cb1190eca977edb9a5\n')
        document.add_page_break()
        return
    def imageDeploymentAndValidation_Switch_Catalyst_9600_17_03_05(self):
        document.add_paragraph("The Image Deployment is when the NCE engineer has verified that the device/devices have passed their relevant analysis checks and are ready to perform an image deployment onto the device.")
        p = document.add_paragraph('Log onto:', style='List Number 2')
        ImageDeploymentAndValidation.add_hyperlink(p, 'https://software.cisco.com/download/navigator.html?mode=home',' https://software.cisco.com/download/navigator.html?mode=home ')
        document.add_paragraph('Go to software download page and map to Download Home and map to the relevant image for the upgrade ', style='List Number 2')
        document.add_paragraph('Take a note of the MD5 checksum and download the new image file. ',   style='List Number 2')
        document.add_paragraph('Run the below command: ',   style='List Number 2')
        ImageDeploymentAndValidation.funTable('Switch# verify /md5 flash:cat9k_iosxe.17.03.05.SPA.bin\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '................................\n'
                                              '...........Done!\n'
                                              'verify /md5 (flash:cat9k_iosxe.17.03.05.SPA.bin) = 62dcae59c73eb99aa54969b4da5a7c84\n')
        document.add_page_break()
        return
    def imageDeploymentAndValidation_Switch_Catalyst_9300_17_03_04(self):
        document.add_paragraph("The Image Deployment is when the NCE engineer has verified that the device/devices have passed their relevant analysis checks and are ready to perform an image deployment onto the device.")
        p = document.add_paragraph('Log onto:', style='List Number 2')
        ImageDeploymentAndValidation.add_hyperlink(p, 'https://software.cisco.com/download/navigator.html?mode=home',' https://software.cisco.com/download/navigator.html?mode=home ')
        document.add_paragraph('Go to software download page and map to Download Home and map to the relevant image for the upgrade ', style='List Number 2')
        document.add_paragraph('Take a note of the MD5 checksum and download the new image file. ',   style='List Number 2')
        document.add_paragraph('Run the below command: ',   style='List Number 2')
        ImageDeploymentAndValidation.funTable('Switch# verify /md5 flash:cat9k_iosxe.17.03.04.SPA.bin\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '................................\n'
                                              '...........Done!\n'
                                              'verify /md5 (flash:cat9k_iosxe.17.03.04.SPA.bin) = 4a64084f8108c9645b2ba264048d7665\n')
        document.add_page_break()
        return
    def imageDeploymentAndValidation_Switch_Catalyst_9500_17_03_04(self):
        document.add_paragraph("The Image Deployment is when the NCE engineer has verified that the device/devices have passed their relevant analysis checks and are ready to perform an image deployment onto the device.")
        p = document.add_paragraph('Log onto:', style='List Number 2')
        ImageDeploymentAndValidation.add_hyperlink(p, 'https://software.cisco.com/download/navigator.html?mode=home',' https://software.cisco.com/download/navigator.html?mode=home ')
        document.add_paragraph('Go to software download page and map to Download Home and map to the relevant image for the upgrade ', style='List Number 2')
        document.add_paragraph('Take a note of the MD5 checksum and download the new image file. ',   style='List Number 2')
        document.add_paragraph('Run the below command: ',   style='List Number 2')
        ImageDeploymentAndValidation.funTable('Switch# verify /md5 flash:cat9k_iosxe.17.03.04.SPA.bin\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '...................................................................................\n'
                                              '................................\n'
                                              '...........Done!\n'
                                              'verify /md5 (flash:cat9k_iosxe.17.03.04.SPA.bin) = 4a64084f8108c9645b2ba264048d7665\n')
        document.add_page_break()
        return

class BackupCurrentConfiguration(Formatter):
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def backupCurrentConfiguration_ASR_1000_17_03_03(self):
        BackupCurrentConfiguration.funTable('Router#copy running-config startup-config'
                                            '\nDestination filename [startup-config]? '
                                            '\nBuilding configuration...'
                                            '\n[OK]'
                                            '\nRouter#copy runn'
                                            '\nRouter#copy start          '
                                            '\nRouter#copy startup-config tftp'
                                            '\nRouter#copy startup-config tftp:'
                                            '\nAddress or name of remote host []? 10.197.65.24'
                                            '\nDestination filename [router-confg]?')
        document.add_page_break()
        return
                                            
class LimitationsAndRestrictions(Formatter):
    def __init__(self,os,version,lr):
        self.os = os
        self.version = version
        self.lr = lr
    def limitationsAndRestrictions_ASR_1000_17_03_03(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/routers/asr1000/release/notes/xe-17-3/asr1000-rel-notes-xe-17-3.html#reference_rnc_tvc_dhb')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#Cisco_Concept\.dita_e96aae30-f4f9-4d68-8262-e46aff0c05d7 > div > table > tbody > tr > td:nth-child(2)')
            href = elem[0].find_all('a', attrs={'href': re.compile("^https://")})
            d = elem[0].text.split()
            new_string = " ".join(d)
            document.add_paragraph(new_string)
            p = document.add_paragraph()
            LimitationsAndRestrictions.add_hyperlink(p,href[0].get('href'),href[0].get('href'))
            document.add_paragraph('\n')
            for i in self.lr:
                p = document.add_paragraph(i,style='List Bullet')     
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Limitations and Restrictions' section.")
        document.add_page_break()
        return
    def limitationsAndRestrictions_Switch_Catalyst_9600_17_06_02(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-6/release_notes/ol-17-6-9600.html#concept_ycv_jdf_3mb')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#concept_yyd_5gg_3rb > div')
            l = elem[0].text.splitlines()
            res = []
            for ele in l:
                if ele.strip():
                    res.append(ele)
            table = document.add_table(rows=1, cols=2, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = res[0]
            row[1].text = res[1]
            row = table.add_row().cells
            row[0].text = res[2]
            row[1].text = res[3]+res[4].strip()+'\n'+res[5]+'\n\n'+res[6]
            document.add_paragraph('\n')
            for i in self.lr:
                p = document.add_paragraph(i,style='List Bullet')
            
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Limitations and Restrictions' section.")
        document.add_page_break()
        return
    def limitationsAndRestrictions_Switch_Catalyst_9600_17_03_05(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-3/release_notes/ol-17-3-9600.html#concept_dcb_fd2_3mb')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#concept_dcb_fd2_3mb > div > table > tbody > tr:nth-child(1) > td:nth-child(2)')
            href = elem[0].find_all('a', attrs={'href': re.compile("^https://")})
            d = elem[0].text.split()
            new_string = " ".join(d)
            document.add_paragraph(new_string)
            p = document.add_paragraph()
            LimitationsAndRestrictions.add_hyperlink(p,href[0].get('href'),href[0].get('href'))
            document.add_paragraph('\n')
            for i in self.lr:
                p = document.add_paragraph(i,style='List Bullet')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Limitations and Restrictions' section.")
        document.add_page_break()
        return
    def limitationsAndRestrictions_Switch_Catalyst_9300_17_03_04(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9300/software/release/17-3/release_notes/ol-17-3-9300.html')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#concept_ymv_q2f_3mb > div')
            l = elem[0].text.splitlines()
            res = []
            for ele in l:
                if ele.strip():
                    res.append(ele.strip())
            #document.add_paragraph(res[0])
            document.add_paragraph(res[0],style='List Bullet')
            document.add_paragraph(res[1],style='List Bullet')
            document.add_paragraph(res[2],style='List Bullet')
            document.add_paragraph(res[3],style='List Bullet 2')
            document.add_paragraph(res[4],style='List Bullet 2')
            document.add_paragraph(res[5],style='List Bullet 2')
            document.add_paragraph(res[6],style='List Bullet')
            document.add_paragraph(res[7],style='List Bullet 2')
            document.add_paragraph(res[8],style='List Bullet 2')
            document.add_paragraph(res[9],style='List Bullet 2')
            document.add_paragraph(res[10],style='List Bullet 2')
            document.add_paragraph(res[11],style='List Bullet')
            document.add_paragraph(res[12],style='List Bullet 2')
            document.add_paragraph(res[13]+res[14],style='List Bullet 2')
            document.add_paragraph(res[15]+res[16]+res[17],style='List Bullet 2')
            document.add_paragraph(res[24],style='List Bullet')
            document.add_paragraph(res[25],style='List Bullet 2')
            document.add_paragraph(res[26],style='List Bullet 2')
            document.add_paragraph(res[27],style='List Bullet 2')
            document.add_paragraph(res[28]+res[29],style='List Bullet 2')
            document.add_paragraph(res[30],style='List Bullet 2')
            document.add_paragraph(res[31]+' '+res[32],style='List Bullet')
            document.add_paragraph(res[33]+' '+res[34]+' '+res[35],style='List Bullet')
            document.add_paragraph(res[51],style='List Bullet')
            document.add_paragraph(res[52],style='List Bullet')
            document.add_paragraph('\n')
            for i in self.lr:
                p = document.add_paragraph(i,style='List Bullet')
            
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Limitations and Restrictions' section.")
        document.add_page_break()
        return
    def limitationsAndRestrictions_Switch_Catalyst_9500_17_03_04(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9500/software/release/17-3/release_notes/ol-17-3-9500.html')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#concept_ymv_q2f_3mb > div')
            l = elem[0].text.splitlines()
            res = []
            for ele in l:
                if ele.strip():
                    res.append(ele.strip())
            document.add_paragraph(res[0])
            document.add_paragraph(res[3],style='List Bullet')
            document.add_paragraph('\t'+res[4]+' '+res[5])
            document.add_paragraph('\t'+res[6]+' '+res[7])
            document.add_paragraph(res[8],style='List Bullet')
            document.add_paragraph(res[9],style='List Bullet')
            document.add_paragraph(res[10]+' '+res[11]+' '+res[12],style='List Bullet 2')
            document.add_paragraph(res[13],style='List Bullet 2')
            document.add_paragraph(res[14]+' '+res[15],style='List Bullet 3')
            document.add_paragraph(res[16]+' '+res[17],style='List Bullet 3')
            document.add_paragraph(res[18]+' '+res[19],style='List Bullet 3')
            document.add_paragraph(res[20],style='List Bullet')
            document.add_paragraph(res[21],style='List Bullet')
            document.add_paragraph(res[22],style='List Bullet 2')
            document.add_paragraph(res[23],style='List Bullet 2')
            document.add_paragraph(res[24],style='List Bullet 2')
            document.add_paragraph(res[25],style='List Bullet')
            document.add_paragraph(res[26],style='List Bullet 2')
            document.add_paragraph(res[27],style='List Bullet 2')
            document.add_paragraph(res[28]+' '+res[29],style='List Bullet 2')
            document.add_paragraph(res[30],style='List Bullet 3')
            document.add_paragraph(res[31],style='List Bullet 3')
            document.add_paragraph(res[32],style='List Bullet 3')
            document.add_paragraph(res[33],style='List Bullet 3')
            document.add_paragraph(res[34]+' '+res[35],style='List Bullet 2')
            document.add_paragraph(res[36]+' '+res[37]+' '+res[38],style='List Bullet 2')
            document.add_paragraph(res[39]+' '+res[40],style='List Bullet 2')
            document.add_paragraph(res[41]+' '+res[42]+' '+res[43]+' '+res[44],style='List Bullet')
            document.add_paragraph(res[45],style='List Bullet')
            document.add_paragraph(res[46],style='List Bullet 2')
            document.add_paragraph(res[47]+ ' '+res[48],style='List Bullet 2')
            document.add_paragraph(res[49],style='List Bullet 2')
            document.add_paragraph(res[50]+' '+res[51],style='List Bullet 2')
            document.add_paragraph(res[52]+' '+res[53],style='List Bullet 2')
            document.add_paragraph(res[54],style='List Bullet')
            document.add_paragraph(res[55],style='List Bullet 2')
            document.add_paragraph(res[56],style='List Bullet 2')
            document.add_paragraph(res[57],style='List Bullet 2')
            document.add_paragraph(res[58],style='List Bullet')
            document.add_paragraph(res[59],style='List Bullet 2')
            document.add_paragraph(res[60]+' '+res[61],style='List Bullet 2')
            document.add_paragraph('\t'+res[62]+' '+res[63]+' '+res[64])
            document.add_paragraph(res[71]+' '+res[72],style='List Bullet')
            document.add_paragraph(res[73]+' '+res[74]+' '+res[75]+' '+res[76],style='List Bullet')
            document.add_paragraph(res[91],style='List Bullet')
            document.add_paragraph(res[92],style='List Bullet')
            document.add_paragraph('\n')
            for i in self.lr:
                p = document.add_paragraph(i,style='List Bullet')
            
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Limitations and Restrictions' section.")
        document.add_page_break()
        return

#Class for 3.1 Pre-install verification
class PreInstallVerification(Formatter):
    def __init__(self,os,version):
        self.os = os
        self.version = version

    def preInstallVerification_ASR_1000_17_03_03(self):
        PreInstallVerification.funTable('!\n'
                                        'terminal length 0\n'
                                        'show run\n'
                                        'show version\n'
                                        'show cdp neighbors\n'
                                        'show cdp neighbors detail\n'
                                        'show etherchannel summary\n'
                                        'show ip interface brief\n'
                                        'show ip interface brief | inc up\n'
                                        'show interface status\n'
                                        'show interface status | i connected\n'
                                        'show interface\n'
                                        'Show bfd neighbor\n'
                                        'show bfd drops\n'
                                        'show mac address-table\n'
                                        'show ip arp\n'
                                        'show vlan\n'
                                        'show vlan brief\n'
                                        'show spanning-tree\n'
                                        'show interfaces trunk\n'
                                        'show standby bri\n'
                                        'Show ip route\n'
                                        'show ip route summ\n'
                                        'Show ip ospf nei\n'
                                        'Show ip ospf interface\n'
                                        'show ip ospf database\n'
                                        'show ip bgp\n'
                                        'show ip bgp summ\n'
                                        'show udld neighbors\n'
                                        'show ntp status\n'
                                        'show controllers\n'
                                        'show env status\n'
                                        'show redundancy\n'
                                        'show platform\n'
                                        'Show inventory\n'
                                        'show diagnostic status\n'
                                        'show diagnostic events\n'
                                        'show proc cpu history\n'
                                        'show proc cpu sorted\n'
                                        'show proc cpu\n'
                                        'show log\n'
                                        'term len 30')
        document.add_paragraph('\n')
        PreInstallVerification.funTable('ASR1002-1#sh bootflash: | i .bin'
                   '\n18 393419388 Oct 09 2019 21:56:37 +00:00 /bootflash/asr1000rp1-advipservicesk9.03.16.06.S.155-3.S6-ext.bin'
                   '\n24 389286512 Jun 21 2021 22:38:49 +00:00 /bootflash/asr1000rp1-adventerprise.03.16.10.S.155-3.S10-ext.bin'
                   '\n27 395158140 Aug 31 2020 14:29:31 +00:00 /bootflash/asr1000rp1-adventerprisek9.03.16.08.S.155-3.S8-ext.bin'
                   '\n28 371004028 Oct 27 2020 12:43:12 +00:00 /bootflash/asr1000rp1-adventerprisek9.03.13.04.S.154-3.S4-ext.bin'
                   '\n29 759926681 Oct 07 2020 20:27:36 +00:00 /bootflash/asr1000-universalk9.16.09.05.SPA.bin'
                   '\n32 387019376 Jan 27 2021 07:05:15 +00:00 /bootflash/asr1000rp1-adventerprise.03.16.02.S.155-3.S2-ext.bin'
                   '\n33 459892152 Feb 16 2021 21:12:03 +00:00 /bootflash/asr1002x-universal.03.16.09.S.155-3.S9-ext.SPA.bin'
                   '\n34 347124348 Dec 19 2020 18:49:26 +00:00 /bootflash/asr1000rp1-adventerprisek9.03.10.05.S.153-3.S5-ext.bin'
                   '\n35 395094652 Apr 16 2021 11:54:36 +00:00 /bootflash/asr1000rp1-adventerprisek9.03.16.10.S.155-3.S10-ext.bin'
                   '\n42 666764156 Apr 19 2021 01:52:10 +00:00 /bootflash/asr1000rp2-advipservices.03.16.10.S.155-3.S10-ext.bin')
        document.add_page_break()
        return
    def preInstallVerification_Switch_Catalyst_9600_17_06_02(self):
        document.add_paragraph('It’s important to collect a baseline of the device’s state, allowing us to verify its correct operation once the upgrade has been completed. This information can be collected using a set of CLI commands and logged to a text file: as soon as the upgrade is completed, the same set of CLI commands can be run again, and the output compared with that gathered prior to the upgrade.')
        document.add_paragraph('When the install analysis checks have been reviewed by an NCE engineer the Pre verification script can now be ran.')
        PreInstallVerification.funTable('terminal length 0'
                                        '\nshow running-config'
                                        '\nshow cdp neighbors'
                                        '\nshow stackwise-virtual'
                                        '\nshow stackwise-virtual dual-active-detection'
                                        '\nshow stackwise-virtual link'
                                        '\nshow stackwise-virtual switch'
                                        '\nshow interface status'
                                        '\nshow ip interface brief'
                                        '\nshow vlan summary'
                                        '\nshow spanning-tree'
                                        '\nshow spanning-tree detail'
                                        '\nshow ip protocols'
                                        '\nshow ip bgp'
                                        '\nshow ip bgp summary'
                                        '\nshow ip route summary'
                                        '\nshow ip route'
                                        '\nshow processes cpu history'
                                        '\nshow memory'
                                        '\nshow inventory'
                                        '\nshow logging'
                                        '\nshow version'
                                        '\nshow ntp status')
        return
    def preInstallVerification_Switch_Catalyst_9600_17_03_05(self):
        document.add_paragraph('It’s important to collect a baseline of the device’s state, allowing us to verify its correct operation once the upgrade has been completed. This information can be collected using a set of CLI commands and logged to a text file: as soon as the upgrade is completed, the same set of CLI commands can be run again, and the output compared with that gathered prior to the upgrade.')
        document.add_paragraph('When the install analysis checks have been reviewed by an NCE engineer the Pre verification script can now be ran.')
        PreInstallVerification.funTable('terminal length 0'
                                        '\nshow running-config'
                                        '\nshow cdp neighbors'
                                        '\nshow stackwise-virtual'
                                        '\nshow stackwise-virtual dual-active-detection'
                                        '\nshow stackwise-virtual link'
                                        '\nshow stackwise-virtual switch'
                                        '\nshow interface status'
                                        '\nshow ip interface brief'
                                        '\nshow vlan summary'
                                        '\nshow spanning-tree'
                                        '\nshow spanning-tree detail'
                                        '\nshow ip protocols'
                                        '\nshow ip bgp'
                                        '\nshow ip bgp summary'
                                        '\nshow ip route summary'
                                        '\nshow ip route'
                                        '\nshow processes cpu history'
                                        '\nshow memory'
                                        '\nshow inventory'
                                        '\nshow logging'
                                        '\nshow version'
                                        '\nshow ntp status')
        return
    def preInstallVerification_Switch_Catalyst_9300_17_03_04(self):
        document.add_paragraph('It’s important to collect a baseline of the device’s state, allowing us to verify its correct operation once the upgrade has been completed. This information can be collected using a set of CLI commands and logged to a text file: as soon as the upgrade is completed, the same set of CLI commands can be run again, and the output compared with that gathered prior to the upgrade.')
        document.add_paragraph('When the install analysis checks have been reviewed by an NCE engineer the Pre verification script can now be ran.')
        PreInstallVerification.funTable('terminal length 0'
                                        '\nshow running-config'
                                        '\nshow cdp neighbors'
                                        '\nshow stackwise-virtual'
                                        '\nshow stackwise-virtual dual-active-detection'
                                        '\nshow stackwise-virtual link'
                                        '\nshow stackwise-virtual switch'
                                        '\nshow interface status'
                                        '\nshow ip interface brief'
                                        '\nshow vlan summary'
                                        '\nshow spanning-tree'
                                        '\nshow spanning-tree detail'
                                        '\nshow ip protocols'
                                        '\nshow ip bgp'
                                        '\nshow ip bgp summary'
                                        '\nshow ip route summary'
                                        '\nshow ip route'
                                        '\nshow processes cpu history'
                                        '\nshow memory'
                                        '\nshow inventory'
                                        '\nshow logging'
                                        '\nshow version'
                                        '\nshow ntp status')
        return
    def preInstallVerification_Switch_Catalyst_9500_17_03_04(self):
        document.add_paragraph('It’s important to collect a baseline of the device’s state, allowing us to verify its correct operation once the upgrade has been completed. This information can be collected using a set of CLI commands and logged to a text file: as soon as the upgrade is completed, the same set of CLI commands can be run again, and the output compared with that gathered prior to the upgrade.')
        document.add_paragraph('When the install analysis checks have been reviewed by an NCE engineer the Pre verification script can now be ran.')
        PreInstallVerification.funTable('terminal length 0'
                                        '\nshow running-config'
                                        '\nshow cdp neighbors'
                                        '\nshow stackwise-virtual'
                                        '\nshow stackwise-virtual dual-active-detection'
                                        '\nshow stackwise-virtual link'
                                        '\nshow stackwise-virtual switch'
                                        '\nshow interface status'
                                        '\nshow ip interface brief'
                                        '\nshow vlan summary'
                                        '\nshow spanning-tree'
                                        '\nshow spanning-tree detail'
                                        '\nshow ip protocols'
                                        '\nshow ip bgp'
                                        '\nshow ip bgp summary'
                                        '\nshow ip route summary'
                                        '\nshow ip route'
                                        '\nshow processes cpu history'
                                        '\nshow memory'
                                        '\nshow inventory'
                                        '\nshow logging'
                                        '\nshow version'
                                        '\nshow ntp status')
        return

class InstallAndReload:
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def installAndReload_ASR_1000_17_03_03(self):
        document.add_paragraph("After the Pre-install verification script checks have been completed and reviewed by an NCE engineer,the Install and reload script is the next step for completing the upgrade.")
        r = document.add_paragraph().add_run()
        try:
            r.add_picture('timesaver.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text(' The install process takes around 20-30 minutes which is the time taken for the network to be disruptive.')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
        return
    def installAndReload_Switch_Catalyst_9600_17_06_02(self):
        document.add_paragraph("After the Pre-install verification checks have been completed and reviewed by an NCE engineer,the Install and reload is the next step for completing the upgrade.")
        r = document.add_paragraph().add_run()
        try:
            r.add_picture('timesaver.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text(' The install process takes around 20-30 minutes which is the time taken for the network to be disruptive.')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
        return
    def installAndReload_Switch_Catalyst_9600_17_03_05(self):
        document.add_paragraph("After the Pre-install verification checks have been completed and reviewed by an NCE engineer,the Install and reload is the next step for completing the upgrade.")
        r = document.add_paragraph().add_run()
        try:
            r.add_picture('timesaver.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text(' The install process takes around 20-30 minutes which is the time taken for the network to be disruptive.')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
        return
    def installAndReload_Switch_Catalyst_9300_17_03_04(self):
        document.add_paragraph("After the Pre-install verification checks have been completed and reviewed by an NCE engineer,the Install and reload is the next step for completing the upgrade.")
        r = document.add_paragraph().add_run()
        try:
            r.add_picture('timesaver.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text(' The install process takes around 20-30 minutes which is the time taken for the network to be disruptive.')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
        return
    def installAndReload_Switch_Catalyst_9500_17_03_04(self):
        document.add_paragraph("After the Pre-install verification checks have been completed and reviewed by an NCE engineer,the Install and reload is the next step for completing the upgrade.")
        r = document.add_paragraph().add_run()
        try:
            r.add_picture('timesaver.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text(' The install process takes around 20-30 minutes which is the time taken for the network to be disruptive.')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
        return

#class for Upgrade switch software
class UpgradeSwitch:
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def upgradeSwitch_Switch_Catalyst_9600_17_06_02(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-6/release_notes/ol-17-6-9600.html#concept_ycv_jdf_3mb')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#concept_kk5_xbf_3mb')
            l = elem[0].text.splitlines()
            res = []
            for ele in l:
                if ele.strip():
                    res.append(ele)
            para = document.add_paragraph()
            para.add_run(res[0]).bold = True
            table = document.add_table(rows=5, cols=3, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = res[1]
            row[1].text = res[2]
            row[2].text = res[3]
            a = table.cell(1, 0)
            b = table.cell(2, 0)
            A = a.merge(b)
            A.text = res[4]
            a1 = table.cell(1, 1)
            a1.text = res[5]
            a2 = table.cell(1, 2)
            a2.text = res[6]
            a3 = table.cell(2, 1)
            a3.text = res[7]
            a4 = table.cell(2, 2)
            a4.text = res[8]
            c = table.cell(3, 0)
            d = table.cell(4, 0)
            B = c.merge(d)
            B.text = res[9]
            a5 = table.cell(3, 1)
            a5.text = res[10]
            a6 = table.cell(3,2)
            a6.text = res[11]
            a7 = table.cell(4, 1)
            a7.text = res[12]
            a8 = table.cell(4, 2)
            a8.text = res[13]
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Upgrading the Switch Software' section")
        return
    def upgradeSwitch_Switch_Catalyst_9600_17_03_05(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-3/release_notes/ol-17-3-9600.html#task_a2s_dfh_jmb')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#concept_kk5_xbf_3mb')
            l = elem[0].text.splitlines()
            res = []
            for ele in l:
                if ele.strip():
                    res.append(ele)
            para = document.add_paragraph()
            para.add_run(res[0]).bold = True
            table = document.add_table(rows=11, cols=3, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = res[1]
            row[1].text = res[2]
            row[2].text = res[3]
            a = table.cell(1, 0)
            b = table.cell(2, 0)
            A = a.merge(b)
            A.text = res[4]
            a1 = table.cell(1, 1)
            a1.text = res[5]
            a2 = table.cell(1, 2)
            a2.text = res[6]
            a3 = table.cell(2, 1)
            a3.text = res[7]
            a4 = table.cell(2, 2)
            a4.text = res[8]
            c = table.cell(3, 0)
            d = table.cell(4, 0)
            B = c.merge(d)
            B.text = res[9]
            a5 = table.cell(3, 1)
            a5.text = res[10]
            a6 = table.cell(3,2)
            a6.text = res[11]
            a7 = table.cell(4, 1)
            a7.text = res[12]
            a8 = table.cell(4, 2)
            a8.text = res[13]
            e = table.cell(5, 0)
            f = table.cell(6, 0)
            B = e.merge(f)
            B.text = res[14]
            a5 = table.cell(5, 1)
            a5.text = res[15]
            a6 = table.cell(5,2)
            a6.text = res[16]
            a7 = table.cell(6, 1)
            a7.text = res[17]
            a8 = table.cell(6, 2)
            a8.text = res[18]
            g = table.cell(7, 0)
            h = table.cell(8, 0)
            B = g.merge(h)
            B.text = res[19]
            a5 = table.cell(7, 1)
            a5.text = res[20]
            a6 = table.cell(7,2)
            a6.text = res[21]
            a7 = table.cell(8, 1)
            a7.text = res[22]
            a8 = table.cell(8, 2)
            a8.text = res[23]
            i = table.cell(9, 0)
            j = table.cell(10, 0)
            B = i.merge(j)
            B.text = res[24]
            a5 = table.cell(9, 1)
            a5.text = res[25]
            a6 = table.cell(9,2)
            a6.text = res[26]
            a7 = table.cell(10, 1)
            a7.text = res[27]
            a8 = table.cell(10, 2)
            a8.text = res[28]
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Upgrading the Switch Software' section")
        return
    def upgradeSwitch_Switch_Catalyst_9300_17_03_04(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9300/software/release/17-3/release_notes/ol-17-3-9300.html')
        soup = BeautifulSoup(res.text, 'lxml')
        para = document.add_paragraph()
        para.add_run('Finding the Software Version').bold = True
        document.add_paragraph('The package files for the Cisco IOS XE software are stored on the system board flash device (flash:).\n'
                               'We can use the show version privileged EXEC command to see the software version that is running on your switch.')
        r = document.add_paragraph().add_run()
        try:
            r.add_picture('note.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('Although the show version output always shows the software image running on the switch, the model name shown at the end of this display is the factory configuration and does not change if we upgrade the software license.')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
        document.add_paragraph('We can also use the dir filesystem: privileged EXEC command to see the directory names of other software images that we might have stored in flash memory')
        try:
            elem = soup.select('#concept_kk5_xbf_3mb')
            l = elem[0].text.splitlines()
            res = []
            for ele in l:
                if ele.strip():
                    res.append(ele)
            para = document.add_paragraph()
            para.add_run(res[0]).bold = True
            table = document.add_table(rows=11, cols=3, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = res[1]
            row[1].text = res[2]
            row[2].text = res[3]
            a = table.cell(1, 0)
            b = table.cell(2, 0)
            A = a.merge(b)
            A.text = res[4]
            a1 = table.cell(1, 1)
            a1.text = res[5]
            a2 = table.cell(1, 2)
            a2.text = res[6]
            a3 = table.cell(2, 1)
            a3.text = res[7]
            a4 = table.cell(2, 2)
            a4.text = res[8]
            c = table.cell(3, 0)
            d = table.cell(4, 0)
            B = c.merge(d)
            B.text = res[9]
            a5 = table.cell(3, 1)
            a5.text = res[10]
            a6 = table.cell(3,2)
            a6.text = res[11]
            a7 = table.cell(4, 1)
            a7.text = res[12]
            a8 = table.cell(4, 2)
            a8.text = res[13]
            e = table.cell(5, 0)
            f = table.cell(6, 0)
            B = e.merge(f)
            B.text = res[14]
            a5 = table.cell(5, 1)
            a5.text = res[15]
            a6 = table.cell(5,2)
            a6.text = res[16]
            a7 = table.cell(6, 1)
            a7.text = res[17]
            a8 = table.cell(6, 2)
            a8.text = res[18]
            g = table.cell(7, 0)
            h = table.cell(8, 0)
            B = g.merge(h)
            B.text = res[19]
            a5 = table.cell(7, 1)
            a5.text = res[20]
            a6 = table.cell(7,2)
            a6.text = res[21]
            a7 = table.cell(8, 1)
            a7.text = res[22]
            a8 = table.cell(8, 2)
            a8.text = res[23]
            i = table.cell(9, 0)
            j = table.cell(10, 0)
            B = i.merge(j)
            B.text = res[24]
            a5 = table.cell(9, 1)
            a5.text = res[25]
            a6 = table.cell(9,2)
            a6.text = res[26]
            a7 = table.cell(10, 1)
            a7.text = res[27]
            a8 = table.cell(10, 2)
            a8.text = res[28]
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Upgrading the Switch Software' section")
        return
    def upgradeSwitch_Switch_Catalyst_9500_17_03_04(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9500/software/release/17-3/release_notes/ol-17-3-9500.html')
        soup = BeautifulSoup(res.text, 'lxml')
        para = document.add_paragraph()
        para.add_run('Finding the Software Version').bold = True
        document.add_paragraph('The package files for the Cisco IOS XE software are stored on the system board flash device (flash:).\n'
                               'We can use the show version privileged EXEC command to see the software version that is running on your switch.')
        r = document.add_paragraph().add_run()
        try:
            r.add_picture('note.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('Although the show version output always shows the software image running on the switch, the model name shown at the end of this display is the factory configuration and does not change if we upgrade the software license.')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
        document.add_paragraph('We can also use the dir filesystem: privileged EXEC command to see the directory names of other software images that we might have stored in flash memory')
        try:
            elem = soup.select('#concept_kk5_xbf_3mb')#concept_kk5_xbf_3mb
            l = elem[0].text.splitlines()
            res = []
            for ele in l:
                if ele.strip():
                    res.append(ele)
            para = document.add_paragraph()
            para.add_run(res[0]).bold = True
            table = document.add_table(rows=11, cols=3, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = res[1]
            row[1].text = res[2]
            row[2].text = res[3]
            a = table.cell(1, 0)
            b = table.cell(2, 0)
            A = a.merge(b)
            A.text = res[4]
            a1 = table.cell(1, 1)
            a1.text = res[5]
            a2 = table.cell(1, 2)
            a2.text = res[6]
            a3 = table.cell(2, 1)
            a3.text = res[7]
            a4 = table.cell(2, 2)
            a4.text = res[8]
            c = table.cell(3, 0)
            d = table.cell(4, 0)
            B = c.merge(d)
            B.text = res[9]
            a5 = table.cell(3, 1)
            a5.text = res[10]
            a6 = table.cell(3,2)
            a6.text = res[11]
            a7 = table.cell(4, 1)
            a7.text = res[12]
            a8 = table.cell(4, 2)
            a8.text = res[13]
            e = table.cell(5, 0)
            f = table.cell(6, 0)
            B = e.merge(f)
            B.text = res[14]
            a5 = table.cell(5, 1)
            a5.text = res[15]
            a6 = table.cell(5,2)
            a6.text = res[16]
            a7 = table.cell(6, 1)
            a7.text = res[17]
            a8 = table.cell(6, 2)
            a8.text = res[18]
            g = table.cell(7, 0)
            h = table.cell(8, 0)
            B = g.merge(h)
            B.text = res[19]
            a5 = table.cell(7, 1)
            a5.text = res[20]
            a6 = table.cell(7,2)
            a6.text = res[21]
            a7 = table.cell(8, 1)
            a7.text = res[22]
            a8 = table.cell(8, 2)
            a8.text = res[23]
            i = table.cell(9, 0)
            j = table.cell(10, 0)
            B = i.merge(j)
            B.text = res[24]
            a5 = table.cell(9, 1)
            a5.text = res[25]
            a6 = table.cell(9,2)
            a6.text = res[26]
            a7 = table.cell(10, 1)
            a7.text = res[27]
            a8 = table.cell(10, 2)
            a8.text = res[28]
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Upgrading the Switch Software' section")
        return
    
#class for Rommon Upgrade 
class RommonUpgrade(Formatter):
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def rommonUpgrade_ASR_1000_17_03_03(self):
        document.add_paragraph('The following sequence of commands is an example of the procedure to upgrade the ROMmon for all the RPs, ESPs, MIPs, and SIPs on a router:')
        resp = requests.get('https://www.cisco.com/c/en/us/td/docs/routers/asr1000/rommon/asr1000-rommon-upg-guide.html#con_46405')
        try:
            soup = BeautifulSoup(resp.text, 'lxml')
            elem = soup.select('#con_59000 > div > pre')
            RommonUpgrade.funTable(elem[0].text)
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Rommon Upgrade' section")
    def rommonUpgrade_Switch_Catalyst_9600_17_06_02(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-6/release_notes/ol-17-6-9600.html#concept_ycv_jdf_3mb')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            p1 = soup.select('#concept_ycv_jdf_3mb__d54e3364 > p:nth-child(1)')
            document.add_paragraph(" ".join(p1[0].text.split()))
            b1 = soup.select('#concept_ycv_jdf_3mb__d54e3364 > ul:nth-child(2)')

            l = b1[0].text.splitlines()
            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele)

            document.add_paragraph(t1[0], style = 'List Bullet 2')
            document.add_paragraph(t1[1]+' '+t1[2].strip(), style = 'List Bullet 2')
            p2 = soup.select('#concept_ycv_jdf_3mb__d54e3364 > p:nth-child(3)')
            document.add_paragraph(" ".join(p2[0].text.split()))
            p3 = soup.select('#concept_ycv_jdf_3mb__d54e3364 > p:nth-child(5)')
            document.add_paragraph(" ".join(p3[0].text.split()))
            p4 = soup.select('#concept_ycv_jdf_3mb__d54e3364 > ul:nth-child(6)')
            document.add_paragraph(" ".join(p4[0].text.split()))
            p5 = soup.select('#concept_ycv_jdf_3mb__d54e3364 > p:nth-child(9)')
            document.add_paragraph(" ".join(p5[0].text.split()))
            table = soup.select('#concept_ycv_jdf_3mb__table_9600_wmc_fj1_1lb')
            l = table[0].text.splitlines()
            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele)
            table = document.add_table(rows=1, cols=2, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = t1[0]
            row[1].text = t1[1]
            row = table.add_row().cells
            row[0].text = t1[2]
            row[1].text = t1[3]+'\n\n'+t1[4].strip()+'\n\n'+t1[5]+'\n'+t1[6]+'\n\n'+t1[7]+t1[8].strip()+'\n\n'+t1[9]+'\n'+t1[10]+'\n'+t1[11]+'\n'+t1[12]
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Rommon Upgrade' section")
    
    def rommonUpgrade_Switch_Catalyst_9600_17_03_05(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-3/release_notes/ol-17-3-9600.html#concept_ycv_jdf_3mb')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            p1 = soup.select('#concept_ycv_jdf_3mb__d54e3704 > p:nth-child(1)')
            document.add_paragraph(" ".join(p1[0].text.split()))
            b1 = soup.select('#concept_ycv_jdf_3mb__d54e3704 > ul:nth-child(2)')

            l = b1[0].text.splitlines()
            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele)

            document.add_paragraph(t1[0], style = 'List Bullet 2')
            document.add_paragraph(t1[1]+' '+t1[2].strip(), style = 'List Bullet 2')
            p2 = soup.select('#concept_ycv_jdf_3mb__d54e3704 > p:nth-child(3)')
            document.add_paragraph(" ".join(p2[0].text.split()))
            p3 = soup.select('#concept_ycv_jdf_3mb__d54e3704 > p:nth-child(5)')
            document.add_paragraph(" ".join(p3[0].text.split()))
            p4 = soup.select('#concept_ycv_jdf_3mb__d54e3704 > ul:nth-child(6)')
            document.add_paragraph(" ".join(p4[0].text.split()))
            p5 = soup.select('#concept_ycv_jdf_3mb__d54e3704 > p:nth-child(9)')
            document.add_paragraph(" ".join(p5[0].text.split()))
            table = soup.select('#concept_ycv_jdf_3mb__table_9600_wmc_fj1_1lb')
            l = table[0].text.splitlines()
            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele)
            table = document.add_table(rows=1, cols=2, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = t1[0]
            row[1].text = t1[1]
            row = table.add_row().cells
            row[0].text = t1[2]
            row[1].text = t1[3]+'\n\n'+t1[4].strip()+'\n\n'+t1[5]+'\n'+t1[6]+'\n\n'+t1[7]+'\n'+t1[8].strip()+'\n\n'+t1[9]+'\n'+t1[10]
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Rommon Upgrade' section")
        return
    def rommonUpgrade_Switch_Catalyst_9300_17_03_04(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9300/software/release/17-3/release_notes/ol-17-3-9300.html')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            p1 = soup.select('#concept_ycv_jdf_3mb__d54e5427 > p:nth-child(1)')
            document.add_paragraph(" ".join(p1[0].text.split()))
            b1 = soup.select('#concept_ycv_jdf_3mb__d54e5427 > ul:nth-child(2)')
            l = b1[0].text.splitlines()
            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele)
            document.add_paragraph(t1[0], style = 'List Bullet 2')
            document.add_paragraph(t1[1]+' '+t1[2].strip(), style = 'List Bullet 2')
            p2 = soup.select('#concept_ycv_jdf_3mb__d54e5427 > p:nth-child(3)')
            document.add_paragraph(" ".join(p2[0].text.split()))
            p3 = soup.select('#concept_ycv_jdf_3mb__d54e5427 > p:nth-child(5)')
            document.add_paragraph(" ".join(p3[0].text.split()))
            p4 = soup.select('#concept_ycv_jdf_3mb__d54e5427 > ul:nth-child(6)')
            l = p4[0].text.splitlines()
            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele.strip())
            document.add_paragraph(t1[0], style = 'List Bullet 2')
            document.add_paragraph('\t'+t1[1]+' '+ t1[2]+' '+t1[3]+' '+t1[4])
            document.add_paragraph(t1[5], style = 'List Bullet 2')
            document.add_paragraph('\t'+t1[6])
            p = document.add_paragraph('You must manually upgrade this ROMMON. Enter the ')
            r = p.add_run('upgrade rom-monitor capsule golden switch')
            r.bold = True
            r.italic = True
            p.add_run(' command in privileged EXEC mode.')
            try:
                r = document.add_paragraph().add_run()
                r.add_picture('note.jpg', width=Inches(0.3), height=Inches(.3))
                r.add_text('In case of a switch stack, perform the upgrade on the active switch and all members of the stack.')
            except Exception as e:
                exc_type, exc_obj, exc_tb = sys.exc_info()
                print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
            p5 = soup.select('#concept_ycv_jdf_3mb__d54e5427 > p:nth-child(9)')
            document.add_paragraph(" ".join(p5[0].text.split()))
            table = soup.select('#concept_ycv_jdf_3mb__table_9300_gs5_lj1_1lb')
            l = table[0].text.splitlines()
            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele.strip())
            table = document.add_table(rows=1, cols=2, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = t1[0]
            row[1].text = t1[1]
            row = table.add_row().cells
            row[0].text = t1[2]
            row[1].text = t1[3]+t1[4]+t1[5]+'\n'+t1[6]+'\n'+t1[7]+'\n'+t1[8]+t1[9]+t1[10]+t1[11]+'\n'+t1[12]+'\n'+t1[13]+'\n'+t1[14]+'\n\n'+t1[15]+'\n\n'+t1[16]+'\n'+t1[17]+'\n'+t1[18]+'\n'+t1[19]+'\n'+t1[20]
            row = table.add_row().cells
            row[0].text = t1[21]
            row[1].text = t1[22]+t1[23]+t1[24]+'\n'+t1[25]+'\n'+t1[26]+'\n'+t1[27]+t1[28]+t1[29]+t1[30]+'\n'+t1[31]+'\n'+t1[32]+'\n'+t1[33]+'\n\n'+t1[34]+'\n\n'+t1[35]+'\n'+t1[36]+'\n'+t1[37]+'\n'+t1[38]+'\n'+t1[39]+'\n'+t1[40]
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Rommon Upgrade' section")
        return
    def rommonUpgrade_Switch_Catalyst_9500_17_03_04(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9500/software/release/17-3/release_notes/ol-17-3-9500.html')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            p1 = soup.select('#concept_ycv_jdf_3mb__d54e5023 > p:nth-child(1)')
            document.add_paragraph(" ".join(p1[0].text.split()))
            b1 = soup.select('#concept_ycv_jdf_3mb__d54e5023 > ul:nth-child(2)')
            l = b1[0].text.splitlines()
            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele)
            document.add_paragraph(t1[0], style = 'List Bullet 2')
            document.add_paragraph(t1[1]+' '+t1[2].strip(), style = 'List Bullet 2')
            p2 = soup.select('#concept_ycv_jdf_3mb__d54e5023 > p:nth-child(3)')
            document.add_paragraph(" ".join(p2[0].text.split()))
            p3 = soup.select('#concept_ycv_jdf_3mb__d54e5023 > p:nth-child(5)')
            document.add_paragraph(" ".join(p3[0].text.split()))
            p4 = soup.select('#concept_ycv_jdf_3mb__d54e5023 > ul:nth-child(6)')
            l = p4[0].text.splitlines()
            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele.strip())
            document.add_paragraph(t1[0], style = 'List Bullet 2')
            document.add_paragraph('\t'+t1[1]+'\n\n'+ t1[2]+' '+t1[3]+' '+t1[4]+''+t1[5])
            document.add_paragraph(t1[6], style = 'List Bullet 2')
            document.add_paragraph('\t'+t1[7])
            p = document.add_paragraph('You must manually upgrade this ROMMON. The manual upgrade applies to all models in the series. Enter the ')
            r = p.add_run('upgrade rom-monitor capsule golden switch')
            r.bold = True
            r.italic = True
            p.add_run(' command in privileged EXEC mode.')
            
            try:
                r = document.add_paragraph().add_run()
                r.add_picture('note.jpg', width=Inches(0.3), height=Inches(.3))
                r.add_text('In case of a switch stack, perform the upgrade on the active switch and all members of the stack.')
            except Exception as e:
                exc_type, exc_obj, exc_tb = sys.exc_info()
                print(f"{e}(line {exc_tb.tb_lineno}): Image not Found")
            p5 = soup.select('#concept_ycv_jdf_3mb__d54e5023 > p:nth-child(9)')
            document.add_paragraph(" ".join(p5[0].text.split()))
            table = soup.select('#concept_ycv_jdf_3mb__table_9500_xph_gj1_1lb')
            l = table[0].text.splitlines()
            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele.strip())
            table = document.add_table(rows=1, cols=2, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = t1[0]
            row[1].text = t1[1]
            row = table.add_row().cells
            row[0].text = t1[2]
            row[1].text = t1[3]+t1[4]+t1[5]+'\n'+t1[6]+'\n'+t1[7]+'\n'+t1[8]+t1[9]+t1[10]+t1[11]+'\n'+t1[12]+'\n'+t1[13]+'\n'+t1[14]+'\n\n'+t1[15]+'\n\n'+t1[16]+'\n'+t1[17]+'\n'+t1[18]
            
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Rommon Upgrade' section")
        return


#class for Upgrade in Install Mode
class UpgradeInInstallMode(Formatter):
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def upgradeInInstallMode_Switch_Catalyst_9600_17_06_02(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-6/release_notes/ol-17-6-9600.html#concept_ycv_jdf_3mb')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#task_a2s_dfh_jmb__d54e3778 > p')
            table1 = soup.select('#task_a2s_dfh_jmb__d54e3691 > table:nth-child(4)')
            l = table1[0].text.splitlines()

            res = []
            for ele in l:
                if ele.strip():
                    res.append(ele)
            elem2 = soup.select('#task_a2s_dfh_jmb > div > table')
            document.add_paragraph(elem[0].text)
            table = document.add_table(rows=1, cols=2, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = res[0]
            row[1].text = res[1]
            for index, elem in enumerate(res[2:]):
                if (index+1 < len(l) and index - 1 >= 0) and elem != "":
                    row = table.add_row().cells
                    row[0].text = str(res[index+1])
                    row[1].text = str(elem)
            step1 = soup.select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(1)')
            document.add_paragraph("Step 1 "+step1[0].find('p', class_="ph cmd").text)
            document.add_paragraph(step1[0].find('span', class_="ph synph").text)
            document.add_paragraph(re.sub(' +', ' ', step1[0].find('p', class_='p').text.replace('\n','')))
            document.add_paragraph('The following sample output displays the cleaning up of unused files, by using the install remove inactive command:')
            document.add_paragraph(UpgradeInInstallMode.funTable(step1[0].find('pre',class_="pre codeblock").text))

            step2 = soup.select("#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(2) > td:nth-child(2)")
            document.add_paragraph("Step 2 "+step2[0].find('p', class_="ph cmd").text)
            li1 = step2[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(2) > td:nth-child(2) > ol > li:nth-child(1)')
            document.add_paragraph('a.   '+li1[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to copy the new image from a TFTP server to flash memory. The location is either an IP address or a host name. The filename is specified relative to the directory used for file transfers. Skip this step if you want to use the new image from a TFTP server.')
            document.add_paragraph(UpgradeInInstallMode.funTable(li1[0].find('pre',class_="pre codeblock").text))

            li2 = step2[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(2) > td:nth-child(2) > ol > li:nth-child(2)')
            document.add_paragraph('b.   '+li2[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to confirm that the image has been successfully copied to flash.')
            document.add_paragraph(UpgradeInInstallMode.funTable(li2[0].find('pre',class_="pre codeblock").text))

            step3 = soup.select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2)')
            document.add_paragraph("Step 3 "+step3[0].find('p', class_="ph cmd").text)
            li1 = step3[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(1)')
            document.add_paragraph('a.   '+li1[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to set the boot variable to flash:packages.conf .')
            document.add_paragraph(UpgradeInInstallMode.funTable(li1[0].find('pre',class_="pre codeblock").text))
            li2 = step3[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(2)')
            document.add_paragraph('b.   '+li2[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to configure the switch to auto-boot. Settings are synchronized with the standby switch, if applicable.')
            document.add_paragraph(UpgradeInInstallMode.funTable(li2[0].find('pre',class_="pre codeblock").text))
            li3 = step3[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(3)')
            document.add_paragraph('c.   '+li3[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to save boot settings.')
            document.add_paragraph(UpgradeInInstallMode.funTable(li3[0].find('pre',class_="pre codeblock").text))
            li4 = step3[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(4)')
            document.add_paragraph('d.   '+li4[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to verify the boot variable (packages.conf) and manual boot setting (no):')
            document.add_paragraph(UpgradeInInstallMode.funTable(li4[0].find('pre',class_="pre codeblock").text))

            step4 = soup.select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(4) > td:nth-child(2)')
            document.add_paragraph("Step 4 "+step4[0].find('p', class_="ph cmd").text)
            document.add_paragraph('Install image to flash\n'
                                   'install add file activate commit\n'
                                   'Use this command to install the image.\n'
                                   'We recommend that you point to the source image on a TFTP server or the flash , if you have copied the image to flash memory.\n\n'
                                   'The following sample output displays installation of the Cisco IOS XE Bengaluru 17.6.1 software image to flash:')
            document.add_paragraph(UpgradeInInstallMode.funTable(step4[0].find('pre',class_="pre codeblock").text))
            step5 = soup.select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(5) > td:nth-child(2)')
            document.add_paragraph("Step 5 "+step5[0].find('p', class_="ph cmd").text)
            document.add_paragraph(step5[0].find('p', class_="p").text)
            li1 = step5[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(5) > td:nth-child(2) > ol > li:nth-child(1)')
            document.add_paragraph('a.   '+li1[0].find('span', class_="ph synph").text)
            document.add_paragraph('The following is sample output of the dir flash:*.pkg command:')
            document.add_paragraph(UpgradeInInstallMode.funTable(li1[0].find('pre',class_="pre codeblock").text))

            li2 = step5[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(5) > td:nth-child(2) > ol > li:nth-child(2)')
            document.add_paragraph('b.   '+li2[0].find('span', class_="ph synph").text)
            document.add_paragraph('The following is sample output of the dir flash:*.conf command. It displays the .conf files in the flash partition; note the two .conf files:\n'
                                   'a.    packages.conf—the file that has been re-written with the newly installed .pkg files.  \n' 
                                   'b.    cat9k_iosxe.17.06.01.SPA.conf— a backup copy of the newly installed packages.conf file.')
            document.add_paragraph(UpgradeInInstallMode.funTable(li2[0].find('pre',class_="pre codeblock").text))

            step6 = soup.select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(6) > td:nth-child(2)')
            document.add_paragraph("Step 6 "+step6[0].find('p', class_="ph cmd").text)
            document.add_paragraph(step6[0].find('p', class_="p").text)
            document.add_paragraph('After the image boots up, use this command to verify the version of the new image.\n')
            document.add_paragraph('The following sample output of the show version command displays the Cisco IOS XE Bengaluru 17.6.1 image on the device:')
            document.add_paragraph(UpgradeInInstallMode.funTable(step6[0].find('pre',class_="pre codeblock").text))
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Upgrading In Install Mode' section")
        document.add_page_break()
        return
        
    def upgradeInInstallMode_Switch_Catalyst_9600_17_03_05(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-3/release_notes/ol-17-3-9600.html#task_a2s_dfh_jmb')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-3/release_notes/ol-17-3-9600.html#concept_ycv_jdf_3mb')
            soup = BeautifulSoup(res.text, 'lxml')
            elem = soup.select('#task_a2s_dfh_jmb__d54e4166 > p')
            table1 = soup.select('#task_a2s_dfh_jmb__d54e4054 > table:nth-child(4)')
            l = table1[0].text.splitlines()
            res = []
            for ele in l:
                if ele.strip():
                    res.append(ele)
            elem2 = soup.select('#task_a2s_dfh_jmb > div > table')
            document.add_paragraph(elem[0].text)
            table = document.add_table(rows=1, cols=2, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = res[0]
            row[1].text = res[1]
            for index, elem in enumerate(res[2:]):
                if (index+1 < len(l) and index - 1 >= 0) and elem != "":
                    row = table.add_row().cells
                    row[0].text = str(res[index+1])
                    row[1].text = str(elem)
            step1 = soup.select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(1)')
            p = document.add_paragraph("\n")
            p.add_run('Step 1 ').bold = True
            p.add_run(step1[0].find('p', class_="ph cmd").text)
            document.add_paragraph(step1[0].find('span', class_="ph synph").text)
            document.add_paragraph(re.sub(' +', ' ', step1[0].find('p', class_='p').text.replace('\n','')))
            document.add_paragraph('The following sample output displays the cleaning up of unused files, by using the install remove inactive command:')
            document.add_paragraph(UpgradeInInstallMode.funTable(step1[0].find('pre',class_="pre codeblock").text))
            step2 = soup.select("#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(2) > td:nth-child(2)")
            p = document.add_paragraph("\n")
            p.add_run('Step 2 ').bold = True
            p.add_run(step2[0].find('p', class_="ph cmd").text)
            li1 = step2[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(2) > td:nth-child(2) > ol > li:nth-child(1)')
            document.add_paragraph('a.   '+li1[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to copy the new image from a TFTP server to flash memory. The location is either an IP address or a host name. The filename is specified relative to the directory used for file transfers. Skip this step if you want to use the new image from a TFTP server.')
            document.add_paragraph(UpgradeInInstallMode.funTable("Switch#copy tftp: flash:\n"
                                                                 "Address or name of remote host []? 10.31.104.72\n"
                                                                 "Source filename []? cat9k_iosxe.17.03.05.SPA.bin\n"
                                                                 "Destination filename [cat9k_iosxe.17.03.05.SPA.bin]?\n"
                                                                 "Accessing tftp://10.31.104.72/cat9k_iosxe.17.03.05.SPA.bin...\n"
                                                                 "Loading cat9k_iosxe.17.03.05.SPA.bin from 10.31.104.72 (via GigabitEthernet0/0): !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n"
                                                                 "[OK - 906518831 bytes]\n"
                                                                 "\n"
                                                                 "906518831 bytes copied in 56.878 secs (15937952 bytes/sec)"))
            li2 = step2[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(2) > td:nth-child(2) > ol > li:nth-child(2)')
            document.add_paragraph('b.   '+li2[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to confirm that the image has been successfully copied to flash.')
            document.add_paragraph(UpgradeInInstallMode.funTable("Switch#dir flash:*.bin\n"
                                                                 "Directory of bootflash:/*.bin\n"
                                                                 "\n"
                                                                 "Directory of bootflash:/\n"
                                                                 "\n"
                                                                 "421849 -rw- 906518831 Mar 16 2022 04:31:34 +00:00 cat9k_iosxe.17.03.05.SPA.bin\n"
                                                                 "421852 -rw- 907450095 Feb 24 2022 17:40:57 +00:00 cat9k_iosxe.17.03.04.SPA.bin\n"
                                                                 "11250098176 bytes total (7851098112 bytes free)"))
            step3 = soup.select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 3 ').bold = True
            p.add_run(step3[0].find('p', class_="ph cmd").text)
            li1 = step3[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(1)')
            document.add_paragraph('a.   '+li1[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to set the boot variable to flash:packages.conf .')
            document.add_paragraph(UpgradeInInstallMode.funTable(li1[0].find('pre',class_="pre codeblock").text))
            li2 = step3[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(2)')
            document.add_paragraph('b.   '+li2[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to configure the switch to auto-boot. Settings are synchronized with the standby switch, if applicable.')
            document.add_paragraph(UpgradeInInstallMode.funTable(li2[0].find('pre',class_="pre codeblock").text))
            li3 = step3[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(3)')
            document.add_paragraph('c.   '+li3[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to save boot settings.')
            document.add_paragraph(UpgradeInInstallMode.funTable(li3[0].find('pre',class_="pre codeblock").text))
            li4 = step3[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(4)')
            document.add_paragraph('d.   '+li4[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to verify the boot variable (packages.conf) and manual boot setting (no):')
            document.add_paragraph(UpgradeInInstallMode.funTable(li4[0].find('pre',class_="pre codeblock").text))
            step4 = soup.select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(4) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 4 ').bold = True
            p.add_run(step4[0].find('p', class_="ph cmd").text)
            document.add_paragraph('Install image to flash\n'
                                   'install add file activate commit\n'
                                   'Use this command to install the image.\n'
                                   'We recommend that you point to the source image on a TFTP server or the flash , if you have copied the image to flash memory.\n\n'
                                   'The following sample output displays installation of the Cisco IOS XE Bengaluru 17.3.5 software image to flash:')
            document.add_paragraph(UpgradeInInstallMode.funTable("Switch#install add file flash:cat9k_iosxe.17.03.05.SPA.bin activate commit\n"
                                                                 "install_add_activate_commit: START Wed Mar 16 17:33:12 UTC 2022\n"
                                                                 "install_add_activate_commit: Adding PACKAGE\n"
                                                                 "install_add_activate_commit: Checking whether new add is allowed ....\n"
                                                                 "\n"
                                                                 "--- Starting initial file syncing ---\n"
                                                                 "Copying image file: bootflash:cat9k_iosxe.17.03.05.SPA.bin to standby\n"
                                                                 "Info: Finished copying bootflash:cat9k_iosxe.17.03.05.SPA.bin to standby\n"
                                                                 "Finished initial file syncing\n"
                                                                 "\n"
                                                                 "--- Starting Add ---\n"
                                                                 "Performing Add on Active/Standby\n"
                                                                 "[1] Add package(s) on R0\n"
                                                                 "[1] Finished Add on R0\n"
                                                                 "[1] Add package(s) on R1\n"
                                                                 "[1] Finished Add on R1\n"
                                                                 "Checking status of Add on [R0 R1]\n"
                                                                 "Add: Passed on [R0 R1]\n"
                                                                 "Finished Add\n"
                                                                 "\n"
                                                                 "Image added. Version: 17.03.05.0.6600\n"
                                                                 "install_add_activate_commit: Activating PACKAGE\n"
                                                                 "Following packages shall be activated:\n"
                                                                 "/bootflash/cat9k-wlc.17.03.05.SPA.pkg\n"
                                                                 "/bootflash/cat9k-webui.17.03.05.SPA.pkg\n"
                                                                 "/bootflash/cat9k-srdriver.17.03.05.SPA.pkg\n"
                                                                 "/bootflash/cat9k-sipspa.17.03.05.SPA.pkg\n"
                                                                 "/bootflash/cat9k-sipbase.17.03.05.SPA.pkg\n"
                                                                 "/bootflash/cat9k-rpboot.17.03.05.SPA.pkg\n"
                                                                 "/bootflash/cat9k-rpbase.17.03.05.SPA.pkg\n"
                                                                 "/bootflash/cat9k-lni.17.03.05.SPA.pkg\n"
                                                                 "/bootflash/cat9k-guestshell.17.03.05.SPA.pkg\n"
                                                                 "/bootflash/cat9k-espbase.17.03.05.SPA.pkg\n"
                                                                 "/bootflash/cat9k-cc_srdriver.17.03.05.SPA.pkg\n"
                                                                 "\n"
                                                                 "This operation may require a reload of the system. Do you want to proceed? [y/n]y\n"
                                                                 "--- Starting Activate ---\n"
                                                                 "Performing Activate on Active/Standby\n"
                                                                 "[1] Activate package(s) on R0\n"
                                                                 "--- Starting list of software package changes ---\n"
                                                                 "Old files list:\n"
                                                                 "Removed cat9k-cc_srdriver.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-espbase.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-guestshell.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-lni.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-rpbase.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-rpboot.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-sipbase.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-sipspa.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-srdriver.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-webui.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-wlc.17.03.04.SPA.pkg\n"
                                                                 "New files list:\n"
                                                                 "Added cat9k-cc_srdriver.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-espbase.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-guestshell.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-lni.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-rpbase.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-rpboot.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-sipbase.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-sipspa.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-srdriver.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-webui.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-wlc.17.03.05.SPA.pkg\n"
                                                                 "Finished list of software package changes\n"
                                                                 "[1] Finished Activate on R0\n"
                                                                 "[1] Activate package(s) on R1\n"
                                                                 "--- Starting list of software package changes ---\n"
                                                                 "Old files list:\n"
                                                                 "Removed cat9k-cc_srdriver.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-espbase.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-guestshell.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-lni.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-rpbase.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-rpboot.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-sipbase.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-sipspa.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-srdriver.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-webui.17.03.04.SPA.pkg\n"
                                                                 "Removed cat9k-wlc.17.03.04.SPA.pkg\n"
                                                                 "New files list:\n"
                                                                 "Added cat9k-cc_srdriver.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-espbase.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-guestshell.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-lni.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-rpbase.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-rpboot.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-sipbase.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-sipspa.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-srdriver.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-webui.17.03.05.SPA.pkg\n"
                                                                 "Added cat9k-wlc.17.03.05.SPA.pkg\n"
                                                                 "Finished list of software package changes\n"
                                                                 "[1] Finished Activate on R1\n"
                                                                 "Checking status of Activate on [R0 R1]\n"
                                                                 "Activate: Passed on [R0 R1]\n"
                                                                 "Finished Activate\n"
                                                                 "\n"
                                                                 "--- Starting Commit ---\n"
                                                                 "Performing Commit on Active/Standby\n"
                                                                 "[1] Commit package(s) on R0\n"
                                                                 "[1] Finished Commit on R0\n"
                                                                 "[1] Commit package(s) on R1\n"
                                                                 "[1] Finished Commit on R1\n"
                                                                 "Checking status of Commit on [R0 R1]\n"
                                                                 "Commit: Passed on [R0 R1]\n"
                                                                 "Finished Commit\n"
                                                                 "\n"
                                                                 "Send model notification for install_add_activate_commit before reload\n"
                                                                 "Install will reload the system now!\n"
                                                                 "SUCCESS: install_add_activate_commit Wed Mar 16 17:43:46 UTC 2022\n"))
            step5 = soup.select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(5) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 5 ').bold = True
            p.add_run(step5[0].find('p', class_="ph cmd").text)
            document.add_paragraph(step5[0].find('p', class_="p").text)
            li1 = step5[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(5) > td:nth-child(2) > ol > li:nth-child(1)')
            document.add_paragraph('a.   '+li1[0].find('span', class_="ph synph").text)
            document.add_paragraph('The following is sample output of the dir flash:*.pkg command:')
            document.add_paragraph(UpgradeInInstallMode.funTable("Switch# dir flash:*.pkg\n"
                                                                 "Directory of flash:/*.pkg\n"
                                                                 "Directory of flash:/\n"
                                                                 "475140 -rw- 2012104 Mar 31 2020 09:52:41 -07:00 cat9k-cc_srdriver.17.02.01.SPA.pkg\n"
                                                                 "475141 -rw- 70333380 Mar 31 2020 09:52:44 -07:00 cat9k-espbase.17.02.01.SPA.pkg\n"
                                                                 "475142 -rw- 13256 Mar 31 2020 09:52:44 -07:00 cat9k-guestshell.17.02.01.SPA.pkg\n"
                                                                 "475143 -rw- 349635524 Mar 31 2020 09:52:54 -07:00 cat9k-rpbase.17.02.01.SPA.pkg\n"
                                                                 "475149 -rw- 24248187 Mar 31 2020 09:53:02 -07:00 cat9k-rpboot.17.02.01.SPA.pkg\n"
                                                                 "475144 -rw- 25285572 Mar 31 2020 09:52:55 -07:00 cat9k-sipbase.17.02.01.SPA.pkg\n"
                                                                 "475145 -rw- 20947908 Mar 31 2020 09:52:55 -07:00 cat9k-sipspa.17.02.01.SPA.pkg\n"
                                                                 "475146 -rw- 2962372 Mar 31 2020 09:52:56 -07:00 cat9k-srdriver.17.02.01.SPA.pkg\n"
                                                                 "475147 -rw- 13284288 Mar 31 2020 09:52:56 -07:00 cat9k-webui.17.02.01.SPA.pkg\n"
                                                                 "475148 -rw- 13248 Mar 31 2020 09:52:56 -07:00 cat9k-wlc.17.02.01.SPA.pkg\n"
                                                                 "\n"
                                                                 "491524 -rw- 25711568 Mar 17 2022 11:49:33 -07:00 cat9k-cc_srdriver.17.03.05.SPA.pkg\n"
                                                                 "491525 -rw- 78484428 Mar 17 2022 11:49:35 -07:00 cat9k-espbase.17.03.05.SPA.pkg\n"
                                                                 "491526 -rw- 1598412 Mar 17 2022 11:49:35 -07:00 cat9k-guestshell.17.03.05.SPA.pkg\n"
                                                                 "491527 -rw- 404153288 Mar 17 2022 11:49:47 -07:00 cat9k-rpbase.17.03.05.SPA.pkg\n"
                                                                 "491533 -rw- 31657374 Mar 17 2022 11:50:09 -07:00 cat9k-rpboot.17.03.05.SPA.pkg\n"
                                                                 "491528 -rw- 27681740 Mar 17 2022 11:49:48 -07:00 cat9k-sipbase.17.03.05.SPA.pkg\n"
                                                                 "491529 -rw- 52224968 Mar 17 2022 11:49:49 -07:00 cat9k-sipspa.17.03.05.SPA.pkg\n"
                                                                 "491530 -rw- 31130572 Mar 17 2022 11:49:50 -07:00 cat9k-srdriver.17.03.05.SPA.pkg\n"
                                                                 "491531 -rw- 14783432 Mar 17 2022 11:49:51 -07:00 cat9k-webui.17.03.05.SPA.pkg\n"
                                                                 "491532 -rw- 9160 Mar 17 2022 11:49:51 -07:00 cat9k-wlc.17.03.05.SPA.pkg\n"
                                                                 "11353194496 bytes total (8963174400 bytes free)"))
            li2 = step5[0].select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(5) > td:nth-child(2) > ol > li:nth-child(2)')
            document.add_paragraph('b.   '+li2[0].find('span', class_="ph synph").text)
            document.add_paragraph('The following is sample output of the dir flash:*.conf command. It displays the .conf files in the flash partition; note the two .conf files:\n'
                                   'a.    packages.conf—the file that has been re-written with the newly installed .pkg files.  \n' 
                                   'b.    cat9k_iosxe.17.03.05.SPA.conf— a backup copy of the newly installed packages.conf file.')
            document.add_paragraph(UpgradeInInstallMode.funTable("Switch# dir flash:*.conf\n"
                                                                 "\n"
                                                                 "Directory of flash:/*.conf\n"
                                                                 "Directory of flash:/\n"
                                                                 "\n"
                                                                 "16631 -rw- 4882 Mar 17 2022 05:39:42 +00:00 packages.conf\n"
                                                                 "16634 -rw- 4882 Mar 17 2022 05:34:06 +00:00 cat9k_iosxe.17.03.05.SPA.conf"))
            step6 = soup.select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(6) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 6 ').bold = True
            p.add_run(step6[0].find('p', class_="ph cmd").text)
            document.add_paragraph(step6[0].find('span',class_='synph').text)
            document.add_paragraph(step6[0].find('p', class_="p").text.replace("17.3.1","17.3.5"))
            document.add_paragraph('In case of a high availability set up or a Cisco StackWise Virtual set up, remember to upgrade the active and standby.')
            document.add_paragraph(UpgradeInInstallMode.funTable(step6[0].find('pre',class_="pre codeblock").text))

            step7 = soup.select('#task_a2s_dfh_jmb > div > table > tbody > tr:nth-child(7) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 7 ').bold = True
            p.add_run(step7[0].find('p', class_="ph cmd").text)
            document.add_paragraph(step7[0].find('p', class_="p").text)
            document.add_paragraph('After the image boots up, use this command to verify the version of the new image.\n')
            document.add_paragraph('The following sample output of the show version command displays the Cisco IOS XE Bengaluru 17.3.5 image on the device:')
            document.add_paragraph(UpgradeInInstallMode.funTable("Switch# show version\n"
                                                                 "Cisco IOS XE Software, Version 17.03.05\n"
                                                                 "Cisco IOS Software [Amsterdam], Catalyst L3 Switch Software (CAT9K_IOSXE), Version 17.3.5, RELEASE SOFTWARE (fc2)\n"
                                                                 "Technical Support: http://www.cisco.com/techsupport\n"
                                                                 "Copyright (c) 1986-2022 by Cisco Systems, Inc..\n"
                                                                 "<output truncated>\n"))
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Upgrading In Install Mode' section")
        document.add_page_break()
        return    
    def upgradeInInstallMode_Switch_Catalyst_9300_17_03_04(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9300/software/release/17-3/release_notes/ol-17-3-9300.html')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#task_xpb_3tx_xkb__d54e6074 > p')
            table1 = soup.select('#task_xpb_3tx_xkb__d54e5964')
            l = table1[0].text.splitlines()

            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele)
            document.add_paragraph(elem[0].text.replace('\n',''))
            p = document.add_paragraph()
            p.add_run(t1[0]).bold =True
            document.add_paragraph(t1[1])
            table = document.add_table(rows=1, cols=3, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = t1[2]
            row[1].text = t1[3]
            row[2].text = t1[4]
            row = table.add_row().cells
            row[0].text = t1[5]
            row[1].text = t1[6]
            row = table.add_row().cells
            row[0].text = t1[8]
            row[1].text = t1[9]
            a=table.cell(1,2)
            b= table.cell(2,2)
            A = a.merge(b)
            A.text = t1[7]
            document.add_paragraph(t1[10].replace('5','')+'\n'+t1[11])
            step1 = soup.select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(1) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 1 ').bold = True
            p.add_run(step1[0].find('p', class_="ph cmd").text)
            document.add_paragraph(step1[0].find('span', class_="ph synph").text)
            document.add_paragraph(re.sub(' +', ' ', step1[0].find('p', class_='p').text.replace('\n','')))
            document.add_paragraph('The following sample output displays the cleaning up of unused files, by using the install remove inactive command:')
            document.add_paragraph(UpgradeInInstallMode.funTable(step1[0].find('pre',class_="pre codeblock").text))
            step2 = soup.select("#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(2) > td:nth-child(2)")
            p = document.add_paragraph("\n")
            p.add_run('Step 2 ').bold = True
            p.add_run(step2[0].find('p', class_="ph cmd").text)
            li1 = step2[0].select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(2) > td:nth-child(2) > ol > li:nth-child(1)')
            document.add_paragraph('a.   '+li1[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to copy the new image from a TFTP server to flash memory. The location is either an IP address or a host name. The filename is specified relative to the directory used for file transfers. Skip this step if you want to use the new image from a TFTP server.')
            document.add_paragraph(UpgradeInInstallMode.funTable('Switch# copy tftp://10.8.0.6/image/cat9k_iosxe.17.03.04.SPA.bin flash:\n'
                                                                 'destination filename [cat9k_iosxe.17.03.04.SPA.bin]?\n'
                                                                 'Accessing tftp://10.8.0.6/image/cat9k_iosxe.17.03.04.SPA.bin...\n'
                                                                 'Loading /cat9k_iosxe.17.03.04.SPA.bin from 10.8.0.6 (via GigabitEthernet0/0):\n'
                                                                 '!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n'
                                                                 '[OK - 907450095 bytes]\n'
                                                                 '\n'
                                                                 '907450095 bytes copied in 50.649 secs (11870255 bytes/sec)\n'))
            li2 = step2[0].select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(2) > td:nth-child(2) > ol > li:nth-child(2)')
            document.add_paragraph('b.   '+li2[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to confirm that the image has been successfully copied to flash.')
            document.add_paragraph(UpgradeInInstallMode.funTable('Switch# dir flash:*.bin\n'
                                                                 'Directory of flash:/*.bin\n'
                                                                 '\n'
                                                                 'Directory of flash:/\n'
                                                                 '\n'
                                                                 '434184 -rw- 907450095 Jul 20 2020 10:18:11 -07:00 cat9k_iosxe.17.03.04.SPA.bin\n'
                                                                 '11353194496 bytes total (8976625664 bytes free)'))
            step3 = soup.select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(3) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 3 ').bold = True
            p.add_run(step3[0].find('p', class_="ph cmd").text)
            li1 = step3[0].select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(1)')
            document.add_paragraph('a.   '+li1[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to set the boot variable to flash:packages.conf .')
            document.add_paragraph(UpgradeInInstallMode.funTable(li1[0].find('pre',class_="pre codeblock").text))
            li2 = step3[0].select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(2)')
            document.add_paragraph('b.   '+li2[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to configure the switch to auto-boot. Settings are synchronized with the standby switch, if applicable.')
            document.add_paragraph(UpgradeInInstallMode.funTable(li2[0].find('pre',class_="pre codeblock").text))
            li3 = step3[0].select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(3)')
            document.add_paragraph('c.   '+li3[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to save boot settings.')
            document.add_paragraph(UpgradeInInstallMode.funTable(li3[0].find('pre',class_="pre codeblock").text))
            li4 = step3[0].select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(4)')
            document.add_paragraph('d.   '+li4[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to verify the boot variable (packages.conf) and manual boot setting (no):')
            document.add_paragraph(UpgradeInInstallMode.funTable('Switch#sh boot\n'
                                   '---------------------------\n'
                                   'Switch 1\n'
                                   '---------------------------\n'
                                   'Current Boot Variables:\n'
                                   'BOOT variable = flash:packages.conf;\n'
                                   '\n'
                                   'Boot Variables on next reload:\n'
                                   'BOOT variable = flash:packages.conf;\n'
                                   'Manual Boot = no\n'
                                   'Enable Break = yes\n'
                                   'Boot Mode = DEVICE\n'
                                   'iPXE Timeout = 0\n'))
            step4 = soup.select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(4) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 4 ').bold = True
            p.add_run(step4[0].find('p', class_="ph cmd").text)
            document.add_paragraph("install add file activate commit\n"
                                   "Use this command to install the image.\n"
                                   "We recommend that you point to the source image on your TFTP server or the flash drive of the active switch, if you have copied the image to flash memory. If you point to an image on the flash or USB drive of a member switch (instead of the active), you must specify the exact flash or USB drive - otherwise installation fails. For example, if the image is on the flash drive of member switch 3 (flash-3): Switch# install add file flash-3:cat9k_iosxe.17.03.04.SPA.bin activate commit.\n"
                                   "The following sample output displays installation of the Cisco IOS XE Amsterdam 17.3.4 software image in the flash memory:\n")
            document.add_paragraph(UpgradeInInstallMode.funTable('Switch#install add file flash:cat9k_iosxe.17.03.04.SPA.bin activate commit\n'
                                                                 'install_add_activate_commit: Adding PACKAGE\n'
                                                                 'install_add_activate_commit: Checking whether new add is allowed ....\n'
                                                                 'install_add_activate_commit: START Mon Apr 4 08:57:20 UTC 2022\n'
                                                                 '\n'
                                                                 '--- Starting initial file syncing ---\n'
                                                                 'Info: Finished copying flash:cat9k_iosxe.17.03.04.SPA.bin to the selected switch(es)\n'
                                                                 'Finished initial file syncing\n'
                                                                 '\n'
                                                                 '--- Starting Add ---\n'
                                                                 'Performing Add on all members\n'
                                                                 '[1] Add package(s) on switch 1\n'
                                                                 '[1] Finished Add on switch 1\n'
                                                                 'Checking status of Add on [1]\n'
                                                                 'Add: Passed on [1]\n'
                                                                 'Finished Add\n'
                                                                 '\n'
                                                                 'Image added. Version: 17.03.04.0.5557\n'
                                                                 'install_add_activate_commit: Activating PACKAGE\n'
                                                                 'Following packages shall be activated:\n'
                                                                 '/flash/cat9k-wlc.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-webui.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-srdriver.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-sipspa.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-sipbase.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-rpboot.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-rpbase.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-lni.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-guestshell.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-espbase.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-cc_srdriver.17.03.04.SPA.pkg\n'
                                                                 '\n'
                                                                 'This operation may require a reload of the system. Do you want to proceed? [y/n]y\n'
                                                                 '--- Starting Activate ---\n'
                                                                 'Performing Activate on all members\n'
                                                                 '[1] Activate package(s) on switch 1\n'
                                                                 '--- Starting list of software package changes ---\n'
                                                                 'Old files list:\n'
                                                                 'Removed cat9k-cc_srdriver.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-espbase.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-guestshell.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-lni.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-rpbase.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-rpboot.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-sipbase.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-sipspa.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-srdriver.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-webui.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-wlc.16.12.04.SPA.pkg\n'
                                                                 'New files list:\n'
                                                                 'Added cat9k-cc_srdriver.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-espbase.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-guestshell.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-lni.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-rpbase.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-rpboot.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-sipbase.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-sipspa.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-srdriver.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-webui.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-wlc.17.03.04.SPA.pkg\n'
                                                                 'Finished list of software package changes\n'
                                                                 '[1] Finished Activate on switch 1\n'
                                                                 'Checking status of Activate on [1]\n'
                                                                 'Activate: Passed on [1]\n'
                                                                 'Finished Activate\n'
                                                                 '\n'
                                                                 '--- Starting Commit ---\n'
                                                                 'Performing Commit on all members\n'
                                                                 '[1] Commit package(s) on switch 1\n'
                                                                 '[1] Finished Commit on switch 1\n'
                                                                 'Checking status of Commit on [1]\n'
                                                                 'Commit: Passed on [1]\n'
                                                                 'Finished Commit\n'
                                                                 '\n'
                                                                 'Send model notification for install_add_activate_commit before reload\n'
                                                                 '[1]: Performing Upgrade_Service\n'
                                                                 '300+0 records in\n'
                                                                 '300+0 records out\n'
                                                                 '307200 bytes (307 kB, 300 KiB) copied, 0.315729 s, 973 kB/s\n'
                                                                 'SUCCESS: Upgrade_Service finished\n'
                                                                 'Install will reload the system now!\n'
                                                                 'SUCCESS: install_add_activate_commit Mon Apr 4 09:05:27 UTC 2022\n'
                                                                 'C9300-2#\n'
                                                                 'Chassis 1 reloading, reason - Reload command\n'
                                                                 'Apr 4 09:05:31.172: %PMAN-5-EXITACTION: F0/0: pvp: Process manager is exiting: reload fp action requested\n'
                                                                 'Apr 4 09:05:32.182: %PMAN-5-EXITvp: Process manager is exiting: rp processes exit with reload switch code\n'
                                                                 '\n'
                                                                 '\n'
                                                                 '\n'
                                                                 'Initializing Hardware......\n'
                                                                 '\n'
                                                                 'System Bootstrap, Version 17.5.2r, RELEASE SOFTWARE (P)\n'
                                                                 'Compiled Tue 02/02/2021 15:07:07.87 by rel\n'
                                                                 '\n'
                                                                 'Current ROMMON image : Primary\n'
                                                                 'Last reset cause : SoftwareReload\n'
                                                                 'C9300-48U platform with 8388608 Kbytes of main memory\n'
                                                                 '\n'
                                                                 'Preparing to autoboot. [Press Ctrl-C to interrupt] 0\n'
                                                                 'boot: attempting to boot from [flash:packages.conf]\n'
                                                                 'boot: reading file packages.conf\n'
                                                                 '#\n'
                                                                 '#########################################################################################################\n'######################################################################################################################################################################################################################################################################################################################################################################
                                                                 '\n'
                                                                 '\n'
                                                                 'Both links down, not waiting for other switches\n'
                                                                 'Switch number is 1\n'
                                                                 '<snip>\n'))
            step5 = soup.select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(5) > td:nth-child(2)')
            p = document.add_paragraph("\n") 
            p.add_run('Step 5 ').bold = True 
            p.add_run(step5[0].find('p', class_="ph cmd").text) 
            document.add_paragraph(step5[0].find('p', class_="p").text)
            li1 = step5[0].select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(5) > td:nth-child(2) > ol > li:nth-child(1)')
            document.add_paragraph('a.   '+li1[0].find('span', class_="ph synph").text)
            document.add_paragraph('The following is sample output of the dir flash:*.pkg command:')
            document.add_paragraph(UpgradeInInstallMode.funTable('Switch#dir flash:*.pkg\n'
                                                                 'Directory of flash:/*.pkg\n'
                                                                 '\n'
                                                                 'Directory of flash:/\n'
                                                                 '\n'
                                                                 '155671 -rw- 17392652 Apr 4 2022 08:34:48 +00:00 cat9k-cc_srdriver.16.12.04.SPA.pkg\n'
                                                                 '155672 -rw- 104428552 Apr 4 2022 08:34:48 +00:00 cat9k-espbase.16.12.04.SPA.pkg\n'
                                                                 '155673 -rw- 2262024 Apr 4 2022 08:34:48 +00:00 cat9k-guestshell.16.12.04.SPA.pkg\n'
                                                                 '155674 -rw- 5124 Apr 4 2022 08:34:48 +00:00 cat9k-lni.16.12.04.SPA.pkg\n'
                                                                 '155675 -rw- 595178500 Apr 4 2022 08:34:49 +00:00 cat9k-rpbase.16.12.04.SPA.pkg\n'
                                                                 '155682 -rw- 47364227 Apr 4 2022 08:35:48 +00:00 cat9k-rpboot.16.12.04.SPA.pkg\n'
                                                                 '155676 -rw- 34792456 Apr 4 2022 08:34:49 +00:00 cat9k-sipbase.16.12.04.SPA.pkg\n'
                                                                 '155677 -rw- 57529348 Apr 4 2022 08:34:49 +00:00 cat9k-sipspa.16.12.04.SPA.pkg\n'
                                                                 '155678 -rw- 28738568 Apr 4 2022 08:34:49 +00:00 cat9k-srdriver.16.12.04.SPA.pkg\n'
                                                                 '155679 -rw- 14427140 Apr 4 2022 08:34:49 +00:00 cat9k-webui.16.12.04.SPA.pkg\n'
                                                                 '155681 -rw- 9220 Apr 4 2022 08:34:49 +00:00 cat9k-wlc.16.12.04.SPA.pkg\n'
                                                                 '327693 -rw- 5124 Apr 4 2022 08:58:47 +00:00 cat9k-lni.17.03.04.SPA.pkg\n'
                                                                 '327690 -rw- 18289676 Apr 4 2022 08:58:47 +00:00 cat9k-cc_srdriver.17.03.04.SPA.pkg\n'
                                                                 '327691 -rw- 104526856 Apr 4 2022 08:58:47 +00:00 cat9k-espbase.17.03.04.SPA.pkg\n'
                                                                 '327692 -rw- 2262024 Apr 4 2022 08:58:47 +00:00 cat9k-guestshell.17.03.04.SPA.pkg\n'
                                                                 '327694 -rw- 599221252 Apr 4 2022 08:58:48 +00:00 cat9k-rpbase.17.03.04.SPA.pkg\n'
                                                                 '327700 -rw- 47407779 Apr 4 2022 08:59:46 +00:00 cat9k-rpboot.17.03.04.SPA.pkg\n'
                                                                 '327695 -rw- 34874376 Apr 4 2022 08:58:48 +00:00 cat9k-sipbase.17.03.04.SPA.pkg\n'
                                                                 '327696 -rw- 57082884 Apr 4 2022 08:58:48 +00:00 cat9k-sipspa.17.03.04.SPA.pkg\n'
                                                                 '327697 -rw- 29717512 Apr 4 2022 08:58:48 +00:00 cat9k-srdriver.17.03.04.SPA.pkg\n'
                                                                 '327698 -rw- 14431236 Apr 4 2022 08:58:48 +00:00 cat9k-webui.17.03.04.SPA.pkg\n'
                                                                 '327699 -rw- 9220 Apr 4 2022 08:58:48 +00:00 cat9k-wlc.17.03.04.SPA.pkg\n'
                                                                 '11353194496 bytes total (6882414592 bytes free)\n'))
            li2 = step5[0].select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(5) > td:nth-child(2) > ol > li:nth-child(2)')
            document.add_paragraph('b.   '+li2[0].find('span', class_="ph synph").text)
            document.add_paragraph('The following is sample output of the dir flash:*.conf command. It displays the .conf files in the flash partition; note the two .conf files:\n'
                                   'a.    packages.conf—the file that has been re-written with the newly installed .pkg files.  \n' 
                                   'b.    cat9k_iosxe.17.03.04.SPA.conf— a backup copy of the newly installed packages.conf file.')
            document.add_paragraph(UpgradeInInstallMode.funTable('Switch#dir flash:*.conf\n'
                                                                 'Directory of flash:/*.conf\n'
                                                                 '\n'
                                                                 'Directory of flash:/\n'
                                                                 '\n'
                                                                 '155689 -rw- 7715 Apr 4 2022 09:04:35 +00:00 packages.conf\n'
                                                                 '327689 -rw- 7715 Apr 4 2022 08:59:46 +00:00 cat9k_iosxe.17.03.04.SPA.conf\n'
                                                                 '11353194496 bytes total (6882414592 bytes free)\n'))
            step6 = soup.select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(6) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 6 ').bold = True
            p.add_run(step6[0].find('p', class_="ph cmd").text)
            document.add_paragraph(step6[0].find('span',class_='synph').text)
            document.add_paragraph(step6[0].find('p', class_="p").text.replace("17.3.1","17.3.4"))
            document.add_paragraph('A new ROMMON version is available in Cisco IOS XE Amsterdam 17.3.4, for only the C9300 models in the series. After you enter the command, confirm upgrade at the system prompt.')
            document.add_paragraph(UpgradeInInstallMode.funTable(step6[0].find('pre',class_="pre codeblock").text))

            step7 = soup.select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(7) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 7 ').bold = True
            p.add_run(step7[0].find('p', class_="ph cmd").text)
            li1 = step7[0].select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(7) > td:nth-child(2) > ol > li:nth-child(1)')#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(7) > td:nth-child(2) > ol > li:nth-child(1)
            document.add_paragraph('a.   '+li1[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to reload the switch. When the switch reloads after a ROMMON upgrade, the ROMMON version is updated, but not displayed in the output until the next reload.')
            document.add_paragraph(UpgradeInInstallMode.funTable(li1[0].find('pre',class_="pre codeblock").text))
            li2 = step7[0].select('#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(7) > td:nth-child(2) > ol > li:nth-child(2)')
            document.add_paragraph('b.   '+li2[0].find('span', class_="ph synph").text)
            document.add_paragraph(step7[0].find('p', class_="p").text)
            document.add_paragraph('The following sample output of the show version command displays the Cisco IOS XE Amsterdam 17.3.4 image on the device:')
            document.add_paragraph(UpgradeInInstallMode.funTable("Switch# show version\n"
                                                                 "Cisco IOS XE Software, Version 17.03.04\n"
                                                                 "Cisco IOS Software [Amsterdam], Catalyst L3 Switch Software (CAT9K_IOSXE), Version 17.3.4, RELEASE SOFTWARE (fc2)\n"
                                                                 "Technical Support: http://www.cisco.com/techsupport\n"
                                                                 "Copyright (c) 1986-2022 by Cisco Systems, Inc..\n"
                                                                 "<output truncated>\n"))
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Upgrading In Install Mode' section")
        document.add_page_break()
        return  
    def upgradeInInstallMode_Switch_Catalyst_9500_17_03_04(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9500/software/release/17-3/release_notes/ol-17-3-9500.html')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#task_p3r_f21_jmb__d54e5695 > p')
            table1 = soup.select('#task_p3r_f21_jmb__d54e5558')
            l = table1[0].text.splitlines()

            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele)
            document.add_paragraph(elem[0].text.replace('\n',''))
            p = document.add_paragraph()
            p.add_run(t1[0]).bold =True
            table = document.add_table(rows=1, cols=3, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = t1[1]
            row[1].text = t1[2]
            row[2].text = t1[3]
            row = table.add_row().cells
            row[0].text = t1[4]
            row[1].text = t1[5]
            row = table.add_row().cells
            row[0].text = t1[7]
            row[1].text = t1[8].replace('5','')+'\n'+t1[9].replace('6','')
            a=table.cell(1,2)
            b= table.cell(2,2)
            A = a.merge(b)
            A.text = t1[6]
            document.add_paragraph(t1[10].replace('5','')+'\n')
            document.add_paragraph(t1[12])
            document.add_paragraph(t1[13],style='List Bullet 2')
            document.add_paragraph(t1[14],style='List Bullet 2')
            document.add_paragraph(t1[15],style='List Bullet 2')
            document.add_paragraph(t1[16])
            step1 = soup.select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(1) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 1 ').bold = True
            document.add_paragraph(step1[0].find('p', class_="ph cmd").text)
            document.add_paragraph(step1[0].find('span', class_="ph synph").text)
            document.add_paragraph(re.sub(' +', ' ', step1[0].find('p', class_='p').text.replace('\n','')))
            document.add_paragraph('The following sample output displays the cleaning up of unused files, by using the install remove inactive command:')
            document.add_paragraph(UpgradeInInstallMode.funTable(step1[0].find('pre',class_="pre codeblock").text))
            step2 = soup.select("#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(2) > td:nth-child(2)")
            p = document.add_paragraph("\n")
            p.add_run('Step 2 ').bold = True
            document.add_paragraph(step2[0].find('p', class_="ph cmd").text)
            li1 = step2[0].select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(2) > td:nth-child(2) > ol > li:nth-child(1)')
            document.add_paragraph('a.   '+li1[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to copy the new image from a TFTP server to flash memory. The location is either an IP address or a host name. The filename is specified relative to the directory used for file transfers. Skip this step if you want to use the new image from a TFTP server.')
            document.add_paragraph(UpgradeInInstallMode.funTable('Switch# copy tftp://10.8.0.6/image/cat9k_iosxe.17.03.04.SPA.bin flash:\n'
                                                                 'destination filename [cat9k_iosxe.17.03.04.SPA.bin]?\n'
                                                                 'Accessing tftp://10.8.0.6/image/cat9k_iosxe.17.03.04.SPA.bin...\n'
                                                                 'Loading /cat9k_iosxe.17.03.04.SPA.bin from 10.8.0.6 (via GigabitEthernet0/0):\n'
                                                                 '!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n'
                                                                 '[OK - 907450095 bytes]\n'
                                                                 '\n'
                                                                 '907450095 bytes copied in 50.649 secs (11870255 bytes/sec)\n'))
            li2 = step2[0].select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(2) > td:nth-child(2) > ol > li:nth-child(2)')
            document.add_paragraph('b.   '+li2[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to confirm that the image has been successfully copied to flash.')
            document.add_paragraph(UpgradeInInstallMode.funTable('Switch# dir flash:*.bin\n'
                                                                 'Directory of flash:/*.bin\n'
                                                                 '\n'
                                                                 'Directory of flash:/\n'
                                                                 '\n'
                                                                 '434184 -rw- 907450095 Jul 20 2020 10:18:11 -07:00 cat9k_iosxe.17.03.04.SPA.bin\n'
                                                                 '11353194496 bytes total (8976625664 bytes free)'))
            step3 = soup.select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 3 ').bold = True
            p.add_run(step3[0].find('p', class_="ph cmd").text)
            li1 = step3[0].select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(1)')
            document.add_paragraph('a.   '+li1[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to set the boot variable to flash:packages.conf .')
            document.add_paragraph(UpgradeInInstallMode.funTable(li1[0].find('pre',class_="pre codeblock").text))
            li2 = step3[0].select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(2)')
            document.add_paragraph('b.   '+li2[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to configure the switch to auto-boot. Settings are synchronized with the standby switch, if applicable.')
            document.add_paragraph(UpgradeInInstallMode.funTable(li2[0].find('pre',class_="pre codeblock").text))
            li3 = step3[0].select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(3)')
            document.add_paragraph('c.   '+li3[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to save boot settings.')
            document.add_paragraph(UpgradeInInstallMode.funTable(li3[0].find('pre',class_="pre codeblock").text))
            li4 = step3[0].select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > ol > li:nth-child(4)')
            document.add_paragraph('d.   '+li4[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to verify the boot variable (packages.conf) and manual boot setting (no):')
            document.add_paragraph(UpgradeInInstallMode.funTable('Switch#show boot\n'
                                                                 '---------------------------\n'
                                                                 'Switch 1\n'
                                                                 '---------------------------\n'
                                                                 'Current Boot Variables:\n'
                                                                 'BOOT variable = flash:packages.conf;\n'
                                                                 '\n'
                                                                 'Boot Variables on next reload:\n'
                                                                 'BOOT variable = flash:packages.conf;\n'
                                                                 'Manual Boot = no\n'
                                                                 'Enable Break = yes\n'
                                                                 'Boot Mode = DEVICE\n'
                                                                 'iPXE Timeout = 0\n'))
            step4 = soup.select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(4) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 4 ').bold = True
            p.add_run(step4[0].find('p', class_="ph cmd").text)
            document.add_paragraph("install add file activate commit\n"
                                   "Use this command to install the image.\n"
                                   "We recommend that you point to the source image on your TFTP server or the flash drive of the active switch, if you have copied the image to flash memory. If you point to an image on the flash or USB drive of a member switch (instead of the active), you must specify the exact flash or USB drive - otherwise installation fails. For example, if the image is on the flash drive of member switch 3 (flash-3): Switch# install add file flash-3:cat9k_iosxe.17.03.04.SPA.bin activate commit.\n"
                                   "The following sample output displays installation of the Cisco IOS XE Amsterdam 17.3.4 software image in the flash memory:\n")
            document.add_paragraph(UpgradeInInstallMode.funTable('Switch#install add file flash:cat9k_iosxe.17.03.04.SPA.bin activate commit\n'
                                                                 'install_add_activate_commit: START Wed Apr 6 14:17:06 UTC 2022\n'
                                                                 'install_add_activate_commit: Adding PACKAGE\n'
                                                                 'install_add_activate_commit: Checking whether new add is allowed ....\n'
                                                                 '\n'
                                                                 '--- Starting initial file syncing ---\n'
                                                                 'Info: Finished copying flash:cat9k_iosxe.17.03.04.SPA.bin to the selected switch(es)\n'
                                                                 'Finished initial file syncing\n'
                                                                 '\n'
                                                                 '--- Starting Add ---\n'
                                                                 'Performing Add on all members\n'
                                                                 '[1] Add package(s) on switch 1\n'
                                                                 '[1] Finished Add on switch 1\n'
                                                                 'Checking status of Add on [1]\n'
                                                                 'Add: Passed on [1]\n'
                                                                 'Finished Add\n'
                                                                 '\n'
                                                                 'Image added. Version: 17.03.04.0.5557\n'
                                                                 'install_add_activate_commit: Activating PACKAGE\n'
                                                                 'Following packages shall be activated:\n'
                                                                 '/flash/cat9k-wlc.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-webui.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-srdriver.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-sipspa.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-sipbase.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-rpboot.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-rpbase.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-lni.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-guestshell.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-espbase.17.03.04.SPA.pkg\n'
                                                                 '/flash/cat9k-cc_srdriver.17.03.04.SPA.pkg\n'
                                                                 '\n'
                                                                 'This operation may require a reload of the system. Do you want to proceed? [y/n]y\n'
                                                                 '--- Starting Activate ---\n'
                                                                 'Performing Activate on all members\n'
                                                                 '[1] Activate package(s) on switch 1\n'
                                                                 '--- Starting list of software package changes ---\n'
                                                                 'Old files list:\n'
                                                                 'Removed cat9k-cc_srdriver.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-espbase.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-guestshell.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-lni.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-rpbase.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-rpboot.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-sipbase.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-sipspa.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-srdriver.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-webui.16.12.04.SPA.pkg\n'
                                                                 'Removed cat9k-wlc.16.12.04.SPA.pkg\n'
                                                                 'New files list:\n'
                                                                 'Added cat9k-cc_srdriver.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-espbase.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-guestshell.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-lni.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-rpbase.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-rpboot.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-sipbase.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-sipspa.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-srdriver.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-webui.17.03.04.SPA.pkg\n'
                                                                 'Added cat9k-wlc.17.03.04.SPA.pkg\n'
                                                                 'Finished list of software package changes\n'
                                                                 '[1] Finished Activate on switch 1\n'
                                                                 'Checking status of Activate on [1]\n'
                                                                 'Activate: Passed on [1]\n'
                                                                 'Finished Activate\n'
                                                                 '\n'
                                                                 '--- Starting Commit ---\n'
                                                                 'Performing Commit on all members\n'
                                                                 '[1] Commit package(s) on switch 1\n'
                                                                 '[1] Finished Commit on switch 1\n'
                                                                 'Checking status of Commit on [1]\n'
                                                                 'Commit: Passed on [1]\n'
                                                                 'Finished Commit\n'
                                                                 '\n'
                                                                 'Send model notification for install_add_activate_commit before reload\n'
                                                                 '[1]: Performing Upgrade_Service\n'
                                                                 '300+0 records in\n'
                                                                 '300+0 records out\n'
                                                                 '307200 bytes (307 kB, 300 KiB) copied, 0.320061 s, 960 kB/s\n'
                                                                 'SUCCESS: Upgrade_Service finished\n'
                                                                 'Install will reload the system now!\n'
                                                                 'SUCCESS: install_add_activate_commit Wed Apr 6 14:22:37 UTC 2022\n'
                                                                 'Switch#\n'
                                                                 'Chassis 1 reloading, reason - Reload command\n'
                                                                 'Apr 6 14:22:39.923: %PMAN-5-EXITACTION: F0/0: pvp: Process manager is exiting: reload fp action requested\n'
                                                                 'Apr 6 14:22:41.420: %PMAN-5-EXTACTION: R0/0: pvp: Process manager is exiting: rp processes exit with reload switch code\n'
                                                                 '\n'
                                                                 '\n'
                                                                 '\n'
                                                                 'Initializing Hardware...\n'
                                                                 '\n'
                                                                 'System Bootstrap, Version 16.12.2r, RELEASE SOFTWARE (P)\n'
                                                                 'Compiled Wed 10/23/2019 16:35:17.50 by rel\n'
                                                                 '\n'
                                                                 'Current ROMMON image : Primary\n'
                                                                 'Last reset cause : CpuReset\n'
                                                                 'C9500-16X platform with 16777216 Kbytes of main memory\n'
                                                                 '\n'
                                                                 'Preparing to autoboot. [Press Ctrl-C to interrupt] 0\n'
                                                                 'boot: attempting to boot from [flash:packages.conf]\n'
                                                                 'boot: reading file packages.conf\n'
                                                                 '#\n'
                                                                 '###############################################################################################################################################################################################################################################################################################################################################################################################################################################################################\n'
                                                                 '\n'
                                                                 '\n'
                                                                 'Both links down, not waiting for other switches\n'
                                                                 'Switch number is 1\n'
                                                                 '<snip>\n'))
            step5 = soup.select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(5) > td:nth-child(2)')
            p = document.add_paragraph("\n") 
            p.add_run('Step 5 ').bold = True 
            p.add_run(step5[0].find('p', class_="ph cmd").text) 
            document.add_paragraph(step5[0].find('p', class_="p").text)
            li1 = step5[0].select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(5) > td:nth-child(2) > ol > li:nth-child(1)')
            document.add_paragraph('a.   '+li1[0].find('span', class_="ph synph").text)
            document.add_paragraph('The following is sample output of the dir flash:*.pkg command:')
            document.add_paragraph(UpgradeInInstallMode.funTable('Switch#dir flash:*.pkg\n'
                                                                 'Directory of flash:/*.pkg\n'
                                                                 '\n'
                                                                 'Directory of flash:/\n'
                                                                 '\n'
                                                                 '155671 -rw- 17392652 Apr 4 2022 08:34:48 +00:00 cat9k-cc_srdriver.16.12.04.SPA.pkg\n'
                                                                 '155672 -rw- 104428552 Apr 4 2022 08:34:48 +00:00 cat9k-espbase.16.12.04.SPA.pkg\n'
                                                                 '155673 -rw- 2262024 Apr 4 2022 08:34:48 +00:00 cat9k-guestshell.16.12.04.SPA.pkg\n'
                                                                 '155674 -rw- 5124 Apr 4 2022 08:34:48 +00:00 cat9k-lni.16.12.04.SPA.pkg\n'
                                                                 '155675 -rw- 595178500 Apr 4 2022 08:34:49 +00:00 cat9k-rpbase.16.12.04.SPA.pkg\n'
                                                                 '155682 -rw- 47364227 Apr 4 2022 08:35:48 +00:00 cat9k-rpboot.16.12.04.SPA.pkg\n'
                                                                 '155676 -rw- 34792456 Apr 4 2022 08:34:49 +00:00 cat9k-sipbase.16.12.04.SPA.pkg\n'
                                                                 '155677 -rw- 57529348 Apr 4 2022 08:34:49 +00:00 cat9k-sipspa.16.12.04.SPA.pkg\n'
                                                                 '155678 -rw- 28738568 Apr 4 2022 08:34:49 +00:00 cat9k-srdriver.16.12.04.SPA.pkg\n'
                                                                 '155679 -rw- 14427140 Apr 4 2022 08:34:49 +00:00 cat9k-webui.16.12.04.SPA.pkg\n'
                                                                 '155681 -rw- 9220 Apr 4 2022 08:34:49 +00:00 cat9k-wlc.16.12.04.SPA.pkg\n'
                                                                 '327693 -rw- 5124 Apr 4 2022 08:58:47 +00:00 cat9k-lni.17.03.04.SPA.pkg\n'
                                                                 '327690 -rw- 18289676 Apr 4 2022 08:58:47 +00:00 cat9k-cc_srdriver.17.03.04.SPA.pkg\n'
                                                                 '327691 -rw- 104526856 Apr 4 2022 08:58:47 +00:00 cat9k-espbase.17.03.04.SPA.pkg\n'
                                                                 '327692 -rw- 2262024 Apr 4 2022 08:58:47 +00:00 cat9k-guestshell.17.03.04.SPA.pkg\n'
                                                                 '327694 -rw- 599221252 Apr 4 2022 08:58:48 +00:00 cat9k-rpbase.17.03.04.SPA.pkg\n'
                                                                 '327700 -rw- 47407779 Apr 4 2022 08:59:46 +00:00 cat9k-rpboot.17.03.04.SPA.pkg\n'
                                                                 '327695 -rw- 34874376 Apr 4 2022 08:58:48 +00:00 cat9k-sipbase.17.03.04.SPA.pkg\n'
                                                                 '327696 -rw- 57082884 Apr 4 2022 08:58:48 +00:00 cat9k-sipspa.17.03.04.SPA.pkg\n'
                                                                 '327697 -rw- 29717512 Apr 4 2022 08:58:48 +00:00 cat9k-srdriver.17.03.04.SPA.pkg\n'
                                                                 '327698 -rw- 14431236 Apr 4 2022 08:58:48 +00:00 cat9k-webui.17.03.04.SPA.pkg\n'
                                                                 '327699 -rw- 9220 Apr 4 2022 08:58:48 +00:00 cat9k-wlc.17.03.04.SPA.pkg\n'
                                                                 '11353194496 bytes total (6882414592 bytes free)\n'))
            li2 = step5[0].select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(5) > td:nth-child(2) > ol > li:nth-child(2)')
            document.add_paragraph('b.   '+li2[0].find('span', class_="ph synph").text)
            document.add_paragraph('The following is sample output of the dir flash:*.conf command. It displays the .conf files in the flash partition; note the two .conf files:\n'
                                   'a.    packages.conf—the file that has been re-written with the newly installed .pkg files.  \n' 
                                   'b.    cat9k_iosxe.17.03.04.SPA.conf— a backup copy of the newly installed packages.conf file.')
            document.add_paragraph(UpgradeInInstallMode.funTable('Switch#dir flash:*.conf\n'
                                                                 'Directory of flash:/*.conf\n'
                                                                 '\n'
                                                                 'Directory of flash:/\n'
                                                                 '\n'
                                                                 '155689 -rw- 7715 Apr 4 2022 09:04:35 +00:00 packages.conf\n'
                                                                 '327689 -rw- 7715 Apr 4 2022 08:59:46 +00:00 cat9k_iosxe.17.03.04.SPA.conf\n'
                                                                 '11353194496 bytes total (6882414592 bytes free)\n'))
            step6 = soup.select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(6) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 6 ').bold = True
            p.add_run(step6[0].find('p', class_="ph cmd").text)
            document.add_paragraph(step6[0].find('span',class_='synph').text)
            document.add_paragraph(step6[0].find('p', class_="p").text.replace("17.3.1","17.3.4").replace('\n',''))
            document.add_paragraph('In case of a Cisco StackWise Virtual setup, remember to upgrade the active and standby')
            points = soup.select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(6) > td:nth-child(2) > ul')
            l = points[0].text.splitlines()
            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele)
            document.add_paragraph(t1[0],style='List Bullet 2')
            document.add_paragraph(t1[1],style='List Bullet 2')
            document.add_paragraph(UpgradeInInstallMode.funTable(step6[0].find('pre',class_="pre codeblock").text))

            step7 = soup.select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(7) > td:nth-child(2)')
            p = document.add_paragraph("\n")
            p.add_run('Step 7 ').bold = True
            p.add_run(step7[0].find('p', class_="ph cmd").text)
            li1 = step7[0].select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(7) > td:nth-child(2) > ol > li:nth-child(1)')#task_xpb_3tx_xkb > div > table > tbody > tr:nth-child(7) > td:nth-child(2) > ol > li:nth-child(1)
            document.add_paragraph('a.   '+li1[0].find('span', class_="ph synph").text)
            document.add_paragraph('Use this command to reload the switch. When the switch reloads after a ROMMON upgrade, the ROMMON version is updated, but not displayed in the output until the next reload.')
            document.add_paragraph(UpgradeInInstallMode.funTable(li1[0].find('pre',class_="pre codeblock").text))
            li2 = step7[0].select('#task_p3r_f21_jmb > div > table > tbody > tr:nth-child(7) > td:nth-child(2) > ol > li:nth-child(2)')
            document.add_paragraph('b.   '+li2[0].find('span', class_="ph synph").text)
            document.add_paragraph(step7[0].find('p', class_="p").text)
            document.add_paragraph('The following sample output of the show version command displays the Cisco IOS XE Amsterdam 17.3.4 image on the device:')
            document.add_paragraph(UpgradeInInstallMode.funTable("Switch# show version\n"
                                                                 "Cisco IOS XE Software, Version 17.03.04\n"
                                                                 "Cisco IOS Software [Amsterdam], Catalyst L3 Switch Software (CAT9K_IOSXE), Version 17.3.4, RELEASE SOFTWARE (fc2)\n"
                                                                 "Technical Support: http://www.cisco.com/techsupport\n"
                                                                 "Copyright (c) 1986-2022 by Cisco Systems, Inc..\n"
                                                                 "<output truncated>\n"))
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some data is missing in 'Upgrading In Install Mode' section")
        document.add_page_break()
        return   
class OsUpgrade(Formatter):
    def __init__(self,os,version):
        self.os =os
        self.version = version
    def osUpgrade_ASR_1000_17_03_03(self):
        p = document.add_paragraph("\n")
        p.add_run('Step 1').bold = True
        p.add_run(" Copy new image to flash\nCopy the system image to the master switch flash memory with a transfer protocol. You can use ftp:,\ntftp:, scp:, or sftp:. The examples in this procedure uses TFTP:").italic = False 
        OsUpgrade.funTable("Router#copy tftp: bootflash:\n"
                           "Address or name of remote host [10.76.78.236]?\n"
                           "Source filename [asr1002x-universalk9.17.03.02.SPA.bin]? asr1002x-universalk9.17.03.03.SPA.bin\n"
                           "Destination filename [asr1002x-universalk9.17.03.03.SPA.bin]?\n"
                           "Accessing tftp://10.76.78.236/asr1002x-universalk9.17.03.03.SPA.bin...\n"
                           "Loading asr1002x-universalk9.17.03.03.SPA.bin from 10.76.78.236 (via GigabitEthernet0): !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!OO!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n"
                           "[OK - 859647475 bytes]")
        document.add_paragraph("Use dir flash: to confirm that the image has been successfully copied to flash.") 
        OsUpgrade.funTable("Router#dir flash:*.bin\n"
                    "Directory of bootflash:/*.bin\n"
                    "Directory of bootflash:/\n"
                    "   16  -rw-   678623418   Jul 8 2021 07:03:16 +00:00  asr1002x-universalk9.17.05.01a.SPA.bin\n"
                    "   18  -rw-   716879226  Nov 28 2019 01:00:59 +00:00  asr1002x-universalk9.16.09.03.SPA.bin\n"
                    "   19  -rw-   680899549  Nov 28 2019 01:20:00 +00:00  asr1002x-universalk9.16.07.03.SPA.bin\n"
                    "   23  -rw-   716473272  May 11 2021 06:43:02 +00:00  asr1002x-universalk9.16.09.02.SPA.bin\n"
                    "   29  -rw-   463878176  Dec 14 2018 19:06:20 +00:00  asr1002x-universalk9.03.16.05.S.155-3.S5-ext.SPA.bin\n"
                    "48583  -rw-   669180599  Dec 19 2018 01:34:59 +00:00  asr1002x-universalk9.16.06.05.SPA.bin\n"
                    "   13  -rw-   859647475  Dec 17 2021 01:19:29 +00:00  asr1002x-universalk9.17.03.03.SPA.bin\n"
                    "6646632448 bytes total (495079424 bytes free)")          
        p = document.add_paragraph("\n")
        p.add_run('Step 2').bold = True
        p.add_run(" Set boot variable").italic = False
        OsUpgrade.funTable('Router(config)# boot system flash bootflash:asr1002x-universalk9.17.03.03.SPA.bin')
        r = document.add_paragraph('\n').add_run()
        try:
            r.add_picture('note.jpg',width=Inches(0.3), height=Inches(.3))
            r.add_text('We can add once again the boot statement for the old version in order to have a backup version for the device to use in case of it not able to boot with the new one.')
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found...")
        OsUpgrade.funTable("Router(config)# boot system flash bootflash:asr1002x-universalk9.16.09.03.SPA.bin")
        p = document.add_paragraph('\n')
        p.add_run('Step 3').bold = True
        p.add_run(" write memory\nUse this command to save boot settings.").italic = False 
        OsUpgrade.funTable("Router#write memory\n"
                           "Building configuration...\n"
                           "[OK]")
        p = document.add_paragraph("\n")
        p.add_run('Step 4').bold = True
        p.add_run(" show bootvar or show boot\nUse this command to verify the boot variable (packages.conf) and manual boot setting (no):").italic = False 
        OsUpgrade.funTable("Router#sh bootvar\n"
                           "BOOT variable = bootflash:asr1002x-universalk9.17.03.03.SPA.bin,1;bootflash:asr1002x-universalk9.16.09.03.SPA.bin,1;\n"
                           "CONFIG_FILE variable does not exist\n"
                           "BOOTLDR variable does not exist\n"
                           "Configuration register is 0x2102\n"
                           "Standby not ready to show bootvar")
        p = document.add_paragraph("\n")
        p.add_run('Step 5').bold = True
        p.add_run(" reload").italic = False
        OsUpgrade.funTable("Router#reload\n"
                           "Proceed with reload? [confirm]\n"
                           "Dec 17 01:28:04.333: %PMAN-5-E\n\n"
                           "Initializing Hardware ...\n"
                           "System integrity status: 00000610\n\n"
                           "System Bootstrap, Version 17.3(1r), RELEASE SOFTWARE\n"
                           "Copyright (c) 1994-2020 by cisco Systems, Inc.\n"
                           "Current image running: Boot ROM1\n"
                           "Last reset cause: LocalSoft\n\n"
                           "SPI Flash 4KB Descriptor Area Checksum = 0x6f5a0db2\n"
                           "ASR1002-X platform with 8388608 Kbytes of main memory\n"
                           "File size is 0x333d2df3\n"
                           "Located asr1002x-universalk9.17.03.03.SPA.bin\n"
                           "Image size 859647475 inode num 13, bks cnt 209875 blk size 8*512\n"
                           "###################################################################################################################\n"
                           "<snip>\n"
                           "Boot image size = 859647475 (0x333d2df3) bytes\n"
                           "ROM:RSA Self Test Passed\n"
                           "ROM:Sha512 Self Test Passed\n"
                           "Package header rev 1 structure detected\n"
                           "Validating main package signatures\n"
                           "RSA Signed RELEASE Image Signature Verification Successful.\n"
                           "Image validated\n"
                           "Dec 17 01:31:36.621: %BOOT-5-OPMODE_LOG: R0/0: binos: System booted in CONTROLLER mode\n"
                           "              Restricted Rights Legend\n"
                           "Use, duplication, or disclosure by the Government is\n"
                           "subject to restrictions as set forth in subparagraph\n"
                           "(c) of the Commercial Computer Software - Restricted\n"
                           "Rights clause at FAR sec. 52.227-19 and subparagraph\n"
                           "(c) (1) (ii) of the Rights in Technical Data and Computer\n"
                           "Software clause at DFARS sec. 252.227-7013.\n\n"
                           "           Cisco Systems, Inc.\n"
                           "           170 West Tasman Drive\n"
                           "           San Jose, California 95134-1706\n\n\n"
                           "Cisco IOS Software [Amsterdam], ASR1000 Software (X86_64_LINUX_IOSD-UNIVERSALK9-M), Version 17.3.3, RELEASE SOFTWARE (fc7)\n"
                           "Technical Support: http://www.cisco.com/techsupport\n"
                           "Copyright (c) 1986-2021 by Cisco Systems, Inc.\n"
                           "Compiled Thu 04-Mar-21 12:36 by mcpre\n\n\n"
                           "PLEASE READ THE FOLLOWING TERMS CAREFULLY. INSTALLING THE LICENSE OR\n"
                           "LICENSE KEY PROVIDED FOR ANY CISCO SOFTWARE PRODUCT, PRODUCT FEATURE,\n"
                           "AND/OR SUBSEQUENTLY PROVIDED SOFTWARE FEATURES (COLLECTIVELY, THE SOFTWARE"
                           "), AND/OR USING SUCH SOFTWARE CONSTITUTES YOUR FULL\n"
                           "ACCEPTANCE OF THE FOLLOWING TERMS. YOU MUST NOT PROCEED FURTHER IF YOU\n"
                           "ARE NOT WILLING TO BE BOUND BY ALL THE TERMS SET FORTH HEREIN.\n\n"
                           "Your use of the Software is subject to the Cisco End User License Agreement\n"
                           "(EULA) and any relevant supplemental terms (SEULA) found at\n"
                           "http://www.cisco.com/c/en/us/about/legal/cloud-and-software/software-terms.html.\n\n"
                           "You hereby acknowledge and agree that certain Software and/or features are\n"
                           "licensed for a particular term, that the license to such Software and/or\n"
                           "features is valid only for the applicable term and that such Software and/or\n"
                           "features may be shut down or otherwise terminated by Cisco after expiration\n"
                           "of the applicable license term (e.g., 90-day trial period). Cisco reserves\n"
                           "the right to terminate any such Software feature electronically or by any\n"
                           "other means available. While Cisco may provide alerts, it is your sole\n"
                           "responsibility to monitor your usage of any such term Software feature to\n"
                           "ensure that your systems and networks are prepared for a shutdown of the\n"
                           "Software feature.\n\n\n\n"
                           "All TCP AO KDF Tests Pass\n"
                           "cisco ASR1002-X (2RU-X) processor (revision 2KP) with 3756810K/6147K bytes of memory\n."
                           "Processor board ID FOX2114P78M\n"
                           "Router operating mode: Controller-Managed\n"
                           "6 Gigabit Ethernet interfaces\n"
                           "32768K bytes of non-volatile configuration memory.\n"
                           "8388608K bytes of physical memory.\n"
                           "6594559K bytes of eUSB flash at bootflash:.\n\n\n"
                           "Press RETURN to get started!")
    def osUpgrade_Switch_Catalyst_9600_17_06_02(self):
        return


class RollbackProcedure(Formatter):
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def rollbackProcedure_ASR_1000_17_03_03(self):
        document.add_paragraph('Please follow the outline below if rollback is needed.')
        document.add_paragraph('Remove old boot statement with no boot system',style='List Number 2')
        document.add_paragraph('Update boot statement to reflect original IOS-XE release. ',style='List Number 2')
        document.add_paragraph('Show run to confirm boot statement',style='List Number 2')
        document.add_paragraph('Copy run start',style='List Number 2')
        document.add_paragraph('Verify boot variables with ‘show bootvar’',style='List Number 2')
        document.add_paragraph('Reload the router ',style='List Number 2')
        document.add_page_break()
        return
    def rollbackProcedure_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph('In case of failure, malfunctioning or any kind of network issue after the software upgrade that can not be troubleshooted or remediated due to time constraints or difficulty of doing so, the network engineer must proceed with a software downgrade of the device.')

        rollback = (
            ('Task Number', '1'),
            ('Task Description ', 'Rollback to initial state'),
            ('Task Scheduled Start', ''),
            ('Task Duration','45 minutes'),
            ('Task Owner Pri/Sec',''),
            ('Task Dependencies','Device is operational OOB/Console access to the device')
        )
        table = document.add_table(rows=1, cols=2, style='Colorful Grid Accent 1')
        for term, definition in rollback:
            row_cells = table.add_row().cells
            row_cells[0].text = str(term)
            row_cells[1].text = str(definition)
        try:
            r = document.add_paragraph().add_run()
            r.add_picture('warning.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('For most software releases, direct downgrade will not be possible. One needs to follow the same path that was used for the upgrade in an opposite order. Please verify versions defined in Tasks 2 and 4 in the above')
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found...")
        p = document.add_paragraph()
        p.add_run("Task Detail ").bold = True
        p = document.add_paragraph()

        p.add_run("Step 1 ").bold = True
        p.add_run('Gain OOB access to the device:')
        p = document.add_paragraph()

        p.add_run("Step 2 ").bold = True
        p.add_run('Verify no concurrent configuration sessions are ongoing')
        RollbackProcedure.funTable('# show configuration session summary\n'
                 'There are no active configuration sessions')
        p = document.add_paragraph()

        p.add_run("Step 3 ").bold = True
        p.add_run('Initiate the software downgrade to interim (or previous version, in case rollback occurs before final upgrade)')
        RollbackProcedure.funTable('# install all kickstart n5000-uk9-kickstart.7.1.5.N1.1.bin system n5000-uk9.7.1.5.N1.1.bin')
        p = document.add_paragraph()

        p.add_run("Step 4 ").bold = True
        p.add_run('Confirm to proceed and wait for the procedure to finish')
        p = document.add_paragraph()

        p.add_run("Step 5 ").bold = True
        p.add_run('Verify the process has finished by issuing the commands')
        RollbackProcedure.funTable('# show install all status\n'
                 'This is the log of last installation.\n'
                 '.\n'
                 'Install has been successful.\n'
                 '# show boot')
        p = document.add_paragraph()
        p.add_run("Step 6 ").bold = True
        p.add_run(' Repeat steps 1-5 with the original release for the particular device as stated in site specific addendum to this document')
        try:
            r = document.add_paragraph().add_run()
            r.add_picture('note.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('Note: if downgrading to system image n5000-uk9.5.0.3.N1.1c.bin the following additional steps will be required before downgrading:')
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found...")
        RollbackProcedure.funTable('show incompatibility system n5000-uk9.5.0.3.N1.1c.bin\n'
                 'config terminal\n'
                 'no snmp-server enable traps bridge\m'
                 'no snmp-server enable traps stpx\n'
                 'copy running-config startup-config\n'
                 'end\n'
                 'show incompatibility system n5000-uk9.5.0.3.N1.1c.bin')
        try:
            r = document.add_paragraph().add_run()
            r.add_picture('note.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('Expected output from command in last step - No incompatible configurations')
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found...")
        p = document.add_paragraph()

        p.add_run("Step 7 ").bold = True
        p.add_run('Verify status of the rollback patch using the following command:')
        RollbackProcedure.funTable('# show diff rollback-patch running-config file\n'
                 'bootflash:show_pre_7.3_config')
        p = document.add_paragraph()

        p.add_run("Step 8 ").bold = True
        p.add_run('Execute rollback using the following command:')
        RollbackProcedure.funTable('# rollback running-config file bootflash::<checkpoint_file> verbose.')

        p = document.add_paragraph()

        p.add_run("Step 9 ").bold = True
        p.add_run('Verify rollback status')
        RollbackProcedure.funTable('# show rollback log verify')
        document.add_paragraph('(status should be success)')
        try:
            r = document.add_paragraph().add_run()
            r.add_picture('note.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('Expected warning (no further action required):')
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found...")
        RollbackProcedure.funTable('no lacp suspend-individual config change may not be applied to port-channel in up state')
        p = document.add_paragraph()

        p.add_run("Step 10 ").bold = True
        p.add_run('After moving to the initial version execute the following')
        RollbackProcedure.funTable('show diff rollback-patch running-config file\n'
                 'bootflash:show_pre_7.3_config\n'
                 'rollback running-config file bootflash:<checkpoint_file> verbos\n'
                 'show rollback log verify')
        document.add_paragraph('(status should be success)')
        p = document.add_paragraph()
        p.add_run("Task Success Criteria").bold = True
        document.add_paragraph('\the two needed VLANs are successfully deleted')
        p = document.add_paragraph()
        p.add_run("Failure Procedure").bold = True
        document.add_paragraph('\t1. Resolve the problem, if one discovered, and if possible.\n'
                               '\t2. Call cisco TAC for support\n'
                               '\t3. Proceed to Rollback procedure')
        try:
            r = document.add_paragraph().add_run()
            r.add_picture('caution.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text('FEX serial command was replaced with chassis-serial command When rolling back to 5.0(3)N2(2) from 7.1.(4).N1.(1). The FEX will not move into the online state with this configuration. Removing the chassis-serial command and reapplying the serial command will resolve the issue and the FEX will come online')
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found...")
        document.add_page_break()
        return

class DiffReport:
    def __init__(self,os,version) -> None:
        self.os=os
        self.version=version
    def diffReport_Nexus_5548_7_3_6_N1_1_(self):
        document.add_paragraph("Recollect pre-upgrade check commands as described in Table 5 Pre-check show commands list and create a “Diff Report” based of comparison between the two checks.")
        try:
            r = document.add_paragraph().add_run()
            r.add_picture('note.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text("Diff report severity will be based on the differences found in comparison between "
                       "Pre & Post.\n"
                       "Table 5 has a column named “Exact Match”. If a difference will be found on a query "
                       "marked with “X” the report will consider the Diff as Major issue.\n"
                       "If a difference will be found on other queries the report will considered the Diff as "
                       "Minor issue.")
            r = document.add_paragraph("\n").add_run()
            r.add_picture('note.jpg', width=Inches(0.3), height=Inches(.3))
            r.add_text("The Diff Report will also show any difference in configuration before and after the MW")
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found...")
        return
class TimeTaken:
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def timeTaken_ASR_1000_17_03_03(self):
        data = (
        ("Update boot variable and Rommon upgrade", "10 minutes"),
        ("FPGA upgrade for ASR1000-RP2","15 minutes"),
        ("(Post boot checks)","5 minutes"),
        ("Minimum downtime (for upgrade procedure)","25 minutes"),
        ("Rollback procedure",""),
        ("Update boot variable and Rommon upgrade","10 minutes"),
        ("(Post boot checks)","5 minutes"),
        ("Minimum downtime (during rollback procedure)","10 minutes")
        )
        table = document.add_table(rows=1, cols=2, style = 'Colorful Grid Accent 1')
        row = table.rows[0].cells
        row[0].text = "Section"
        row[1].text = 'Time taken (in minutes)'
        for section,time in data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(section)
            row_cells[1].text = str(time)
        document.add_page_break()
        return

#Class for ISSU creation
class IssuUpgrade(Formatter):
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def issuUpgrade_ASR_1000_17_03_03(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/routers/asr1000/configuration/guide/chassis/asr1000-software-config-guide/issu-asr.html?referring_site=RE&pos=2&page=https://www.cisco.com/c/en/us/support/routers/asr-1000-series-aggregation-services-routers/products-installation-guides-list.html')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#con_1081725 > section')
            data = elem[0].text.strip()
            #d = re.split('Router#',data)
            d = "Router#"
            s =  [e for e in data.split(d) if e]
            document.add_paragraph(s[0])
            for j in s[1:]:
                IssuUpgrade.funTable('Router#'+j)
                document.add_paragraph('\n')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some Data is missing in 'APPENDIX A: IN SERVICE SOFTWARE UPGRADE (ISSU)' section")

        document.add_page_break()
        return 
    def issuUpgrade_Switch_Catalyst_9600_17_06_02(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-6/release_notes/ol-17-6-9600.html#concept_ycv_jdf_3mb')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            elem = soup.select('#task_b1z_hwr_43b__d54e4569')
            elem1 = soup.select('#task_b1z_hwr_43b__d54e4474 > div')
            elem2 = soup.select('#task_b1z_hwr_43b__d54e4474 > p')
            table1 = soup.select('#task_b1z_hwr_43b__d54e4474  > table')
            #table2 =soup.select('#task_b1z_hwr_43b > div > table')
            l = table1[0].text.splitlines()
            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele)
            #l2 = table2[0].text.splitlines()
            t2 = []
            for ele in l:
                if ele.strip():
                    t2.append(ele)
            document.add_paragraph(re.sub(' +', ' ', elem[0].text))
            p = document.add_paragraph()
            p.add_run(elem1[0].text).bold =True
            document.add_paragraph(elem2[0].text)
            table = document.add_table(rows=1, cols=3, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = t1[0]
            row[1].text = t1[1]
            row[2].text = t1[2]
            row = table.add_row().cells
            row[0].text = t1[3]
            row[1].text = t1[4]+' '+t1[5].strip()
            row[2].text = t1[6]
            row = table.add_row().cells
            row[0].text = t1[7]
            row[1].text = t1[8]
            row[2].text = t1[9]

            step1 = soup.select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(1) > td:nth-child(2)')
            document.add_paragraph("Step 1 "+step1[0].find('p', class_="ph cmd").text)
            #document.add_paragraph(step1[0].find('span', class_="ph synph").text)
            document.add_paragraph('Enables privileged EXEC mode. Enter your password if prompted.')
            document.add_paragraph(IssuUpgrade.funTable(step1[0].find('pre',class_="pre codeblock").text))

            step2 = soup.select("#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(2) > td:nth-child(2)")
            document.add_paragraph("Step 2 "+step2[0].find('p', class_="ph cmd").text.strip())
            div1 = step2[0].select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(2) > td:nth-child(2) > div:nth-child(2)')
            document.add_paragraph(" ".join(div1[0].find('p', class_="p").text.split()))
            document.add_paragraph(IssuUpgrade.funTable(div1[0].find('pre',class_="pre codeblock").text))

            div2 = step2[0].select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(2) > td:nth-child(2) > div:nth-child(3)')
            document.add_paragraph("The following sample output displays installation of Cisco IOS XE Amsterdam 17.06.01 software image with ISSU procedure.")
            document.add_paragraph(IssuUpgrade.funTable(div2[0].find('pre',class_="pre codeblock").text))

            step3 = soup.select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(3) > td:nth-child(2)')
            document.add_paragraph("Step 3 "+step3[0].find('p', class_="ph cmd").text)
            div1 = step3[0].select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > div:nth-child(2)')
            document.add_paragraph(div1[0].find('p', class_="p").text)

            div2 = step3[0].select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > div:nth-child(3)')
            document.add_paragraph("The following sample output of the show version command displays the Cisco IOS XE Amsterdam 17.06.01 image on the device:")
            document.add_paragraph(IssuUpgrade.funTable(div2[0].find('pre',class_="pre codeblock").text))



            step4 = soup.select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(4) > td:nth-child(2)')
            document.add_paragraph("Step 4 "+step4[0].find('p', class_="ph cmd").text)
            div1 = step4[0].select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(4) > td:nth-child(2) > div')
            document.add_paragraph(div1[0].find('p', class_="p").text)
            document.add_paragraph(IssuUpgrade.funTable(div1[0].find('pre',class_="pre codeblock").text))

            step5 = soup.select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(5) > td:nth-child(2)')
            document.add_paragraph("Step 5 "+step5[0].find('p', class_="ph cmd").text)
            document.add_paragraph(step5[0].find('p', class_="p").text)
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some Data is missing in 'APPENDIX A: IN SERVICE SOFTWARE UPGRADE (ISSU)' section")
        document.add_page_break()
        return
    def issuUpgrade_Switch_Catalyst_9600_17_03_05(self):
        res = requests.get('https://www.cisco.com/c/en/us/td/docs/switches/lan/catalyst9600/software/release/17-3/release_notes/ol-17-3-9600.html')
        soup = BeautifulSoup(res.text, 'lxml')
        try:
            
            elem = soup.select('#task_b1z_hwr_43b__d54e5028')
            elem1 = soup.select('#task_b1z_hwr_43b__d54e4933 > div')
            elem2 = soup.select('#task_b1z_hwr_43b__d54e4933 > p')
            table1 = soup.select('#task_b1z_hwr_43b__d54e4933 > table')
            #table2 =soup.select('#task_b1z_hwr_43b > div > table')
            l = table1[0].text.splitlines()
            t1 = []
            for ele in l:
                if ele.strip():
                    t1.append(ele)
            #l2 = table2[0].text.splitlines()
            t2 = []
            for ele in l:
                if ele.strip():
                    t2.append(ele)
            document.add_paragraph(re.sub(' +', ' ', elem[0].text))
            p = document.add_paragraph()
            p.add_run(elem1[0].text).bold =True
            document.add_paragraph(elem2[0].text)
            table = document.add_table(rows=1, cols=3, style = 'Colorful Grid Accent 1')
            row = table.rows[0].cells
            row[0].text = t1[0]
            row[1].text = t1[1]
            row[2].text = t1[2]
            row = table.add_row().cells
            row[0].text = t1[3]
            row[1].text = t1[4]+' '+t1[5].strip()
            row[2].text = t1[6]
            row = table.add_row().cells
            row[0].text = t1[7]
            row[1].text = t1[8]
            row[2].text = t1[9]

            step1 = soup.select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(1) > td:nth-child(2)')
            p = document.add_paragraph()
            p.add_run("Step 1 ").bold =True
            p.add_run(step1[0].find('p', class_="ph cmd").text)
            document.add_paragraph('Enables privileged EXEC mode. Enter your password if prompted.')
            document.add_paragraph(IssuUpgrade.funTable(step1[0].find('pre',class_="pre codeblock").text))

            step2 = soup.select("#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(2) > td:nth-child(2)")
            p = document.add_paragraph()
            p.add_run("Step 2 ").bold =True
            p.add_run(step2[0].find('p', class_="ph cmd").text.strip())
            div1 = step2[0].select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(2) > td:nth-child(2) > div:nth-child(2)')
            document.add_paragraph(" ".join(div1[0].find('p', class_="p").text.split()))
            document.add_paragraph(IssuUpgrade.funTable(div1[0].find('pre',class_="pre codeblock").text))

            div2 = step2[0].select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(2) > td:nth-child(2) > div:nth-child(3)')
            document.add_paragraph("The following sample output displays installation of Cisco IOS XE Amsterdam 17.3.2a software image with ISSU procedure.")
            document.add_paragraph(IssuUpgrade.funTable(div2[0].find('pre',class_="pre codeblock").text))

            step3 = soup.select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(3) > td:nth-child(2)')
            p = document.add_paragraph()
            p.add_run("Step 3 ").bold =True
            p.add_run(step3[0].find('p', class_="ph cmd").text)
            div1 = step3[0].select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > div:nth-child(2)')
            document.add_paragraph(div1[0].find('p', class_="p").text)
            div2 = step3[0].select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(3) > td:nth-child(2) > div:nth-child(3)')
            document.add_paragraph("The following sample output of the show version command displays the Cisco IOS XE Amsterdam 17.03.05 image on the device:")
            document.add_paragraph(IssuUpgrade.funTable(div2[0].find('pre',class_="pre codeblock").text))



            step4 = soup.select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(4) > td:nth-child(2)')
            p = document.add_paragraph()
            p.add_run("Step 4 ").bold =True
            p.add_run(step4[0].find('p', class_="ph cmd").text)
            div1 = step4[0].select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(4) > td:nth-child(2) > div')
            document.add_paragraph(div1[0].find('p', class_="p").text)
            document.add_paragraph(IssuUpgrade.funTable(div1[0].find('pre',class_="pre codeblock").text))

            step5 = soup.select('#task_b1z_hwr_43b > div > table > tbody > tr:nth-child(5) > td:nth-child(2)')
            p = document.add_paragraph()
            p.add_run("Step 5 ").bold =True
            p.add_run(step5[0].find('p', class_="ph cmd").text)
            document.add_paragraph(step5[0].find('p', class_="p").text)
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Some Data is missing in 'Issu Upgrade section")
        document.add_page_break()
        return
    


#Acronym Listing
class AcronymListing:
    def __init__(self,os,version):
        self.os = os
        self.version = version
    def acronymListing_ASR_1000_17_03_03(self):
        acronyms = (
                   ('ASR','Aggregation Services Routers'),
                   ('IOS','Internetwork Operating System'),
                   ('ROMMON','Read Only Memory Monitor')
                   )
        table = document.add_table(rows=1, cols=2,style='Colorful Grid Accent 1')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Term'
        hdr_cells[1].text = 'Definition'
        for term,definition in acronyms:
            row_cells = table.add_row().cells
            row_cells[0].text = str(term)
            row_cells[1].text = str(definition)
        #document.add_page_break()
        return
    def acronymListing_Switch_Catalyst_9600_17_06_02(self):
        acronyms = (
                   ('IOS','Internetwork Operating System'),
                   ('ROMMON','Read Only Memory Monitor'),
                   ('C9600','Catalyst Switch model 9600')
                   )
        table = document.add_table(rows=1, cols=2,style='Colorful Grid Accent 1')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Term'
        hdr_cells[1].text = 'Definition'
        for term,definition in acronyms:
            row_cells = table.add_row().cells
            row_cells[0].text = str(term)
            row_cells[1].text = str(definition)
        #document.add_page_break()
        return
    def acronymListing_Switch_Catalyst_9600_17_03_05(self):
        acronyms = (
                   ('IOS','Internetwork Operating System'),
                   ('ROMMON','Read Only Memory Monitor'),
                   ('C9600','Catalyst Switch model 9600')
                   )
        table = document.add_table(rows=1, cols=2,style='Colorful Grid Accent 1')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Term'
        hdr_cells[1].text = 'Definition'
        for term,definition in acronyms:
            row_cells = table.add_row().cells
            row_cells[0].text = str(term)
            row_cells[1].text = str(definition)
        #document.add_page_break()
        return
    def acronymListing_Switch_Catalyst_9300_17_03_04(self):
        acronyms = (
                   ('IOS','Internetwork Operating System'),
                   ('ROMMON','Read Only Memory Monitor'),
                   ('C9300','Catalyst Switch model 9300')
                   )
        table = document.add_table(rows=1, cols=2,style='Colorful Grid Accent 1')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Term'
        hdr_cells[1].text = 'Definition'
        for term,definition in acronyms:
            row_cells = table.add_row().cells
            row_cells[0].text = str(term)
            row_cells[1].text = str(definition)
        #document.add_page_break()
        return
    def acronymListing_Switch_Catalyst_9500_17_03_04(self):
        acronyms = (
                   ('IOS','Internetwork Operating System'),
                   ('ROMMON','Read Only Memory Monitor'),
                   ('C9500','Catalyst Switch model 9500')
                   )
        table = document.add_table(rows=1, cols=2,style='Colorful Grid Accent 1')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Term'
        hdr_cells[1].text = 'Definition'
        for term,definition in acronyms:
            row_cells = table.add_row().cells
            row_cells[0].text = str(term)
            row_cells[1].text = str(definition)
        #document.add_page_break()
        return
class TrademarksandDisclaimer:
    def __init__(self):
        pass
    @classmethod
    def trademarksandDisclaimer(self):
        document.add_paragraph("THE SPECIFICATIONS AND INFORMATION REGARDING THE PRODUCTS IN THIS MANUAL ARE SUBJECT TO CHANGE WITHOUT NOTICE. ALL STATEMENTS, INFORMATION, AND RECOMMENDATIONS IN THIS MANUAL ARE BELIEVED TO BE ACCURATE BUT ARE PRESENTED WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED. USERS MUST TAKE FULL RESPONSIBILITY FOR THEIR APPLICATION OF ANY PRODUCTS."
        "\nTHE SOFTWARE LICENSE AND LIMITED WARRANTY FOR THE ACCOMPANYING PRODUCT ARE SET FORTH IN THE INFORMATION PACKET THAT SHIPPED WITH THE PRODUCT AND ARE INCORPORATED HEREIN BY THIS REFERENCE. IF YOU ARE UNABLE TO LOCATE THE SOFTWARE LICENSE OR LIMITED WARRANTY, CONTACT YOUR CISCO REPRESENTATIVE FOR A COPY.\n"
        "\nNOTWITHSTANDING ANY OTHER WARRANTY HEREIN, ALL DOCUMENT FILES AND SOFTWARE OF THIRD PARTY SUPPLIERS ARE PROVIDED “AS IS” WITH ALL FAULTS. CISCO AND THIRD PARTY SUPPLIERS DISCLAIM ALL WARRANTIES, EXPRESSED OR IMPLIED, INCLUDING, WITHOUT LIMITATION, THOSE OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NON-INFRINGEMENT OR ARISING FROM A COURSE OF DEALING, USAGE, OR TRADE PRACTICE.\n"
        "\nIN NO EVENT SHALL CISCO OR ITS SUPPLIERS BE LIABLE FOR ANY INDIRECT, SPECIAL, CONSEQUENTIAL, OR INCIDENTAL DAMAGES, INCLUDING, WITHOUT LIMITATION, LOST PROFITS OR LOSS OR DAMAGE TO DATA ARISING OUT OF THE USE OR INABILITY TO USE THIS MANUAL, EVEN IF CISCO OR ITS SUPPLIERS HAVE BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.\n"
        "\nCisco and the Cisco Logo are trademarks of Cisco Systems, Inc. and/or its affiliates in the U.S. and other countries. A listing of Cisco's trademarks can be found at www.cisco.com/go/trademarks. Third-party trademarks mentioned are the property of their respective owners. The use of the word partner does not imply a partnership relationship between Cisco and any other company.\n"
        "\nAny Internet Protocol (IP) addresses and phone numbers used in this document are not intended to be actual addresses and phone numbers. Any examples, command display output, network topology diagrams, and other figures included in the document are shown for illustrative purposes only. Any use of actual IP addresses or phone numbers in illustrative content is unintentional and coincidental\n"
        "\n©2022 Cisco Systems, Inc. All rights reserved")
        document.add_page_break()
        return

class DocumentAcceptance:
    def __init__(self):
        pass
    def documentAcceptance(self):
        for i in range(0,3):
            table = document.add_table(rows=0, cols=4)
            row=table.add_row().cells
            c = ['Name', 'Title','Company','Signature','Date']
            for i in c:
                p = row[0].add_paragraph(i)
                p = row[1].add_paragraph('____________')
                p = row[2].add_paragraph(i)
                p = row[3].add_paragraph('____________')
        return

'''
This class is for making the document according to the platform and its version.
So, Please make sure that the Device name, version and heading should be given as same as the class function name.
'''
class Switcher(object):
    def __init__(self):
        pass
    @classmethod
    def str_to_class(self,classname):
        return getattr(sys.modules[__name__], classname)

    def indirect(self,dname,os, area, ver, rd,lr):
        platforms = allProducts['Platforms']
        classes = allProducts['class_list']
        specialChars = ' '
        method_name = area+'_'+str(dname).replace(specialChars,'_')+'_'+re.sub(r'[.()]','_',ver)
        #print(method_name)
        for i in classes:
            if re.search(area, str(i), re.IGNORECASE):
                if 'relatedDocuments' in method_name:
                    print('yes')
                    ob1 = Switcher.str_to_class(i)(os,ver,rd)
                    method = getattr(ob1,method_name,lambda :'Invalid')
                    return method()
                elif 'limitationsAndRestrictions' in method_name:
                    ob1 = Switcher.str_to_class(i)(os,ver,lr)
                    method = getattr(ob1,method_name,lambda :'Invalid')
                    return method()
                else:
                    ob1 = Switcher.str_to_class(i)(os,ver)
                    method = getattr(ob1,method_name,lambda :'Invalid')
                    return method()
                
            else:
                pass
        return

class Dynamic:
    def __init__(self):
        pass
    @classmethod
    def str_to_class(self,classname):
        return getattr(sys.modules[__name__], classname)
    
    def creation_Of_heading_with_content(self,cust,author,devicename,pid,os,version,typeofdevice,lis,rd,lr):
        c = ["Contents","About","History","Review","DocumentConventions","TrademarksandDisclaimer","DocumentAcceptance"]
        num = 0
        headingwithcclass = {
                        "ASR 1000 17.03.03":{
                        "Contents":"Contents",
                        "About This Mop Document":"About",
                        "History":"History",
                        "Review":"Review",
                        "Document Conventions":"DocumentConventions",
                        f"{num} Introduction":"",
                        "Preface":Preface.preface(devicename, pid,typeofdevice),
                        "Audience":Audience.audience(cust),
                        "Scope":Scope.scope(cust, devicename, version, os),
                        "Target Process Steps":"targetProcessSteps",
                        "Assumptions":Assumptions.assumptions(cust),
                        "Related Documents":"relatedDocuments",
                        f"{num} Preparation":"",
                        "Install Analysis":"installAnalysis",
                        f"Verify Current version of {os} & ROMMON":"iosVersion",
                        "Available DRAM on the device":"availableDram",
                        "Verify the Configuration Register and Free storage on Flash":"iosVersion",
                        "Impact of downtime":f"The impact on service availability due to an access {typeofdevice} upgrade may vary depending on network topology and operational factors. Typically, connected devices will experience a service outage of a few minutes while the {typeofdevice} reloads with new software. \n\nConsideration should also be given to what (if any) recovery steps may be required in the event of a device failing to complete an upgrade. This is a particular concern for devices in remote or unmanned locations, where it may require significant time and resources to physically replace equipment. In extreme cases it may simply represent too high an operational risk to upgrade a unit in service.\nAn assessment should be made for each targeted device, or groups of similar devices, and an upgrade strategy for these should be agreed with the customer.",
                        "Software image delivery":"It is anticipated that software images will be delivered to devices over the customer network. However, it may not be possible for all devices to access a single central software repository due to security constraints. As part of the install analysis process, an appropriate repository should be identified for each device.",
                        "Reporting":"The output from this process should be a list of targeted devices, each indicating ‘pass’ or ‘fail’, with any additional caveats and information (for example, identified software repository). \nWhere devices are identified as ‘failed’, these will need to be investigated and remediated. For example, additional space may need to be cleared on the device’s flash memory",
                        "Image Deployment and Validation":"imageDeploymentAndValidation",
                        "Backup Current Configuration - Executed before Installation/change to boot image":"backupCurrentConfiguration",
                        f"{num} Limitations And Restrictions":"limitationsAndRestrictions",
                        f"{num} Upgrade":"",
                        "Pre-install verification":"preInstallVerification",
                        "Install and Reload":"installAndReload",
                        "Upgrading the Device Software":"",
                        "ROMMON Upgrades":"rommonUpgrade",
                        f"Perform the {os} Upgrade":"osUpgrade",
                        "Post-install Verification":"preInstallVerification",
                        f"{num} Rollback Procedure":"rollbackProcedure",
                        f"{num} Time taken for Change activities (Optional)":"timeTaken",
                        f"{num} Appendix A: In Service Software Upgrade (ISSU)":"issuUpgrade",
                        f"{num} Appendix B: Acronym Listing or Full Glossary":"acronymListing",
                        "Trademarks and Disclaimer":"TrademarksandDisclaimer",
                        "Document Acceptance":"DocumentAcceptance",
                        },
                        "Switch Catalyst 9600 17.06.02":{
                        "Contents":"Contents",
                        "About This Mop Document":"About",
                        "History":"History",
                        "Review":"Review",
                        "Document Conventions":"DocumentConventions",
                        f"{num} Introduction":"",
                        "Preface":Preface.preface(devicename, pid,typeofdevice),
                        "Audience":Audience.audience(cust),
                        "Scope":Scope.scope(cust, devicename, version, os),
                        "Target Process Steps":"targetProcessSteps",
                        "Assumptions":Assumptions.assumptions(cust),
                        "Related Documents":"relatedDocuments",
                        f"{num} Preparation":"",
                        "Install Analysis":"installAnalysis",
                        f"Verify the mode of IOS version":"iosVersion",
                        "Available DRAM on the device":"availableDram",
                        "Verify the Configuration Register":"configurationRegister",
                        "Impact of downtime":f"The impact on service availability due to an access {typeofdevice} upgrade may vary depending on network topology and operational factors. Typically, connected devices will experience a service outage of a few minutes while the {typeofdevice} reloads with new software. \n\nConsideration should also be given to what (if any) recovery steps may be required in the event of a device failing to complete an upgrade. This is a particular concern for devices in remote or unmanned locations, where it may require significant time and resources to physically replace equipment. In extreme cases it may simply represent too high an operational risk to upgrade a unit in service.\nAn assessment should be made for each targeted device, or groups of similar devices, and an upgrade strategy for these should be agreed with the customer.",
                        "Software image delivery":"It is anticipated that software images will be delivered to devices over the customer network. However, it may not be possible for all devices to access a single central software repository due to security constraints. As part of the install analysis process, an appropriate repository should be identified for each device.",
                        "Reporting":"The output from this process should be a list of targeted devices, each indicating ‘pass’ or ‘fail’, with any additional caveats and information (for example, identified software repository). \nWhere devices are identified as ‘failed’, these will need to be investigated and remediated. For example, additional space may need to be cleared on the device’s flash memory",
                        "Image Deployment and Validation":"imageDeploymentAndValidation",
                        f"{num} Limitations And Restrictions":"limitationsAndRestrictions",
                        f"{num} Upgrade":"",
                        "Pre-install verification":"preInstallVerification",
                        "Install and Reload":"installAndReload",
                        "Upgrading the Switch Software - 9200/9300/9400/9500":"upgradeSwitch",
                        "ROMMON Upgrades":"rommonUpgrade",
                        "Upgrading in Install Mode":"upgradeInInstallMode",
                        f"{num} Appendix A: In Service Software Upgrade (ISSU)":"issuUpgrade",
                        f"{num} Appendix B: Acronym Listing or Full Glossary":"acronymListing",
                        "Trademarks and Disclaimer":"TrademarksandDisclaimer",
                        "Document Acceptance":"DocumentAcceptance",
                        },
                        "Switch Catalyst 9600 17.03.05":{
                        "Contents":"Contents",
                        "About This Mop Document":"About",
                        "History":"History",
                        "Review":"Review",
                        "Document Conventions":"DocumentConventions",
                        f"{num} Introduction":"",
                        "Preface":Preface.preface(devicename, pid,typeofdevice),
                        "Audience":Audience.audience(cust),
                        "Scope":Scope.scope(cust, devicename, version, os),
                        "Target Process Steps":"targetProcessSteps",
                        "Assumptions":Assumptions.assumptions(cust),
                        "Related Documents":"relatedDocuments",
                        f"{num} Preparation":"",
                        "Install Analysis":"installAnalysis",
                        f"Verify the mode of IOS version":"iosVersion",
                        "Available DRAM on the device":"availableDram",
                        "Impact of downtime":f"The impact on service availability due to an access {typeofdevice} upgrade may vary depending on network topology and operational factors. Typically, connected devices will experience a service outage of a few minutes while the {typeofdevice} reloads with new software. \n\nConsideration should also be given to what (if any) recovery steps may be required in the event of a device failing to complete an upgrade. This is a particular concern for devices in remote or unmanned locations, where it may require significant time and resources to physically replace equipment. In extreme cases it may simply represent too high an operational risk to upgrade a unit in service.\nAn assessment should be made for each targeted device, or groups of similar devices, and an upgrade strategy for these should be agreed with the customer.",
                        "Software image delivery":"It is anticipated that software images will be delivered to devices over the customer network. However, it may not be possible for all devices to access a single central software repository due to security constraints. As part of the install analysis process, an appropriate repository should be identified for each device.",
                        "Reporting":"The output from this process should be a list of targeted devices, each indicating ‘pass’ or ‘fail’, with any additional caveats and information (for example, identified software repository). \nWhere devices are identified as ‘failed’, these will need to be investigated and remediated. For example, additional space may need to be cleared on the device’s flash memory",
                        "Image Deployment and Validation":"imageDeploymentAndValidation",
                        f"{num} Limitations And Restrictions":"limitationsAndRestrictions",
                        f"{num} Upgrade":"",
                        "Pre-install verification":"preInstallVerification",
                        "Install and Reload":"installAndReload",
                        "Upgrading the Switch Software":"upgradeSwitch",
                        "ROMMON Upgrades":"rommonUpgrade",
                        "Upgrading in Install Mode":"upgradeInInstallMode",
                        f"{num} Appendix A: In Service Software Upgrade (ISSU)":"issuUpgrade",
                        f"{num} Appendix B: Acronym Listing or Full Glossary":"acronymListing",
                        "Trademarks and Disclaimer":"TrademarksandDisclaimer",
                        "Document Acceptance":"DocumentAcceptance",
                        },
                        "Switch Catalyst 9300 17.03.04":{
                        "Contents":"Contents",
                        "About This Mop Document":"About",
                        "History":"History",
                        "Review":"Review",
                        "Document Conventions":"DocumentConventions",
                        f"{num} Introduction":"",
                        "Preface":Preface.preface(devicename, pid,typeofdevice),
                        "Audience":Audience.audience(cust),
                        "Scope":Scope.scope(cust, devicename, version, os),
                        "Target Process Steps":"targetProcessSteps",
                        "Assumptions":Assumptions.assumptions(cust),
                        "Related Documents":"relatedDocuments",
                        f"{num} Preparation":"",
                        "Install Analysis":"installAnalysis",
                        f"Verify the mode of IOS version":"iosVersion",
                        "Available DRAM on the device":"availableDram",
                        "Impact of downtime":f"The impact on service availability due to an access {typeofdevice} upgrade may vary depending on network topology and operational factors. Typically, connected devices will experience a service outage of a few minutes while the {typeofdevice} reloads with new software. \n\nConsideration should also be given to what (if any) recovery steps may be required in the event of a device failing to complete an upgrade. This is a particular concern for devices in remote or unmanned locations, where it may require significant time and resources to physically replace equipment. In extreme cases it may simply represent too high an operational risk to upgrade a unit in service.\nAn assessment should be made for each targeted device, or groups of similar devices, and an upgrade strategy for these should be agreed with the customer.",
                        "Software image delivery":"It is anticipated that software images will be delivered to devices over the customer network. However, it may not be possible for all devices to access a single central software repository due to security constraints. As part of the install analysis process, an appropriate repository should be identified for each device.",
                        "Reporting":"The output from this process should be a list of targeted devices, each indicating ‘pass’ or ‘fail’, with any additional caveats and information (for example, identified software repository). \nWhere devices are identified as ‘failed’, these will need to be investigated and remediated. For example, additional space may need to be cleared on the device’s flash memory",
                        "Image Deployment and Validation":"imageDeploymentAndValidation",
                        f"{num} Limitations And Restrictions":"limitationsAndRestrictions",
                        f"{num} Upgrade":"",
                        "Pre-install verification":"preInstallVerification",
                        "Install and Reload":"installAndReload",
                        "Upgrading the Switch Software":"upgradeSwitch",
                        "ROMMON Upgrades":"rommonUpgrade",
                        "Upgrading in Install Mode":"upgradeInInstallMode",
                        f"{num} Appendix A: Acronym Listing or Full Glossary":"acronymListing",
                        "Trademarks and Disclaimer":"TrademarksandDisclaimer",
                        "Document Acceptance":"DocumentAcceptance",
                        },
                        "Switch Catalyst 9500 17.03.04":{
                        "Contents":"Contents",
                        "About This Mop Document":"About",
                        "History":"History",
                        "Review":"Review",
                        "Document Conventions":"DocumentConventions",
                        f"{num} Introduction":"",
                        "Preface":Preface.preface(devicename, pid,typeofdevice),
                        "Audience":Audience.audience(cust),
                        "Scope":Scope.scope(cust, devicename, version, os),
                        "Target Process Steps":"targetProcessSteps",
                        "Assumptions":Assumptions.assumptions(cust),
                        "Related Documents":"relatedDocuments",
                        f"{num} Preparation":"",
                        "Install Analysis":"installAnalysis",
                        f"Verify the mode of IOS version":"iosVersion",
                        "Available DRAM on the device":"availableDram",
                        "Impact of downtime":f"The impact on service availability due to an access {typeofdevice} upgrade may vary depending on network topology and operational factors. Typically, connected devices will experience a service outage of a few minutes while the {typeofdevice} reloads with new software. \n\nConsideration should also be given to what (if any) recovery steps may be required in the event of a device failing to complete an upgrade. This is a particular concern for devices in remote or unmanned locations, where it may require significant time and resources to physically replace equipment. In extreme cases it may simply represent too high an operational risk to upgrade a unit in service.\nAn assessment should be made for each targeted device, or groups of similar devices, and an upgrade strategy for these should be agreed with the customer.",
                        "Software image delivery":"It is anticipated that software images will be delivered to devices over the customer network. However, it may not be possible for all devices to access a single central software repository due to security constraints. As part of the install analysis process, an appropriate repository should be identified for each device.",
                        "Reporting":"The output from this process should be a list of targeted devices, each indicating ‘pass’ or ‘fail’, with any additional caveats and information (for example, identified software repository). \nWhere devices are identified as ‘failed’, these will need to be investigated and remediated. For example, additional space may need to be cleared on the device’s flash memory",
                        "Image Deployment and Validation":"imageDeploymentAndValidation",
                        f"{num} Limitations And Restrictions":"limitationsAndRestrictions",
                        f"{num} Upgrade":"",
                        "Pre-install verification":"preInstallVerification",
                        "Install and Reload":"installAndReload",
                        "Upgrading the Switch Software":"upgradeSwitch",
                        "ROMMON Upgrades":"rommonUpgrade",
                        "Upgrading in Install Mode":"upgradeInInstallMode",
                        f"{num} Appendix A: Acronym Listing or Full Glossary":"acronymListing",
                        "Trademarks and Disclaimer":"TrademarksandDisclaimer",
                        "Document Acceptance":"DocumentAcceptance",
                        },
                        "Nexus 5548 7.3(6)N1(1)":{
                        "Contents":"Contents",
                        "About This Mop Document":"About",
                        "History":"History",
                        "Review":"Review",
                        "Document Conventions":"DocumentConventions",
                        f"{num} Introduction":"",
                        "Preface":Preface.preface(devicename, pid,typeofdevice),
                        "Audience":Audience.audience(cust),
                        "Scope":Scope.scope(cust, devicename, version, os),
                        "Assumptions":Assumptions.assumptions(cust),
                        "Related Documents":"relatedDocuments",
                        f"{num} Project Overview":"",
                        "High Level Project Overview":"highLevelProjectOverview",
                        "Risk Analysis":"riskAnalysis",
                        "Timescales":"timeScales",
                        "Resource requirements":"resourceRequirements",
                        "Responsibilities":"responsibilities",
                        "Escalation":"excalation",
                        "Success Criteria":"successCriteria",
                        f"{num} Prerequisites":"prerequisites",
                        f"{num} Pre-upgrade tasks":"preUpgradeTasks",
                        "Image upload":"imageUpload",
                        "Check Incompatibility of System Image":"checkIncompatibilityOfSystemImage",
                        "Identify the upgrade impact":"identifyTheUpgradeImpact",
                        "Change “admin” user password ":"userPassword",
                        "Hardware state verification":"hwStateVerification",
                        "Services and incident verification":"servicesandIncidentVerification",
                        "Pre-upgrade Checks capture":"preUpgradeChecksCapture",
                        f"{num} Traffic Migration to Alternate path":"In this scenario we will generally not be steering any traffic. The environment is a fabric path one with active use of VPC or dual homing (south-bound). This implies redundancy at all layers of the topology and ability to survive a temporary loss of a path with minimal to no service quality disruption.",
                        f"{num} Upgrade procedure":"upgradeProcedure",
                        f"{num} Rollback procedure":"rollbackProcedure",
                        f"{num} Post Upgrade Activities":"",
                        "Monitoring and customer verification":"Contact application owners and monitoring teams to verify operational state of the network. ",
                        "Diff report":"diffReport",
                        "Remove devices from Maintenance mode and sign off upgrade completion":f"{cust} or their representatives should remove all devices deemed to be successfully upgraded from maintenance mode in the monitoring systems. Communication should be sent to all involved parties informing about the upgrade completion and a sign off of the activity should be provided by {cust}.",
                        "Trademarks and Disclaimer":"TrademarksandDisclaimer",
                        "Document Acceptance":"DocumentAcceptance",
                        }
                        
        }
        prodH = headingwithcclass[f'{devicename} {version}']
        
        def new_lis(af, old_prod):
            new_prodH = {}
            for i, j in list(old_prod.items()):
                if i.find(af[0]) == -1:
                    new_prodH[i] = j
                else:
                    new_prodH[i] = j
                    new_prodH[af[1]] = af[2]
            return new_prodH
        headingwithclass = {}
        if lis: 
            for index,af in enumerate(lis):
                print(index)
                print(af)
                prodH = new_lis(af,prodH)
                print(prodH)
            headingwithclass = prodH
        else:
            headingwithclass = prodH
        s = Switcher()
        class_list = allProducts['class_list']
        count = 0
        c_list = []
        n = 0
        num = 0
        for i,j in headingwithclass.items():
            
            if j in c:
                c_list.append(i)      
            elif i[0].isdigit():
                n+=1
                count = 0
                c_list.append(i.replace('0',str(n)))
            else:
                count+=1
                c_list.append(f'{n}.{count} {i}')

        n = 0
        num = 0
        for i,j in headingwithclass.items():
            
            if j in c:
                ob1 = Dynamic.str_to_class(j)()
                method = getattr(ob1,j[0].lower()+j[1:],lambda :'Invalid')
                if j == 'Contents':
                    method(c_list)
                elif j == 'About':
                    document.add_heading(i,1)
                    method(author)
                elif j == 'TrademarksandDisclaimer':
                    document.add_page_break()
                    document.add_heading(i,1)
                    method()
                elif j == 'DocumentAcceptance':
                    document.add_heading(i,1)
                    method()
                else:
                    document.add_heading(i,2)
                    method()
            elif any(j.lower() in c.lower() for c in class_list):
                if i[0].isdigit():
                    n+=1
                    count= 0
                    document.add_heading(i.replace('0',str(n)),1)
                    s.indirect(devicename,os,j,version,rd,lr)
                else:
                    count+=1
                    document.add_heading(f'{n}.{count} {i}',2)
                    s.indirect(devicename,os,j,version,rd,lr)
            else:
                if i[0].isdigit():
                    n+=1
                    count= 0
                    document.add_heading(i.replace('0',str(n)),1)
                    document.add_paragraph(f"{j}") 
                else:
                    count+=1
                    document.add_heading(f'{n}.{count} {i}',2)
                    document.add_paragraph(f"{j}")
        return c_list

from lxml import etree
#main class
class Device(Formatter):

    def __init__(self, customer, author,dname='anonymous', pid="",sourceVer='16.09.06', targetVer='17.03.03', releaseDate=00-00-00,lis="",rd="",lr=""):
        self.customer = customer
        if author=="":
            self.author="<Author_Name>"
        else:
            self.author = author
        self.dname = dname
        self.pid = pid
        self.sourceVer = sourceVer
        self.targetVer = targetVer
        if releaseDate=="":
            self.releaseDate="<Release_Date>"
        else:
            self.releaseDate = releaseDate
        self.lis = lis
        self.rd = rd
        self.lr = lr
        
    def look_for_platform_version(self):
        platforms = allProducts['Platforms']
    
        if self.dname in platforms.keys():
            dt = platforms[self.dname]
            for item in dt:
                if item["version"] == self.targetVer:
                    return item
            raise SystemExit("Version does not exist")
        else:
            print("Device Does not exist in out database")

    def fetchMainData(self):
        exactP = self.look_for_platform_version()
        devicename = exactP['Devicename']
        os = exactP['os']
        version = exactP['version']
        typeofdevice = exactP['typeofdevice']
        try:
            document.add_picture('a2.jpg', width=Inches(1.0))
            document.add_picture('a1.jpg', width=Inches(6.0))
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(f"{e}(line {exc_tb.tb_lineno}): Image not Found...") 
        paragraph = document.add_paragraph(f'{self.customer} \n Software Upgrade Procedures & Pre-Checks \n MOP Document for {devicename}\n {self.releaseDate} \n Version 1.1', style='TOC Heading')
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.paragraph_format.space_before = Inches(0.3)

        paragraph = document.add_paragraph("Cisco Systems, Inc. "
                                           "\nCorporate Headquarters "
                                           "\n170 West Tasman Drive"
                                           "\nSan Jose, CA 95134-1706 USA "
                                           "\nhttp://www.cisco.com"
                                           "\nTel: 408 526-4000 Toll Free: 800 553-NETS (6387)"
                                           "\nFax: 408 526-4100 ")
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph.paragraph_format.space_before = Inches(0.5)
        section = document.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0].add_run()
        structure_footer = "_________________________________________________________________________________________________________"+f"{self.releaseDate} \t\t {self.dname} {os} {version} \nCisco Highly Confidential. All printed copies and duplicate soft copies are considered uncontrolled and the original online version should be referred to for the latest version.\n Page "
        footer_para.text = structure_footer
        page = Device.add_page_number(document.sections[0].footer.paragraphs[0].add_run())
        document.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.sections[0].different_first_page_header_footer = True
        sectPr = document.sections[0]._sectPr
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(ns.qn('w:start'), "0")
        sectPr.append(pgNumType)
        #document.add_page_break()
        d = Dynamic()
        d.creation_Of_heading_with_content(self.customer,self.author,devicename,pid,os,version,typeofdevice,self.lis,self.rd,self.lr)
        
        document.save(f'{self.customer}_{devicename}_{version}_mop.docx')

        return f"{self.customer}_{devicename}_{version}_mop.docx"

cust = input("Enter Your Name:")
if cust=='':
    cust='Customer_name'
else:
    pass
author = input("Enter author Name:")
Dname = input("Enter Device name:")
pid = input("Enter Product Id:")
sv = input("Enter Source version:")
tv = input("Enter Target version:")
date = input("Enter new version release date:")
lis = []
f = True
count=0
while(f):
    count +=1 
    ext_h = input(f'Enter {count} Heading:')
    if ext_h != "":
        af = input('Enter after which heading so u want this:')
        ext_c = input('Enter content of this heading:')
        lis.append([af,ext_h,ext_c])
    else:
        f = False
#for Related Documents
f = True
count = 0
rd = []
while(f):
    count +=1 
    ext_h = input(f'Enter {count} related Document part(if any):')
    if ext_h != "":
        link = input('Enter link for this head')
        rd.append([ext_h,link])
    else:
        f = False

#For limitations&restriction
lr = []
ext_h = input(f'Enter Limitation & Restrictions part(seperate points with commas):')
if ext_h != "":
    lr = ext_h.split(';')
else:
    pass

ob1 = Device(cust, author,Dname, pid,sv, tv, date, lis,rd,lr)
file_name = ob1.fetchMainData()



def set_updatefields_true(docx_path):
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    doc = Document(docx_path)
    # add child to doc.settings element
    element_updatefields = etree.SubElement(
        doc.settings.element, f"{namespace}updateFields"
    )
    element_updatefields.set(f"{namespace}val", "true")
    doc.save(docx_path)## Heading ##

script_dir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
file_path = os.path.join(script_dir, file_name)
set_updatefields_true(file_path)


    



